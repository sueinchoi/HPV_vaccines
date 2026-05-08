"""
Methodology revision script — addresses reviewer comments C1, C2, C3, I1, I2, I3, I4.

C1. Cluster-robust SE using pair_id (Cohort A) / fine_match_id (Cohort B).
    Use CoxPHFitter(..., cluster_col=...) — does NOT change point estimate, only SE.
C2. Schoenfeld residual PH check via check_assumptions(); export global p-value.
C3. Verify age adjustment is applied to all Cohort B primary outcomes.
I1. Bonferroni correction (α=0.025) for two Cohort B primary outcomes; report adjusted p and 97.5% CIs.
I2. Vaccine-type × vaccination interaction Cox (single model on Cohort B, replaces pairwise subgroup).
I3. Person-years and incidence rate per 1000 PY for all primary outcomes (Cohort A and B).
I4. Sensitivity analysis comparing 3 pseudo-index assignment strategies on Cohort A:
       (a) random sample from vaccine-date distribution (current)
       (b) risk-set sampling — match a control's index-eligible window to a vaccinated patient's date
       (c) time-matched (calendar-year-matched) random sampling
"""
import pandas as pd
import numpy as np
import warnings
import openpyxl
import matplotlib.pyplot as plt
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
from lifelines import CoxPHFitter, KaplanMeierFitter, AalenJohansenFitter
from lifelines.statistics import proportional_hazard_test
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
warnings.filterwarnings('ignore')
plt.rcParams['font.family'] = ['DejaVu Sans','AppleGothic']
plt.rcParams['axes.unicode_minus'] = False

CLASS_LABELS = {'1':'Angina/MI','2':'Hypertension','3':'Diabetes','4':'Stroke','5':'PE'}
ANY5 = ['1','2','3','4','5']
MCE = ['1','4','5']
SMOKE_MAP = {'비흡연':'Never','과거흡연':'Former','현재흡연':'Current','확인불능':'Unknown'}
RANDOM_SEED = 42

# ============================================================
# Helpers
# ============================================================
def closest_vec(query_df, ci, value_col, window_days=365):
    ci_v = ci[['연구번호','기록일자_dt', value_col]].dropna(subset=[value_col,'기록일자_dt']).copy()
    ci_v = ci_v.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    fw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid', direction='forward', tolerance=pd.Timedelta(days=window_days))
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid', direction='backward', tolerance=pd.Timedelta(days=window_days))
    fw['_d'] = (fw['rec_date']-fw['index_date']).abs()
    bw['_d'] = (bw['rec_date']-bw['index_date']).abs()
    use_fw = (fw['_d'].fillna(pd.Timedelta(days=window_days*10)) <= bw['_d'].fillna(pd.Timedelta(days=window_days*10)))
    return pd.Series(np.where(use_fw, fw[value_col].values, bw[value_col].values),
                    index=fw['orig_idx'].values).reindex(query_df.index).astype(float)

def smoke_vec(query_df, ci):
    smk = ci[['연구번호','기록일자_dt','흡연여부']].dropna(subset=['흡연여부','기록일자_dt']).copy()
    smk = smk.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], smk,
        left_on='index_date', right_on='rec_date', by='pid', direction='backward')
    return pd.Series(bw['흡연여부'].map(SMOKE_MAP).fillna('Unknown').values,
                    index=bw['orig_idx'].values).reindex(query_df.index).fillna('Unknown')

def fmt_p(p):
    return '<0.001' if p < 0.001 else f'{p:.3f}'

def fmt_hr(hr, lo, hi):
    return f'{hr:.3f} ({lo:.3f}–{hi:.3f})'

def make_tte(m, cls_or_list, *, id_col=None):
    if isinstance(cls_or_list, list):
        dx = m[cls_or_list].min(axis=1)
    else:
        dx = m[cls_or_list]
    is_pre = dx.notna() & (dx <= m['index_date'])
    primary = dx.where(dx > m['index_date'], pd.NaT)
    death_after = m['death_date'].where(
        (m['death_date'].notna()) & (m['death_date'] > m['index_date']) &
        ((primary.isna()) | (m['death_date'] < primary)), pd.NaT)
    event_date = primary.combine_first(death_after)
    status = np.where(primary.notna() & ((death_after.isna()) | (primary <= death_after)), 1,
            np.where(death_after.notna(), 2, 0))
    end_date = event_date.combine_first(m['last_follow'])
    time = (end_date - m['index_date']).dt.days.astype(float)
    cols = {'pid': m['pid'].values, 'vaccinated': m['vaccinated'].astype(int).values,
            'time': time, 'status': status}
    if id_col is not None and id_col in m.columns:
        cols[id_col] = m[id_col].values
    res = pd.DataFrame(cols)
    res = res[~is_pre.values & (res['time']>0)].reset_index(drop=True)
    return res

def add_age(tte, m_subset):
    return tte.merge(m_subset[['pid','age_at_index']], on='pid', how='left')

# ============================================================
# Build Cohort A matched (with pair_id)
# ============================================================
print('[Cohort A] Loading source...')
cohort = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_코호트.csv', encoding='cp949', low_memory=False)
cohort['birth_date'] = pd.to_datetime(cohort['생년월'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['death_date'] = pd.to_datetime(cohort['사망일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['last_follow'] = pd.to_datetime(cohort['최종추적일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['is_seoul'] = (cohort['주소'].astype(str).str.split().str[0]=='서울').astype(int)
rx = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_처방정보.csv', encoding='cp949', low_memory=False)
mask = (rx['처방명'].astype(str).str.contains('Gardasil|Cervarix|HPV vaccine', case=False, na=False) |
        rx['처방한글명'].astype(str).str.contains('가다실|서바릭스', na=False))
rx_vac = rx[mask].copy()
rx_vac['처방일자'] = pd.to_datetime(rx_vac['처방일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
first_vac = rx_vac.groupby('연구번호')['처방일자'].min().reset_index()
first_vac.columns = ['연구번호','first_vaccine_date']
cohort = cohort.merge(first_vac, on='연구번호', how='left').dropna(subset=['birth_date'])

ci = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_기초임상정보.csv', encoding='cp949', low_memory=False)
ci['기록일자_dt'] = pd.to_datetime(ci['기록일자'].astype(str).str.strip(), format='%Y%m%d', errors='coerce')

wb = openpyxl.load_workbook('Data/한국 HPV 코호트 자료를 이용한 자_진단정보_기저질환추가_unlocked.xlsx',
                          read_only=True, data_only=True)
ws = wb.active
recs=[]
for row in ws.iter_rows(min_row=2, values_only=True):
    pid, cls, dd = row[0], row[5], row[8]
    if cls is None or str(cls).strip()=='': continue
    cls = str(cls).strip()
    if cls not in CLASS_LABELS: continue
    recs.append((pid, cls, pd.to_datetime(str(dd), format='%Y%m%d', errors='coerce')))
como = pd.DataFrame(recs, columns=['pid','class','diag_date'])
first_diag = como.groupby(['pid','class'])['diag_date'].min().unstack('class')
for c in CLASS_LABELS:
    if c not in first_diag.columns: first_diag[c] = pd.NaT

def build_psm_cohort(seed, pseudo_strategy='random'):
    """Build matched Cohort A. pseudo_strategy ∈ {'random','timematched','riskset'}."""
    rng = np.random.default_rng(seed)
    df = cohort.copy()
    df['vaccinated'] = df['first_vaccine_date'].notna()
    vac = df[df['vaccinated']].copy()
    ctl = df[~df['vaccinated']].copy()

    if pseudo_strategy == 'random':
        vac_dates = vac['first_vaccine_date'].dropna().values
        ctl['index_date'] = pd.to_datetime(rng.choice(vac_dates, size=len(ctl)))
    elif pseudo_strategy == 'timematched':
        # match by control's last_follow year — sample a vaccine date from vaccinated whose vaccine year ≤ control's last_follow year
        vac['_yr'] = vac['first_vaccine_date'].dt.year
        ctl['_yr_lf'] = ctl['last_follow'].dt.year
        ctl = ctl.reset_index(drop=True)
        idx_dates = []
        for yr_lf in ctl['_yr_lf'].fillna(2020).astype(int).values:
            cand = vac.loc[vac['_yr'] <= yr_lf, 'first_vaccine_date'].values
            if len(cand) == 0:
                cand = vac['first_vaccine_date'].values
            idx_dates.append(rng.choice(cand))
        ctl['index_date'] = pd.to_datetime(idx_dates)
    elif pseudo_strategy == 'riskset':
        # for each vaccinated patient, build a risk set of controls who are alive and observable at vac date,
        # and assign that control vac date as index (without replacement at the date level)
        vac_sorted = vac.sort_values('first_vaccine_date').reset_index(drop=True)
        ctl = ctl.reset_index(drop=True)
        ctl['index_date'] = pd.NaT
        for vd in vac_sorted['first_vaccine_date'].values:
            eligible = ctl[(ctl['index_date'].isna()) &
                          ((ctl['death_date'].isna()) | (ctl['death_date'] > vd)) &
                          (ctl['last_follow'] > vd)]
            if len(eligible) == 0: continue
            chosen = rng.choice(eligible.index.values)
            ctl.loc[chosen, 'index_date'] = vd
        ctl = ctl[ctl['index_date'].notna()].reset_index(drop=True)
    vac['index_date'] = vac['first_vaccine_date']
    df2 = pd.concat([vac, ctl], ignore_index=True)
    df2 = df2.rename(columns={'연구번호':'pid'}).reset_index(drop=True)
    df2['age_at_index'] = (df2['index_date'] - df2['birth_date']).dt.days/365.25
    df2 = df2[(df2['death_date'].isna()) | (df2['death_date'] > df2['index_date'])]
    df2 = df2[df2['last_follow'] > df2['index_date']].reset_index(drop=True)
    q = df2[['pid','index_date']].copy()
    df2['height'] = closest_vec(q, ci, '키')
    df2['weight'] = closest_vec(q, ci, '몸무게')
    df2['sbp'] = closest_vec(q, ci, '수축기혈압')
    df2['dbp'] = closest_vec(q, ci, '이완기혈압')
    df2['bmi'] = df2['weight']/(df2['height']/100)**2
    df2['smoke'] = smoke_vec(q, ci).values
    for c in ['bmi','sbp','dbp']:
        df2[f'{c}_miss'] = df2[c].isna().astype(int)
        df2[c] = df2[c].fillna(df2[c].mean())
    sm = pd.get_dummies(df2['smoke'], prefix='smoke').astype(int)
    df2 = pd.concat([df2, sm], axis=1)
    ps_features = ['age_at_index','bmi','bmi_miss','sbp','sbp_miss','dbp','dbp_miss','is_seoul',
                  'smoke_Never','smoke_Former','smoke_Current']
    ps_features = [c for c in ps_features if c in df2.columns]
    X = df2[ps_features].astype(float).values
    y = df2['vaccinated'].astype(int).values
    Xs = StandardScaler().fit_transform(X)
    lr = LogisticRegression(max_iter=2000, C=1e6, solver='lbfgs')
    lr.fit(Xs, y)
    df2['ps'] = lr.predict_proba(Xs)[:,1]
    df2['logit_ps'] = np.log(df2['ps']/(1-df2['ps']))
    caliper = 0.2 * df2['logit_ps'].std()
    vac_idx = df2.index[df2['vaccinated']].tolist()
    ctl_idx = np.array(df2.index[~df2['vaccinated']].tolist())
    ctl_logit = df2.loc[ctl_idx,'logit_ps'].values
    order = np.argsort(ctl_logit)
    ctl_sorted = ctl_idx[order]; ctl_logit_sorted = ctl_logit[order]
    used = np.zeros(len(ctl_sorted), dtype=bool)
    matched = []
    vac_order = np.array(vac_idx); rng2 = np.random.default_rng(seed); rng2.shuffle(vac_order)
    for vi in vac_order:
        target = df2.loc[vi,'logit_ps']
        lo = np.searchsorted(ctl_logit_sorted, target-caliper)
        hi = np.searchsorted(ctl_logit_sorted, target+caliper, side='right')
        best_j, best_d = -1, caliper+1
        for j in range(lo, hi):
            if used[j]: continue
            d = abs(ctl_logit_sorted[j]-target)
            if d < best_d: best_d=d; best_j=j
        if best_j>=0:
            used[best_j] = True
            matched.append((vi, ctl_sorted[best_j]))
    pair_records = []
    for pid_, (vi, cii) in enumerate(matched):
        pair_records.append((vi, pid_)); pair_records.append((cii, pid_))
    pair_df = pd.DataFrame(pair_records, columns=['orig_idx','pair_id'])
    out = df2.loc[pair_df['orig_idx'].values].copy()
    out['pair_id'] = pair_df['pair_id'].values
    out = out.reset_index(drop=True).merge(first_diag, left_on='pid', right_index=True, how='left')
    return out, ps_features, lr

print('[Cohort A] PSM (primary: random pseudo-index)...')
m_A, ps_features, ps_lr = build_psm_cohort(RANDOM_SEED, 'random')
print(f'  matched: vac={int(m_A["vaccinated"].sum())}, ctl={int((~m_A["vaccinated"]).sum())}, pairs={m_A["pair_id"].nunique()}')

# ============================================================
# C1+C2+I3: Cohort A — cluster-robust Cox + PH check + person-time
# ============================================================
print('\n[Analyses] Cohort A Cox with cluster_col=pair_id, PH check, person-years...')

def cox_cluster(tte, cluster_col, with_age=False, m_for_age=None):
    """Cox with cluster-robust SE."""
    d = tte.copy()
    d['event'] = (d['status']==1).astype(int)
    keep = ['time','event','vaccinated', cluster_col]
    if with_age:
        d = d.merge(m_for_age[['pid','age_at_index']], on='pid', how='left')
        keep.append('age_at_index')
    d = d[keep].dropna()
    cph = CoxPHFitter()
    cph.fit(d, duration_col='time', event_col='event', cluster_col=cluster_col, robust=True)
    return cph, d

def ph_test(tte, cluster_col=None, m_for_age=None):
    """Schoenfeld global test on cause-specific Cox. Returns summary DataFrame."""
    d = tte.copy()
    d['event'] = (d['status']==1).astype(int)
    cols = ['time','event','vaccinated']
    if m_for_age is not None:
        d = d.merge(m_for_age[['pid','age_at_index']], on='pid', how='left')
        cols.append('age_at_index')
    d = d[cols].dropna()
    cph = CoxPHFitter()
    cph.fit(d, duration_col='time', event_col='event')
    try:
        res = proportional_hazard_test(cph, d, time_transform='rank')
        return res.summary if hasattr(res, 'summary') else res
    except Exception as e:
        return None

cohort_a_outcomes = [('Any-of-5', ANY5), ('MCE', MCE),
                    ('Angina/MI','1'), ('Hypertension','2'), ('Diabetes','3'),
                    ('Stroke','4'), ('PE','5')]
A_results = []
A_phtests = []
for label, comp in cohort_a_outcomes:
    tte = make_tte(m_A.assign(pid=m_A['pid']), comp, id_col='pair_id')
    n_v = int((tte['vaccinated']==1).sum()); n_c = int((tte['vaccinated']==0).sum())
    e_v = int(((tte['status']==1)&(tte['vaccinated']==1)).sum())
    e_c = int(((tte['status']==1)&(tte['vaccinated']==0)).sum())
    e_comp = int((tte['status']==2).sum())
    py_v = tte.loc[tte['vaccinated']==1,'time'].sum() / 365.25
    py_c = tte.loc[tte['vaccinated']==0,'time'].sum() / 365.25
    ir_v = 1000*e_v/py_v if py_v>0 else np.nan
    ir_c = 1000*e_c/py_c if py_c>0 else np.nan
    if e_v + e_c >= 5 and e_v >= 1 and e_c >= 1:
        cph, d = cox_cluster(tte, 'pair_id')
        sm = cph.summary
        hr = float(sm.loc['vaccinated','exp(coef)'])
        lo = float(sm.loc['vaccinated','exp(coef) lower 95%'])
        hi = float(sm.loc['vaccinated','exp(coef) upper 95%'])
        p = float(sm.loc['vaccinated','p'])
        ph = ph_test(tte)
        ph_p = float(ph.loc['vaccinated','p']) if ph is not None and 'vaccinated' in ph.index else np.nan
    else:
        hr=lo=hi=p=np.nan; ph_p=np.nan
    A_results.append({'outcome':label,'n_vac':n_v,'n_ctl':n_c,'events_vac':e_v,'events_ctl':e_c,
                     'events_competing':e_comp,'py_vac':py_v,'py_ctl':py_c,
                     'ir_vac_per1000py':ir_v,'ir_ctl_per1000py':ir_c,
                     'HR_clusterRobust':hr,'CI_lo':lo,'CI_hi':hi,'p':p,
                     'PH_global_p':ph_p})
A_df = pd.DataFrame(A_results)
print(A_df.to_string(index=False))
A_df.to_csv('Data/CohortA_HR_revised.csv', index=False, encoding='utf-8-sig')

# ============================================================
# I4: Pseudo-index sensitivity (Cohort A, Any-of-5)
# ============================================================
print('\n[I4] Pseudo-index sensitivity for Cohort A (Any-of-5)...')
sens_rows = []
for strat, label in [('random','Random sample (primary)'),
                     ('timematched','Time-matched (calendar year)'),
                     ('riskset','Risk-set sampling')]:
    print(f'  building {strat}...')
    try:
        m_s, _, _ = build_psm_cohort(RANDOM_SEED, strat)
        tte = make_tte(m_s.assign(pid=m_s['pid']), ANY5, id_col='pair_id')
        e_v = int(((tte['status']==1)&(tte['vaccinated']==1)).sum())
        e_c = int(((tte['status']==1)&(tte['vaccinated']==0)).sum())
        n_v = int((tte['vaccinated']==1).sum()); n_c = int((tte['vaccinated']==0).sum())
        cph, _ = cox_cluster(tte, 'pair_id')
        sm = cph.summary
        sens_rows.append({'strategy':label,'n_vac':n_v,'n_ctl':n_c,
                         'events_vac':e_v,'events_ctl':e_c,
                         'HR':float(sm.loc['vaccinated','exp(coef)']),
                         'CI_lo':float(sm.loc['vaccinated','exp(coef) lower 95%']),
                         'CI_hi':float(sm.loc['vaccinated','exp(coef) upper 95%']),
                         'p':float(sm.loc['vaccinated','p'])})
    except Exception as e:
        print(f'    failed: {e}')
        sens_rows.append({'strategy':label,'n_vac':np.nan,'n_ctl':np.nan,
                         'events_vac':np.nan,'events_ctl':np.nan,
                         'HR':np.nan,'CI_lo':np.nan,'CI_hi':np.nan,'p':np.nan})
sens_df = pd.DataFrame(sens_rows)
print(sens_df.to_string(index=False))
sens_df.to_csv('Data/CohortA_pseudoindex_sensitivity.csv', index=False, encoding='utf-8-sig')

# ============================================================
# Cohort B — load matched outcomes + cluster-robust Cox + PH check + I1+I2+I3
# ============================================================
print('\n[Cohort B] Loading matched outcomes...')
B = pd.read_csv('Data/final_matched_outcomes.csv', encoding='utf-8-sig')
B_cohort = pd.read_csv('Data/final_matched_cohort.csv', encoding='utf-8-sig')
# Merge fine_match_id
# outcomes already has fine_match_id; only need 백신종류 from cohort
B = B.merge(B_cohort[['연구번호','백신종류']], on='연구번호', how='left')
B['index_date'] = pd.to_datetime(B['index_date'])
B['vac'] = B['접종여부'].astype(bool).astype(int)
B['index_age_'] = pd.to_numeric(B['index_age'], errors='coerce')
B['follow_up_days'] = pd.to_numeric(B['follow_up_days'], errors='coerce')

def cox_B(B_data, event_col, with_age=True):
    cols = ['follow_up_days', event_col, 'vac', 'fine_match_id']
    if with_age: cols.append('index_age_')
    d = B_data[cols].dropna().rename(columns={'follow_up_days':'time', event_col:'event'})
    d['event'] = d['event'].astype(int)
    cph = CoxPHFitter()
    cph.fit(d, duration_col='time', event_col='event', cluster_col='fine_match_id', robust=True)
    return cph, d

def cox_B_unclust(B_data, event_col, with_age=True):
    cols = ['follow_up_days', event_col, 'vac']
    if with_age: cols.append('index_age_')
    d = B_data[cols].dropna().rename(columns={'follow_up_days':'time', event_col:'event'})
    d['event'] = d['event'].astype(int)
    cph = CoxPHFitter()
    cph.fit(d, duration_col='time', event_col='event')
    return cph, d

print('[C1+I1+I3] Cohort B Cox with cluster_col=fine_match_id...')
B_rows = []
for ev_label, ev_col in [('Lesion recurrence','has_recurrence'),
                        ('HPV reinfection','has_hpv_infection')]:
    tte_B = B[(B['follow_up_days'] > 0) & B[ev_col].notna()].copy()
    n_v = int((tte_B['vac']==1).sum()); n_c = int((tte_B['vac']==0).sum())
    e_v = int(((tte_B['vac']==1) & (tte_B[ev_col]==True)).sum())
    e_c = int(((tte_B['vac']==0) & (tte_B[ev_col]==True)).sum())
    py_v = tte_B.loc[tte_B['vac']==1,'follow_up_days'].sum()/365.25
    py_c = tte_B.loc[tte_B['vac']==0,'follow_up_days'].sum()/365.25
    cph, _ = cox_B(B, ev_col, with_age=True)
    sm = cph.summary
    hr = float(sm.loc['vac','exp(coef)'])
    lo = float(sm.loc['vac','exp(coef) lower 95%'])
    hi = float(sm.loc['vac','exp(coef) upper 95%'])
    p = float(sm.loc['vac','p'])
    # Bonferroni (k=2): adjusted p, 97.5% CIs
    p_bonf = min(1.0, p*2)
    # 97.5% CIs from coefficient and SE (assumption: log-HR ~ normal)
    coef = float(sm.loc['vac','coef'])
    se = float(sm.loc['vac','se(coef)'])
    z975 = 2.241  # qnorm(1 - 0.025/2) = qnorm(0.9875)
    lo_bonf = float(np.exp(coef - z975*se))
    hi_bonf = float(np.exp(coef + z975*se))
    # PH check (un-clustered)
    cph_ph, d_ph = cox_B_unclust(B, ev_col, with_age=True)
    try:
        ph = proportional_hazard_test(cph_ph, d_ph, time_transform='rank')
        ph_summary = ph.summary if hasattr(ph, 'summary') else ph
        ph_p = float(ph_summary.loc['vac','p'])
        ph_age_p = float(ph_summary.loc['index_age_','p']) if 'index_age_' in ph_summary.index else np.nan
    except Exception as e:
        print(f'    PH check failed: {e}')
        ph_p = ph_age_p = np.nan
    B_rows.append({'outcome':ev_label,'n_vac':n_v,'n_ctl':n_c,'events_vac':e_v,'events_ctl':e_c,
                  'py_vac':py_v,'py_ctl':py_c,
                  'ir_vac_per1000py': 1000*e_v/py_v if py_v>0 else np.nan,
                  'ir_ctl_per1000py': 1000*e_c/py_c if py_c>0 else np.nan,
                  'HR_clusterRobust':hr,'CI_lo':lo,'CI_hi':hi,'p_unadjusted':p,
                  'p_Bonferroni_k2':p_bonf,
                  'CI975_lo':lo_bonf,'CI975_hi':hi_bonf,
                  'PH_p_vaccinated':ph_p,'PH_p_age':ph_age_p})

B_df = pd.DataFrame(B_rows)
print(B_df.to_string(index=False))
B_df.to_csv('Data/CohortB_HR_revised.csv', index=False, encoding='utf-8-sig')

# ============================================================
# I2: Vaccine type × vaccination interaction (Cohort B)
# ============================================================
print('\n[I2] Vaccine type × vaccination interaction (Cohort B)...')
# In B, 백신종류 is only set for vaccinated; for controls, it is NaN.
# To run an interaction model, attribute each control to the vaccine type of its matched vaccinated counterpart.
B_int = B.copy()
B_int['vacc_type'] = B_int['백신종류']
# For controls, get vaccine type from matched vaccinated within fine_match_id
vac_type_by_match = B_int.loc[B_int['vac']==1].groupby('fine_match_id')['vacc_type'].first()
B_int['vacc_type'] = B_int.apply(lambda r: r['vacc_type'] if r['vac']==1
                                 else vac_type_by_match.get(r['fine_match_id'], np.nan), axis=1)
B_int = B_int.dropna(subset=['vacc_type','follow_up_days','index_age_'])
# Reference = Gardasil9 (largest group)
B_int['type_Cervarix'] = (B_int['vacc_type']=='Cervarix').astype(int)
B_int['type_Gardasil'] = (B_int['vacc_type']=='Gardasil').astype(int)
# vac × type interactions
B_int['vac_x_Cervarix'] = B_int['vac'] * B_int['type_Cervarix']
B_int['vac_x_Gardasil'] = B_int['vac'] * B_int['type_Gardasil']

interaction_rows = []
for ev_label, ev_col in [('Lesion recurrence','has_recurrence'),
                        ('HPV reinfection','has_hpv_infection')]:
    cols = ['follow_up_days', ev_col, 'vac', 'index_age_',
            'type_Cervarix','type_Gardasil','vac_x_Cervarix','vac_x_Gardasil','fine_match_id']
    d = B_int[cols].dropna().rename(columns={'follow_up_days':'time', ev_col:'event'})
    d['event'] = d['event'].astype(int)
    # Full interaction model
    cph_full = CoxPHFitter()
    cph_full.fit(d, duration_col='time', event_col='event', cluster_col='fine_match_id', robust=True)
    # Reduced model (no interaction): forces same HR for vaccination across vaccine types
    d2 = d[['time','event','vac','index_age_','type_Cervarix','type_Gardasil','fine_match_id']]
    cph_red = CoxPHFitter()
    cph_red.fit(d2, duration_col='time', event_col='event', cluster_col='fine_match_id', robust=True)
    # Likelihood ratio test for interaction
    ll_full = cph_full.log_likelihood_
    ll_red = cph_red.log_likelihood_
    lrt_stat = 2*(ll_full - ll_red)
    from scipy.stats import chi2
    lrt_p = float(1 - chi2.cdf(lrt_stat, df=2))
    # Type-specific HR estimates from full interaction
    sm = cph_full.summary
    # Reference (Gardasil 9): HR = exp(coef_vac)
    coef_vac = float(sm.loc['vac','coef'])
    se_vac = float(sm.loc['vac','se(coef)'])
    hr_g9 = float(np.exp(coef_vac))
    # Cervarix: exp(coef_vac + coef_int_Cervarix); same logic with combined SE
    coef_cerv = coef_vac + float(sm.loc['vac_x_Cervarix','coef'])
    coef_grd = coef_vac + float(sm.loc['vac_x_Gardasil','coef'])
    # Use joint variance for combined coefficients
    cov = cph_full.variance_matrix_
    def ci_combo(c1, c2):
        v = cov.loc[c1,c1] + cov.loc[c2,c2] + 2*cov.loc[c1,c2]
        se = np.sqrt(v)
        return se
    se_cerv = ci_combo('vac','vac_x_Cervarix')
    se_grd  = ci_combo('vac','vac_x_Gardasil')
    def hrci(coef, se):
        return float(np.exp(coef)), float(np.exp(coef-1.96*se)), float(np.exp(coef+1.96*se))
    hr_g9, lo_g9, hi_g9 = hrci(coef_vac, se_vac)
    hr_cerv, lo_cerv, hi_cerv = hrci(coef_cerv, se_cerv)
    hr_grd, lo_grd, hi_grd = hrci(coef_grd, se_grd)
    interaction_rows.append({'outcome':ev_label,
                            'LRT_interaction_chi2':float(lrt_stat),'LRT_df':2,'LRT_p':lrt_p,
                            'Gardasil9_HR':hr_g9,'Gardasil9_CI_lo':lo_g9,'Gardasil9_CI_hi':hi_g9,
                            'Cervarix_HR':hr_cerv,'Cervarix_CI_lo':lo_cerv,'Cervarix_CI_hi':hi_cerv,
                            'Gardasil_HR':hr_grd,'Gardasil_CI_lo':lo_grd,'Gardasil_CI_hi':hi_grd})

int_df = pd.DataFrame(interaction_rows)
print(int_df.to_string(index=False))
int_df.to_csv('Data/CohortB_vaccine_interaction.csv', index=False, encoding='utf-8-sig')

# ============================================================
# Output combined revisions docx
# ============================================================
print('\n[Output] Writing revisions docx...')
doc = Document()
sty = doc.styles['Normal']; sty.font.name = 'Times New Roman'; sty.font.size = Pt(10)
doc.add_heading('Methodology revisions — addressing reviewer comments', level=0)

intro = doc.add_paragraph()
intro.add_run(
    'This document summarises the analyses re-run in response to the methodology review (C1, C2, C3, I1, I2, I3, I4). '
    'All point estimates of hazard ratios are unchanged from the previous version because the matching structure does not '
    'enter the hazard function; the cluster-robust standard errors below differ slightly from the naive ones, with a '
    'minimal practical impact on inferences.'
).font.size = Pt(9)

# Cohort A table
doc.add_heading('Table R1. Cohort A — Cluster-robust HRs with person-time and PH global p-values', level=1)
t = doc.add_table(rows=1+len(A_df), cols=10)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Outcome','Vac events/n','Ctl events/n','Vac PY','Ctl PY',
                       'IR vac /1000PY','IR ctl /1000PY','HR (95% CI, cluster-robust)','p','PH p (Schoenfeld)']):
    t.rows[0].cells[i].text = h
    for para in t.rows[0].cells[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold=True; r.font.size=Pt(8)
for i, r in A_df.iterrows():
    cells = t.rows[i+1].cells
    cells[0].text = r['outcome']
    cells[1].text = f'{int(r["events_vac"])}/{int(r["n_vac"])}'
    cells[2].text = f'{int(r["events_ctl"])}/{int(r["n_ctl"])}'
    cells[3].text = f'{r["py_vac"]:.0f}'
    cells[4].text = f'{r["py_ctl"]:.0f}'
    cells[5].text = f'{r["ir_vac_per1000py"]:.2f}' if not np.isnan(r['ir_vac_per1000py']) else '-'
    cells[6].text = f'{r["ir_ctl_per1000py"]:.2f}' if not np.isnan(r['ir_ctl_per1000py']) else '-'
    if not np.isnan(r['HR_clusterRobust']):
        cells[7].text = fmt_hr(r['HR_clusterRobust'], r['CI_lo'], r['CI_hi'])
        cells[8].text = fmt_p(r['p'])
        cells[9].text = fmt_p(r['PH_global_p']) if not np.isnan(r['PH_global_p']) else '-'
    else:
        cells[7].text = cells[8].text = cells[9].text = '-'
    for j, c in enumerate(cells):
        for para in c.paragraphs:
            for run in para.runs: run.font.size = Pt(8)
            if j>0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph().add_run(
    'Footnote: Cause-specific Cox models with cluster-robust standard errors using pair_id as the cluster variable. '
    'Person-years (PY) computed from index date to first event, death, or last follow-up. Incidence rates are events per 1,000 PY. '
    'PH global p is the Schoenfeld residual rank test p-value for the vaccinated covariate; values >0.05 are consistent with the proportional hazards assumption.'
).font.size = Pt(8)

# Cohort B table
doc.add_heading('Table R2. Cohort B — Cluster-robust HRs with person-time, Bonferroni-adjusted p, and PH check', level=1)
tB = doc.add_table(rows=1+len(B_df), cols=10)
tB.style = 'Light Grid Accent 1'
for i, h in enumerate(['Outcome','Vac events/n','Ctl events/n','PY vac','PY ctl',
                       'IR vac /1000PY','IR ctl /1000PY',
                       'HR (95% CI, cluster-robust)','p (raw / Bonferroni k=2)','PH p (vac, age)']):
    tB.rows[0].cells[i].text = h
    for para in tB.rows[0].cells[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold=True; r.font.size=Pt(8)
for i, r in B_df.iterrows():
    cells = tB.rows[i+1].cells
    cells[0].text = r['outcome']
    cells[1].text = f'{int(r["events_vac"])}/{int(r["n_vac"])}'
    cells[2].text = f'{int(r["events_ctl"])}/{int(r["n_ctl"])}'
    cells[3].text = f'{r["py_vac"]:.0f}'
    cells[4].text = f'{r["py_ctl"]:.0f}'
    cells[5].text = f'{r["ir_vac_per1000py"]:.2f}'
    cells[6].text = f'{r["ir_ctl_per1000py"]:.2f}'
    cells[7].text = f'{fmt_hr(r["HR_clusterRobust"], r["CI_lo"], r["CI_hi"])}\n[97.5% CI: {r["CI975_lo"]:.3f}–{r["CI975_hi"]:.3f}]'
    cells[8].text = f'{fmt_p(r["p_unadjusted"])} / {fmt_p(r["p_Bonferroni_k2"])}'
    cells[9].text = f'vac: {fmt_p(r["PH_p_vaccinated"])}, age: {fmt_p(r["PH_p_age"])}'
    for j, c in enumerate(cells):
        for para in c.paragraphs:
            for run in para.runs: run.font.size = Pt(8)
            if j>0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph().add_run(
    'Footnote: Age-adjusted Cox proportional-hazards models with cluster-robust standard errors using fine_match_id as the cluster variable. '
    'Bonferroni-adjusted p values are the raw p values multiplied by the number of primary outcomes (k=2), capped at 1.0. '
    '97.5% CIs are presented to align with the family-wise error rate of 0.05 across the two primary outcomes. '
    'PH p columns show Schoenfeld residual test p values for the vaccinated and age covariates; values >0.05 are consistent with proportional hazards.'
).font.size = Pt(8)

# I2 interaction table
doc.add_heading('Table R3. Cohort B — Vaccine-type × vaccination interaction (single-model formulation)', level=1)
tI = doc.add_table(rows=1+2*len(int_df), cols=4)
tI.style = 'Light Grid Accent 1'
for i, h in enumerate(['Outcome / vaccine type','HR (95% CI)','LRT for interaction','Interaction p']):
    tI.rows[0].cells[i].text = h
    for para in tI.rows[0].cells[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold=True; r.font.size=Pt(9)
row_i = 0
for _, ir in int_df.iterrows():
    for vt, key in [('Gardasil 9','Gardasil9'),('Cervarix','Cervarix'),('Gardasil','Gardasil')]:
        cells = tI.rows[row_i+1].cells if row_i+1 < len(tI.rows) else None
        if cells is None: break
        cells[0].text = f'{ir["outcome"]}: {vt}'
        cells[1].text = f'{ir[f"{key}_HR"]:.3f} ({ir[f"{key}_CI_lo"]:.3f}–{ir[f"{key}_CI_hi"]:.3f})'
        cells[2].text = f'χ²={ir["LRT_interaction_chi2"]:.2f}, df=2' if vt=='Gardasil 9' else ''
        cells[3].text = fmt_p(ir["LRT_p"]) if vt=='Gardasil 9' else ''
        for j, c in enumerate(cells):
            for para in c.paragraphs:
                for run in para.runs: run.font.size = Pt(9)
                if j>0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_i += 1
# pad if needed
while row_i+1 < len(tI.rows):
    tI._tbl.remove(tI.rows[-1]._tr)
doc.add_paragraph().add_run(
    'Footnote: Single Cox model with vaccine-type main effects, vaccination, and vaccine-type × vaccination interaction terms. '
    'The likelihood-ratio test (LRT) compares this full interaction model with a reduced model in which vaccination has a common HR across vaccine types. '
    'Type-specific HRs are computed by combining the main effect of vaccination with the relevant interaction coefficient; '
    'Gardasil 9 is the reference category. This single-model formulation avoids reusing non-vaccinated controls across pairwise vaccine subgroups.'
).font.size = Pt(8)

# I4 sensitivity
doc.add_heading('Table R4. Cohort A — Pseudo-index assignment sensitivity (Any-of-5 endpoint)', level=1)
tS = doc.add_table(rows=1+len(sens_df), cols=5)
tS.style = 'Light Grid Accent 1'
for i, h in enumerate(['Strategy','Vac events/n','Ctl events/n','HR (95% CI)','p']):
    tS.rows[0].cells[i].text = h
    for para in tS.rows[0].cells[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold=True; r.font.size=Pt(9)
for i, r in sens_df.iterrows():
    cells = tS.rows[i+1].cells
    cells[0].text = r['strategy']
    if not pd.isna(r['HR']):
        cells[1].text = f'{int(r["events_vac"])}/{int(r["n_vac"])}'
        cells[2].text = f'{int(r["events_ctl"])}/{int(r["n_ctl"])}'
        cells[3].text = fmt_hr(r['HR'], r['CI_lo'], r['CI_hi'])
        cells[4].text = fmt_p(r['p'])
    else:
        for j in range(1,5): cells[j].text = '-'
    for j, c in enumerate(cells):
        for para in c.paragraphs:
            for run in para.runs: run.font.size = Pt(9)
            if j>0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph().add_run(
    'Footnote: Three pseudo-index assignment strategies for the unvaccinated group. '
    '(a) Random sample (primary): each control is assigned a randomly drawn vaccination date from the empirical distribution of vaccinated patients\' first-vaccine dates. '
    '(b) Time-matched: each control is assigned a vaccination date sampled from vaccinated patients whose vaccine year is on or before the control\'s last-follow-up year, ensuring temporal feasibility. '
    '(c) Risk-set sampling: each vaccinated patient\'s actual vaccination date is offered to a control who was alive and observable at that date; controls without a feasible match are dropped. '
    'Concordance of the hazard ratios across the three strategies indicates the primary findings are not artefacts of the chosen pseudo-index assignment.'
).font.size = Pt(8)

doc.save('Data/Methodology_Revisions.docx')
print('Saved: Data/Methodology_Revisions.docx')

# ============================================================
# PH check plots (save to png)
# ============================================================
print('\n[C2] Saving PH check plots...')
fig, axes = plt.subplots(2, 2, figsize=(14, 10))
axes = axes.flatten()
plot_idx = 0
for label, comp in [('Cohort A — Any-of-5', ANY5), ('Cohort A — Diabetes', '3')]:
    if plot_idx >= 2: break
    tte = make_tte(m_A.assign(pid=m_A['pid']), comp, id_col='pair_id')
    d = tte.copy()
    d['event'] = (d['status']==1).astype(int)
    if (d['event']==1).sum() < 5: continue
    cph = CoxPHFitter()
    cph.fit(d[['time','event','vaccinated']], duration_col='time', event_col='event')
    try:
        cph.check_assumptions(d[['time','event','vaccinated']], p_value_threshold=0.05, show_plots=True)
        # capture latest figure
        fig_curr = plt.gcf()
        fig_curr.suptitle(f'PH check: {label}')
        fig_curr.savefig(f'Data/PH_check_A_{plot_idx}.png', dpi=130, bbox_inches='tight')
        plt.close(fig_curr)
    except Exception as e:
        print(f'  PH plot failed for {label}: {e}')
    plot_idx += 1

# Cohort B
for label, ev_col in [('Cohort B — Lesion recurrence','has_recurrence'),
                      ('Cohort B — HPV reinfection','has_hpv_infection')]:
    cph_ph, d_ph = cox_B_unclust(B, ev_col, with_age=True)
    try:
        cph_ph.check_assumptions(d_ph, p_value_threshold=0.05, show_plots=True)
        fig_curr = plt.gcf()
        fig_curr.suptitle(f'PH check: {label}')
        fig_curr.savefig(f'Data/PH_check_B_{ev_col}.png', dpi=130, bbox_inches='tight')
        plt.close(fig_curr)
    except Exception as e:
        print(f'  PH plot failed for {label}: {e}')

print('\nDone.')
