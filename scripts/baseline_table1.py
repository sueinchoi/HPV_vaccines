"""
Table 1. Baseline Characteristics — Pre/Post matching for Cohort A and Cohort B
출력: docx 표 + csv

Cohort A (Whole-cohort, manage 만성질환 분석)
  - Pre-match: 접종군 2,156 vs 비접종군 30,813
  - Post-match: 접종군 2,155 vs 비접종군 8,620

Cohort B (Surgical efficacy)
  - Pre-match: 수술환자 중 접종군 vs 비접종군 (모두)
  - Post-match: 접종군 241 vs 비접종군 867
"""
import pandas as pd
import numpy as np
from scipy.stats import fisher_exact, ttest_ind, chi2_contingency
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import openpyxl

# ------------------------ helpers ------------------------
def smd_cont(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a) < 2 or len(b) < 2: return np.nan
    pooled = np.sqrt((a.var(ddof=1) + b.var(ddof=1)) / 2)
    return (a.mean() - b.mean()) / pooled if pooled > 0 else np.nan

def smd_bin(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a) == 0 or len(b) == 0: return np.nan
    p1, p2 = a.mean(), b.mean()
    pooled = np.sqrt((p1*(1-p1) + p2*(1-p2)) / 2)
    return (p1 - p2) / pooled if pooled > 0 else np.nan

def fmt_cont(s):
    s = pd.Series(s).dropna()
    if len(s) == 0: return '-'
    return f'{s.mean():.2f} ± {s.std():.2f}'

def fmt_bin(s, total):
    n = int(pd.Series(s).fillna(0).astype(bool).sum())
    return f'{n} ({100*n/total:.1f}%)' if total > 0 else '-'

def p_cont(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a) < 2 or len(b) < 2: return np.nan
    return ttest_ind(a, b, equal_var=False).pvalue

def p_bin(a_yes, a_n, b_yes, b_n):
    if a_n == 0 or b_n == 0: return np.nan
    table = [[a_yes, a_n - a_yes], [b_yes, b_n - b_yes]]
    return fisher_exact(table)[1]

def fmt_p(p):
    if pd.isna(p): return '-'
    if p < 0.001: return '<0.001'
    return f'{p:.3f}'

def fmt_smd(s):
    if pd.isna(s): return '-'
    return f'{abs(s):.3f}'

# ------------------------ data load ------------------------
print('Loading source data...')
cohort = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_코호트.csv',
                    encoding='cp949', low_memory=False)
cohort['birth_date'] = pd.to_datetime(cohort['생년월'].astype('Int64').astype(str),
                                     format='%Y%m%d', errors='coerce')
cohort['death_date'] = pd.to_datetime(cohort['사망일자'].astype('Int64').astype(str),
                                     format='%Y%m%d', errors='coerce')
cohort['last_follow'] = pd.to_datetime(cohort['최종추적일자'].astype('Int64').astype(str),
                                       format='%Y%m%d', errors='coerce')
cohort['region'] = cohort['주소'].astype(str).str.split().str[0]

# Vaccinated
rx = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_처방정보.csv',
                encoding='cp949', low_memory=False)
mask = (rx['처방명'].astype(str).str.contains('Gardasil|Cervarix|HPV vaccine', case=False, na=False) |
        rx['처방한글명'].astype(str).str.contains('가다실|서바릭스', na=False))
rx_vac = rx[mask].copy()
rx_vac['처방일자'] = pd.to_datetime(rx_vac['처방일자'].astype('Int64').astype(str),
                                  format='%Y%m%d', errors='coerce')
first_vac = rx_vac.groupby('연구번호')['처방일자'].min().reset_index()
first_vac.columns = ['연구번호','first_vaccine_date']
cohort = cohort.merge(first_vac, on='연구번호', how='left')

# Comorbidities
print('Loading comorbidities...')
wb = openpyxl.load_workbook(
    'Data/한국 HPV 코호트 자료를 이용한 자_진단정보_기저질환추가_unlocked.xlsx',
    read_only=True, data_only=True)
ws = wb.active
recs = []
CLASS_LABELS = {'1':'협심증/심근경색', '2':'고혈압', '3':'당뇨', '4':'뇌출혈/뇌경색', '5':'폐색전증'}
for row in ws.iter_rows(min_row=2, values_only=True):
    pid, cls, dd = row[0], row[5], row[8]
    if cls is None or str(cls).strip() == '': continue
    cls = str(cls).strip()
    if cls not in CLASS_LABELS: continue
    d = pd.to_datetime(str(dd), format='%Y%m%d', errors='coerce')
    recs.append((pid, cls, d))
como = pd.DataFrame(recs, columns=['pid','class','diag_date'])
first_diag = como.groupby(['pid','class'])['diag_date'].min().unstack('class')
for c in CLASS_LABELS:
    if c not in first_diag.columns:
        first_diag[c] = pd.NaT

# Surgery
print('Loading surgery...')
surg = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_수술처방_수술종류구분완료.csv',
                  encoding='cp949', low_memory=False)
surg['수술처방일자'] = pd.to_datetime(surg['수술처방일자'].astype('Int64').astype(str),
                                  format='%Y%m%d', errors='coerce')
surg = surg[surg['수술 종류'].astype(str).isin(['1','3'])]
SURG_TYPE = {'1':'원추절제술', '3':'자궁절제술'}
surg['수술방법'] = surg['수술 종류'].astype(str).map(SURG_TYPE)
first_surg = surg.sort_values('수술처방일자').groupby('연구번호').first()[['수술처방일자','수술방법']].reset_index()
first_surg.columns = ['연구번호','first_surg_date','first_surg_type']
print(f'  surgery patients (excl. 제외): {len(first_surg)}')

# ------------------------ Cohort A ------------------------
print('\n=== Cohort A ===')
# Pre-match: full cohort split by vaccination
A_pre = cohort.copy()
A_pre['vaccinated'] = A_pre['first_vaccine_date'].notna()
# Post-match
A_post = pd.read_csv('Data/full_cohort_age_matched.csv', encoding='utf-8-sig')
A_post['index_date'] = pd.to_datetime(A_post['index_date'])
A_post = A_post.merge(cohort[['연구번호','birth_date','death_date','last_follow','region']],
                     left_on='pid', right_on='연구번호', how='left')
A_post['age_at_index'] = (A_post['index_date'] - A_post['birth_date']).dt.days/365.25
A_post['follow_up_days'] = (A_post['last_follow'] - A_post['index_date']).dt.days
A_post['died'] = A_post['death_date'].notna() & (A_post['death_date'] >= A_post['index_date'])

# Comorbidity columns (post-match: pre-index 진단)
A_post = A_post.merge(first_diag, left_on='pid', right_index=True, how='left')
for c in CLASS_LABELS:
    A_post[f'baseline_{c}'] = (A_post[c].notna()) & (A_post[c] <= A_post['index_date'])
A_post['baseline_any'] = A_post[[f'baseline_{c}' for c in CLASS_LABELS]].any(axis=1)

# For PRE-match (no index_date for controls): use any-time prevalence
A_pre_with_como = A_pre.merge(first_diag, left_on='연구번호', right_index=True, how='left')
for c in CLASS_LABELS:
    A_pre_with_como[f'lifetime_{c}'] = A_pre_with_como[c].notna()
A_pre_with_como['lifetime_any'] = A_pre_with_como[[f'lifetime_{c}' for c in CLASS_LABELS]].any(axis=1)
A_pre_with_como['birth_year'] = A_pre_with_como['birth_date'].dt.year
A_pre_with_como['follow_total_days'] = (A_pre_with_como['last_follow'] - A_pre_with_como['birth_date']).dt.days
A_pre_with_como['died_flag'] = A_pre_with_como['death_date'].notna()

# ------------------------ Cohort B ------------------------
print('=== Cohort B ===')
# Pre-match: surgery patients split by vaccination
B_pre = first_surg.merge(cohort[['연구번호','birth_date','death_date','last_follow','region','first_vaccine_date']], on='연구번호', how='left')
B_pre['vaccinated'] = B_pre['first_vaccine_date'].notna()
B_pre['birth_year'] = B_pre['birth_date'].dt.year
B_pre['surg_year'] = B_pre['first_surg_date'].dt.year
B_pre['age_at_surgery'] = (B_pre['first_surg_date'] - B_pre['birth_date']).dt.days/365.25
B_pre = B_pre.merge(first_diag, left_on='연구번호', right_index=True, how='left')
for c in CLASS_LABELS:
    B_pre[f'baseline_{c}'] = (B_pre[c].notna()) & (B_pre[c] <= B_pre['first_surg_date'])
B_pre['baseline_any'] = B_pre[[f'baseline_{c}' for c in CLASS_LABELS]].any(axis=1)
B_pre['died_flag'] = B_pre['death_date'].notna()

# Post-match Cohort B
B_post = pd.read_csv('Data/final_matched_cohort.csv', encoding='utf-8-sig')
B_post['index_date'] = pd.to_datetime(B_post['index_date'])
B_post['첫수술일자'] = pd.to_datetime(B_post['첫수술일자'])
B_post['follow_up_days'] = (pd.to_datetime(B_post['최종추적일자']) - B_post['index_date']).dt.days
B_post = B_post.merge(cohort[['연구번호','birth_date','region','death_date']], on='연구번호', how='left')
B_post['birth_year'] = B_post['birth_date'].dt.year
B_post['died'] = B_post['death_date'].notna() & (B_post['death_date'] >= B_post['index_date'])
B_post = B_post.merge(first_diag, left_on='연구번호', right_index=True, how='left')
for c in CLASS_LABELS:
    B_post[f'baseline_{c}'] = (B_post[c].notna()) & (B_post[c] <= B_post['index_date'])
B_post['baseline_any'] = B_post[[f'baseline_{c}' for c in CLASS_LABELS]].any(axis=1)

# ------------------------ Build rows ------------------------
def make_rows(df, vac_col, vars_cont, vars_bin, vars_cat=None):
    """vars_cont: list of (label, colname); vars_bin: list of (label, colname); vars_cat: list of (label, colname, categories)"""
    vac = df[df[vac_col]==True]
    ctl = df[df[vac_col]==False]
    nv, nc = len(vac), len(ctl)
    rows = [('N', f'{nv}', f'{nc}', '-', '-')]
    for label, col in vars_cont:
        if col not in df.columns:
            rows.append((label, '-', '-', '-', '-'))
            continue
        v = vac[col]; c = ctl[col]
        rows.append((label, fmt_cont(v), fmt_cont(c), fmt_p(p_cont(v, c)), fmt_smd(smd_cont(v, c))))
    for label, col in vars_bin:
        if col not in df.columns:
            rows.append((label, '-', '-', '-', '-'))
            continue
        v_yes = int(vac[col].fillna(0).astype(bool).sum())
        c_yes = int(ctl[col].fillna(0).astype(bool).sum())
        rows.append((label,
                     f'{v_yes} ({100*v_yes/nv:.1f}%)' if nv>0 else '-',
                     f'{c_yes} ({100*c_yes/nc:.1f}%)' if nc>0 else '-',
                     fmt_p(p_bin(v_yes, nv, c_yes, nc)),
                     fmt_smd(smd_bin(vac[col].astype(float), ctl[col].astype(float)))))
    if vars_cat:
        for label, col, cats in vars_cat:
            rows.append((label, '', '', '', ''))
            for cat in cats:
                v_yes = int((vac[col]==cat).sum()); c_yes = int((ctl[col]==cat).sum())
                rows.append((f'  {cat}',
                             f'{v_yes} ({100*v_yes/nv:.1f}%)' if nv>0 else '-',
                             f'{c_yes} ({100*c_yes/nc:.1f}%)' if nc>0 else '-',
                             fmt_p(p_bin(v_yes, nv, c_yes, nc)),
                             fmt_smd(smd_bin((vac[col]==cat).astype(float), (ctl[col]==cat).astype(float)))))
    return rows

# Cohort A pre/post variable lists
A_pre_cont = [('Birth year', 'birth_year')]
A_pre_bin = [('Mortality (ever)', 'died_flag'),
             ('협심증/심근경색 (lifetime)', 'lifetime_1'),
             ('고혈압 (lifetime)', 'lifetime_2'),
             ('당뇨 (lifetime)', 'lifetime_3'),
             ('뇌출혈/뇌경색 (lifetime)', 'lifetime_4'),
             ('폐색전증 (lifetime)', 'lifetime_5'),
             ('Composite (any of 1-5, lifetime)', 'lifetime_any')]

A_post_cont = [('Age at index, years', 'age_at_index'),
               ('Birth year', 'birth_year') if False else ('Birth year', 'birth_year'),
               ('Index year', 'index_year') if 'index_year' in A_post.columns else ('Birth year (post)', 'birth_year'),
               ('Follow-up, days', 'follow_up_days')]
# rebuild without ternary issues:
A_post['birth_year'] = A_post['birth_date'].dt.year
A_post['index_year'] = A_post['index_date'].dt.year
A_post_cont = [('Age at index, years', 'age_at_index'),
               ('Birth year', 'birth_year'),
               ('Index year', 'index_year'),
               ('Follow-up, days', 'follow_up_days')]
A_post_bin = [('Mortality during follow-up', 'died'),
              ('협심증/심근경색 (baseline)', 'baseline_1'),
              ('고혈압 (baseline)', 'baseline_2'),
              ('당뇨 (baseline)', 'baseline_3'),
              ('뇌출혈/뇌경색 (baseline)', 'baseline_4'),
              ('폐색전증 (baseline)', 'baseline_5'),
              ('Composite (any of 1-5, baseline)', 'baseline_any')]

A_pre_rows = make_rows(A_pre_with_como, 'vaccinated', A_pre_cont, A_pre_bin)
A_post_rows = make_rows(A_post, 'vaccinated', A_post_cont, A_post_bin)

# Cohort B
B_pre_cont = [('Age at surgery, years', 'age_at_surgery'),
              ('Birth year', 'birth_year'),
              ('Surgery year', 'surg_year')]
B_pre_bin = [('Mortality (ever)', 'died_flag'),
             ('협심증/심근경색 (pre-surgery)', 'baseline_1'),
             ('고혈압 (pre-surgery)', 'baseline_2'),
             ('당뇨 (pre-surgery)', 'baseline_3'),
             ('뇌출혈/뇌경색 (pre-surgery)', 'baseline_4'),
             ('폐색전증 (pre-surgery)', 'baseline_5'),
             ('Composite (any of 1-5)', 'baseline_any')]
B_pre_cat = [('Surgery type', 'first_surg_type', ['원추절제술','자궁절제술'])]
B_pre_rows = make_rows(B_pre, 'vaccinated', B_pre_cont, B_pre_bin, B_pre_cat)

B_post['vaccinated'] = B_post['접종여부'].astype(bool)
B_post['index_year'] = B_post['index_date'].dt.year
B_post_cont = [('Age at index, years', 'index_age'),
               ('Birth year', 'birth_year'),
               ('Index year', 'index_year'),
               ('BMI, kg/m²', 'closest_bmi'),
               ('Follow-up, days', 'follow_up_days')]
B_post_bin = [('Mortality during follow-up', 'died'),
              ('협심증/심근경색 (baseline)', 'baseline_1'),
              ('고혈압 (baseline)', 'baseline_2'),
              ('당뇨 (baseline)', 'baseline_3'),
              ('뇌출혈/뇌경색 (baseline)', 'baseline_4'),
              ('폐색전증 (baseline)', 'baseline_5'),
              ('Composite (any of 1-5, baseline)', 'baseline_any')]
B_post_cat = [('Surgery type', '수술방법', ['원추절제술','자궁절제술'])]
B_post_rows = make_rows(B_post, 'vaccinated', B_post_cont, B_post_bin, B_post_cat)

# ------------------------ build docx ------------------------
print('\nWriting docx...')
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_table_block(title, header_rows, rows):
    """header_rows: list of (line1, line2) lists for 2-row header
    rows: list of (varname, vac, ctl, p, smd)"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)

    table = doc.add_table(rows=2 + len(rows), cols=5)
    table.style = 'Light Grid Accent 1'
    table.autofit = False

    # Row 0: section labels
    hdr0 = table.rows[0].cells
    hdr0[0].text = 'Characteristic'
    hdr0[1].text = header_rows[0][0]
    hdr0[2].text = header_rows[0][1]
    hdr0[3].text = 'p-value'
    hdr0[4].text = '|SMD|'
    # Row 1: n
    hdr1 = table.rows[1].cells
    hdr1[0].text = ''
    hdr1[1].text = header_rows[1][0]
    hdr1[2].text = header_rows[1][1]
    hdr1[3].text = ''
    hdr1[4].text = ''

    for c in list(hdr0) + list(hdr1):
        for para in c.paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (var, v, ctl, pv, smd) in enumerate(rows):
        cells = table.rows[i+2].cells
        cells[0].text = var
        cells[1].text = v
        cells[2].text = ctl
        cells[3].text = pv
        cells[4].text = smd
        for j, c in enumerate(cells):
            for para in c.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)
                if j > 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

# Title
title = doc.add_heading('Table 1. Baseline Characteristics of Study Cohorts', level=0)

# Cohort A
add_heading('Cohort A — Whole-cohort comorbidity analysis', level=1)
nv_pre = int(A_pre_with_como['vaccinated'].sum())
nc_pre = int((~A_pre_with_como['vaccinated']).sum())
nv_post = int(A_post['vaccinated'].sum())
nc_post = int((~A_post['vaccinated']).sum())
add_table_block(
    f'(a) Pre-matching (Vaccinated n={nv_pre} vs Non-vaccinated n={nc_pre})',
    [('Vaccinated', 'Non-vaccinated'), (f'(n={nv_pre})', f'(n={nc_pre})')],
    A_pre_rows[1:]  # drop redundant N row since shown in header
)
add_table_block(
    f'(b) Post-matching (Vaccinated n={nv_post} vs Non-vaccinated n={nc_post}, 1:4 age-matched)',
    [('Vaccinated', 'Non-vaccinated'), (f'(n={nv_post})', f'(n={nc_post})')],
    A_post_rows[1:]
)

# Cohort B
add_heading('Cohort B — Surgical efficacy analysis', level=1)
nv_bpre = int(B_pre['vaccinated'].sum())
nc_bpre = int((~B_pre['vaccinated']).sum())
nv_bpost = int(B_post['vaccinated'].sum())
nc_bpost = int((~B_post['vaccinated']).sum())
add_table_block(
    f'(c) Pre-matching (Vaccinated n={nv_bpre} vs Non-vaccinated n={nc_bpre}, all surgery patients)',
    [('Vaccinated', 'Non-vaccinated'), (f'(n={nv_bpre})', f'(n={nc_bpre})')],
    B_pre_rows[1:]
)
add_table_block(
    f'(d) Post-matching (Vaccinated n={nv_bpost} vs Non-vaccinated n={nc_bpost}, fine-matched 1:4)',
    [('Vaccinated', 'Non-vaccinated'), (f'(n={nv_bpost})', f'(n={nc_bpost})')],
    B_post_rows[1:]
)

# Footnote
foot = doc.add_paragraph()
foot.add_run('Footnote: ').bold = True
foot.add_run('Continuous variables are presented as mean ± SD; categorical/binary variables as n (%). '
            'p-values from independent-sample t-test (continuous) or Fisher\'s exact test (categorical). '
            '|SMD|: absolute standardized mean difference; values <0.10 indicate good balance. '
            'Pre-matching for Cohort A uses lifetime prevalence as no index date is available for non-vaccinated controls; '
            'post-matching uses pre-index baseline prevalence. '
            'Cohort B pre-matching uses pre-surgery baseline; post-matching uses pre-index baseline.').font.size = Pt(9)

out_path = 'Data/Table1_BaselineCharacteristics.docx'
doc.save(out_path)
print(f'Saved: {out_path}')

# Also save CSV
all_rows = []
for label, rows in [('CohortA_pre', A_pre_rows), ('CohortA_post', A_post_rows),
                    ('CohortB_pre', B_pre_rows), ('CohortB_post', B_post_rows)]:
    for r in rows:
        all_rows.append({'block': label, 'variable': r[0], 'vaccinated': r[1],
                        'non_vaccinated': r[2], 'p_value': r[3], 'abs_SMD': r[4]})
pd.DataFrame(all_rows).to_csv('Data/Table1_BaselineCharacteristics.csv', index=False, encoding='utf-8-sig')
print('Saved: Data/Table1_BaselineCharacteristics.csv')
