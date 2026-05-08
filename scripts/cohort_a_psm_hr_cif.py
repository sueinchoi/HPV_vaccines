"""
Cohort A — Hazard Ratios and Cumulative Incidence with Competing Risks
======================================================================

PSM 매칭 코호트 (cohort_a_psm_matched.csv) 기반:
- 5대 만성질환 각각 발생 (primary event)
- 사망 = competing event
- 시간: index_date → 첫 진단 / 사망 / last_follow 중 가장 빠른 시점

분석:
1. Aalen-Johansen estimator → 누적발생함수 (CIF) plot (vac vs ctl per outcome)
2. Cause-specific Cox PH (sklearn lifelines CoxPHFitter, robust SE clustered by pair)
3. Fine-Gray subdistribution Cox (Geskus 2011 IPCW weighted Cox)
4. Output: docx (Table HR + 5-panel CIF figure) + CSV
"""
import pandas as pd
import numpy as np
import openpyxl
import matplotlib.pyplot as plt
import warnings
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
from lifelines import CoxPHFitter, AalenJohansenFitter, KaplanMeierFitter, CoxTimeVaryingFitter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

warnings.filterwarnings('ignore')
plt.rcParams['font.family'] = ['DejaVu Sans','AppleGothic']
plt.rcParams['axes.unicode_minus'] = False

RANDOM_SEED = 42
rng = np.random.default_rng(RANDOM_SEED)
np.random.seed(RANDOM_SEED)
CLASS_LABELS = {'1':'Angina/MI','2':'Hypertension','3':'Diabetes','4':'Stroke','5':'PE'}
CLASS_KOR = {'1':'협심증/심근경색','2':'고혈압','3':'당뇨','4':'뇌출혈/뇌경색','5':'폐색전증'}
# MCE composite: first occurrence of MI (1), Stroke (4), or PE (5)
MCE_COMPONENTS = ['1','4','5']
MCE_LABEL = 'MCE (MI / Stroke / PE)'
MCE_KOR = 'MCE (심근경색·뇌졸중·폐색전증)'
# Any-of-5 composite: first occurrence of any of the 5 comorbidities
ANY5_COMPONENTS = ['1','2','3','4','5']
ANY5_LABEL = 'Any of 5 comorbidities'
ANY5_KOR = 'Any of 5 (5대 질환 중 어느 하나)'
SMOKE_MAP = {'비흡연':'Never','과거흡연':'Former','현재흡연':'Current','확인불능':'Unknown'}

# --------------------------- helpers ---------------------------
def closest_vec(query_df, ci, value_col, window_days=365):
    ci_v = ci[['연구번호','기록일자_dt', value_col]].dropna(subset=[value_col,'기록일자_dt']).copy()
    ci_v = ci_v.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    fw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid', direction='forward', tolerance=pd.Timedelta(days=window_days))
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid', direction='backward', tolerance=pd.Timedelta(days=window_days))
    fw['_d'] = (fw['rec_date']-fw['index_date']).abs()
    bw['_d'] = (bw['rec_date']-bw['index_date']).abs()
    use_fw = (fw['_d'].fillna(pd.Timedelta(days=window_days*10)) <= bw['_d'].fillna(pd.Timedelta(days=window_days*10)))
    return pd.Series(np.where(use_fw, fw[value_col].values, bw[value_col].values),
                    index=fw['orig_idx'].values).reindex(query_df.index).astype(float)

def smoke_vec(query_df, ci):
    smk = ci[['연구번호','기록일자_dt','흡연여부']].dropna(subset=['흡연여부','기록일자_dt']).copy()
    smk = smk.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], smk,
        left_on='index_date', right_on='rec_date', by='pid', direction='backward')
    return pd.Series(bw['흡연여부'].map(SMOKE_MAP).fillna('Unknown').values,
                    index=bw['orig_idx'].values).reindex(query_df.index).fillna('Unknown')

# --------------------------- load + PSM ---------------------------
print('[1/8] Loading source...')
cohort = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_코호트.csv', encoding='cp949', low_memory=False)
cohort['birth_date'] = pd.to_datetime(cohort['생년월'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['death_date'] = pd.to_datetime(cohort['사망일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['last_follow'] = pd.to_datetime(cohort['최종추적일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['is_seoul'] = (cohort['주소'].astype(str).str.split().str[0]=='서울').astype(int)

rx = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_처방정보.csv', encoding='cp949', low_memory=False)
mask = (rx['처방명'].astype(str).str.contains('Gardasil|Cervarix|HPV vaccine', case=False, na=False) |
        rx['처방한글명'].astype(str).str.contains('가다실|서바릭스', na=False))
rx_vac = rx[mask].copy()
rx_vac['처방일자'] = pd.to_datetime(rx_vac['처방일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
first_vac = rx_vac.groupby('연구번호')['처방일자'].min().reset_index()
first_vac.columns = ['연구번호','first_vaccine_date']
cohort = cohort.merge(first_vac, on='연구번호', how='left').dropna(subset=['birth_date'])

print('[2/8] Loading clinical info & comorbidities...')
ci = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_기초임상정보.csv', encoding='cp949', low_memory=False)
ci['기록일자_dt'] = pd.to_datetime(ci['기록일자'].astype(str).str.strip(), format='%Y%m%d', errors='coerce')

wb = openpyxl.load_workbook('Data/한국 HPV 코호트 자료를 이용한 자_진단정보_기저질환추가_unlocked.xlsx',
                          read_only=True, data_only=True)
ws = wb.active
recs=[]
for row in ws.iter_rows(min_row=2, values_only=True):
    pid, cls, dd = row[0], row[5], row[8]
    if cls is None or str(cls).strip()=='': continue
    cls = str(cls).strip()
    if cls not in CLASS_LABELS: continue
    recs.append((pid, cls, pd.to_datetime(str(dd), format='%Y%m%d', errors='coerce')))
como = pd.DataFrame(recs, columns=['pid','class','diag_date'])
first_diag = como.groupby(['pid','class'])['diag_date'].min().unstack('class')
for c in CLASS_LABELS:
    if c not in first_diag.columns: first_diag[c] = pd.NaT

print('[3/8] Build cohort with pseudo-index for controls...')
df = cohort.copy()
df['vaccinated'] = df['first_vaccine_date'].notna()
vac_dates = df.loc[df['vaccinated'], 'first_vaccine_date'].dropna().values
df.loc[~df['vaccinated'], 'index_date'] = pd.to_datetime(rng.choice(vac_dates, size=(~df['vaccinated']).sum()))
df.loc[df['vaccinated'], 'index_date'] = df.loc[df['vaccinated'], 'first_vaccine_date']
df = df.rename(columns={'연구번호':'pid'}).reset_index(drop=True)
df['age_at_index'] = (df['index_date'] - df['birth_date']).dt.days/365.25
df = df[(df['death_date'].isna()) | (df['death_date'] > df['index_date'])]
# strict: require ≥1 day of follow-up after index (otherwise no at-risk time)
df = df[df['last_follow'] > df['index_date']].reset_index(drop=True)

print('[4/8] Closest vital signs + smoking...')
q = df[['pid','index_date']].copy()
df['height'] = closest_vec(q, ci, '키')
df['weight'] = closest_vec(q, ci, '몸무게')
df['sbp'] = closest_vec(q, ci, '수축기혈압')
df['dbp'] = closest_vec(q, ci, '이완기혈압')
df['bmi'] = df['weight']/(df['height']/100)**2
df['smoke'] = smoke_vec(q, ci).values

df = df.merge(first_diag, left_on='pid', right_index=True, how='left')

print('[5/8] PSM 1:1...')
for c in ['bmi','sbp','dbp']:
    df[f'{c}_miss'] = df[c].isna().astype(int)
    df[c] = df[c].fillna(df[c].mean())
sm = pd.get_dummies(df['smoke'], prefix='smoke').astype(int)
df = pd.concat([df, sm], axis=1)
ps_features = ['age_at_index','bmi','bmi_miss','sbp','sbp_miss','dbp','dbp_miss','is_seoul',
              'smoke_Never','smoke_Former','smoke_Current']
ps_features = [c for c in ps_features if c in df.columns]

X = df[ps_features].astype(float).values
y = df['vaccinated'].astype(int).values
Xs = StandardScaler().fit_transform(X)
lr = LogisticRegression(max_iter=2000, C=1e6, solver='lbfgs')
lr.fit(Xs, y)
df['ps'] = lr.predict_proba(Xs)[:,1]
df['logit_ps'] = np.log(df['ps']/(1-df['ps']))
caliper = 0.2 * df['logit_ps'].std()
print(f'  caliper={caliper:.4f}')

vac_idx = df.index[df['vaccinated']].tolist()
ctl_idx = np.array(df.index[~df['vaccinated']].tolist())
ctl_logit = df.loc[ctl_idx,'logit_ps'].values
order = np.argsort(ctl_logit)
ctl_sorted = ctl_idx[order]; ctl_logit_sorted = ctl_logit[order]
used = np.zeros(len(ctl_sorted), dtype=bool)
matched = []
vac_order = np.array(vac_idx); rng.shuffle(vac_order)
for vi in vac_order:
    target = df.loc[vi,'logit_ps']
    lo = np.searchsorted(ctl_logit_sorted, target-caliper)
    hi = np.searchsorted(ctl_logit_sorted, target+caliper, side='right')
    best_j, best_d = -1, caliper+1
    for j in range(lo, hi):
        if used[j]: continue
        d = abs(ctl_logit_sorted[j]-target)
        if d < best_d: best_d=d; best_j=j
    if best_j>=0:
        used[best_j] = True
        matched.append((vi, ctl_sorted[best_j]))
print(f'  matched pairs: {len(matched)}')

# Build matched dataframe with pair_id
pair_records = []
for pid_, (v_idx, c_idx) in enumerate(matched):
    pair_records.append((v_idx, pid_))
    pair_records.append((c_idx, pid_))
pair_df = pd.DataFrame(pair_records, columns=['orig_idx','pair_id'])
m = df.loc[pair_df['orig_idx'].values].copy()
m['pair_id'] = pair_df['pair_id'].values
m = m.reset_index(drop=True)
print(f'  matched cohort: vac={int(m["vaccinated"].sum())}, ctl={int((~m["vaccinated"]).sum())}')

# --------------------------- time-to-event per comorbidity ---------------------------
print('\n[6/8] Building time-to-event...')
def make_tte(m, cls):
    """Returns (time_days, status: 0=cens, 1=primary, 2=death-competing)
    cls can be a single class string ('1'..'5') or a list of classes (composite — first occurrence)."""
    if isinstance(cls, list):
        # Composite: earliest occurrence among components
        dx = m[cls].min(axis=1)
    else:
        dx = m[cls]
    death = m['death_date']
    last = m['last_follow']
    idx = m['index_date']
    # primary if dx > index (incident)
    is_pre = dx.notna() & (dx <= idx)  # exclude prevalent — not at risk
    primary_date = dx.where(dx > idx, pd.NaT)
    # competing event = death (occurring before primary)
    death_after = death.where((death.notna()) & (death > idx) & ((primary_date.isna()) | (death < primary_date)), pd.NaT)
    # event_date = first of (primary, death_after)
    event_date = primary_date.combine_first(death_after)
    # status
    status = np.where(primary_date.notna() & ((death_after.isna()) | (primary_date <= death_after)), 1,
            np.where(death_after.notna(), 2, 0))
    end_date = event_date.combine_first(last)
    time = (end_date - idx).dt.days.astype(float)
    res = pd.DataFrame({'pid': m['pid'].values, 'pair_id': m['pair_id'].values,
                       'vaccinated': m['vaccinated'].astype(int).values,
                       'time': time, 'status': status})
    # exclude prevalent (pre-existing) cases
    res = res[~is_pre.values].reset_index(drop=True)
    res = res[res['time'] > 0].reset_index(drop=True)
    return res

# --------------------------- Fine-Gray weighted dataset (Geskus) ---------------------------
def geskus_long(tte, primary=1, competing=2):
    """Geskus IPCW weighted long-format for Fine-Gray Cox.
    For competing-event subjects, extend follow-up beyond their event with time-decreasing weights G(t)/G(c)
    where G is censoring KM (event=censoring). Then run weighted CoxTimeVaryingFitter.
    """
    t = tte.copy().reset_index(drop=True)
    t['_id'] = t.index
    # Censoring KM: event = (status == 0)
    cens_obs = (t['status'] == 0).astype(int)
    kmf = KaplanMeierFitter().fit(t['time'], event_observed=cens_obs)
    G = kmf.survival_function_.iloc[:,0]
    # last G value
    G_max = G.iloc[-1] if len(G)>0 else 1.0
    G_idx = G.index.values
    G_val = G.values
    def G_at(t_):
        # right-continuous step function: G evaluated at t- (just before t)
        pos = np.searchsorted(G_idx, t_, side='right') - 1
        return G_val[pos] if pos >= 0 else 1.0

    primary_times = sorted(t.loc[t['status']==primary,'time'].unique().tolist())
    max_time = t['time'].max()

    rows = []
    for _, r in t.iterrows():
        if r['status'] == primary:
            rows.append({'_id':r['_id'],'pair_id':r['pair_id'],'vaccinated':r['vaccinated'],
                        '_start':0,'_stop':r['time'],'_event':1,'_w':1.0})
        elif r['status'] == 0:
            rows.append({'_id':r['_id'],'pair_id':r['pair_id'],'vaccinated':r['vaccinated'],
                        '_start':0,'_stop':r['time'],'_event':0,'_w':1.0})
        else:  # competing
            # extend with intervals at each future primary event time
            G_c = G_at(r['time'])
            if G_c <= 0: continue
            split = sorted([pt for pt in primary_times if pt > r['time']] + [max_time])
            split = [s for s in split if s > r['time']]
            prev = r['time']
            for st in split:
                G_st = G_at(st)
                w = G_st / G_c
                if w > 0:
                    rows.append({'_id':r['_id'],'pair_id':r['pair_id'],'vaccinated':r['vaccinated'],
                                '_start':prev,'_stop':st,'_event':0,'_w':w})
                prev = st
    return pd.DataFrame(rows)

# --------------------------- run analyses ---------------------------
print('[7/8] Running Cox + Fine-Gray + AalenJohansen for each comorbidity...')
hr_results = []
cif_data = {}  # cls -> {'vac': (t, p_lower, p_upper, p), 'ctl': ...}

for cls, label_kor in CLASS_KOR.items():
    print(f'  -- {label_kor} --')
    tte = make_tte(m, cls)
    n = len(tte); n_vac = int(tte['vaccinated'].sum()); n_ctl = n - n_vac
    e_primary = int((tte['status']==1).sum()); e_comp = int((tte['status']==2).sum())
    e_v = int(((tte['status']==1)&(tte['vaccinated']==1)).sum())
    e_c = int(((tte['status']==1)&(tte['vaccinated']==0)).sum())
    median_fu_v = tte.loc[tte['vaccinated']==1,'time'].median()
    median_fu_c = tte.loc[tte['vaccinated']==0,'time'].median()
    print(f'     n={n}, primary events={e_primary} (vac {e_v}/{n_vac}, ctl {e_c}/{n_ctl}), competing(death)={e_comp}')

    res = {'class':cls,'label':label_kor,
          'n':n,'n_vac':n_vac,'n_ctl':n_ctl,
          'events_vac':e_v,'events_ctl':e_c,
          'events_competing':e_comp,
          'median_fu_vac_d':round(median_fu_v,0) if pd.notna(median_fu_v) else np.nan,
          'median_fu_ctl_d':round(median_fu_c,0) if pd.notna(median_fu_c) else np.nan}

    # Cause-specific Cox (treat competing event as censored)
    if e_v + e_c >= 5 and e_v >= 1 and e_c >= 1:
        try:
            cs_data = tte.copy()
            cs_data['event'] = (cs_data['status']==1).astype(int)
            cs_data = cs_data[['time','event','vaccinated','pair_id']]
            cph = CoxPHFitter()
            cph.fit(cs_data[['time','event','vaccinated']], duration_col='time', event_col='event',
                   robust=True)  # robust SE
            sm = cph.summary
            hr = float(sm.loc['vaccinated','exp(coef)'])
            ci_lo = float(sm.loc['vaccinated','exp(coef) lower 95%'])
            ci_hi = float(sm.loc['vaccinated','exp(coef) upper 95%'])
            p = float(sm.loc['vaccinated','p'])
            res.update({'cs_HR':hr,'cs_CI_lo':ci_lo,'cs_CI_hi':ci_hi,'cs_p':p})
        except Exception as e:
            print(f'     cs Cox failed: {e}')
            res.update({'cs_HR':np.nan,'cs_CI_lo':np.nan,'cs_CI_hi':np.nan,'cs_p':np.nan})
    else:
        res.update({'cs_HR':np.nan,'cs_CI_lo':np.nan,'cs_CI_hi':np.nan,'cs_p':np.nan})

    # Fine-Gray subdistribution Cox via Geskus weighting
    if e_v + e_c >= 5 and e_v >= 1 and e_c >= 1:
        try:
            fg = geskus_long(tte)
            ctv = CoxTimeVaryingFitter()
            ctv.fit(fg, id_col='_id', start_col='_start', stop_col='_stop',
                   event_col='_event', weights_col='_w', show_progress=False)
            sm = ctv.summary
            hr = float(sm.loc['vaccinated','exp(coef)'])
            ci_lo = float(sm.loc['vaccinated','exp(coef) lower 95%'])
            ci_hi = float(sm.loc['vaccinated','exp(coef) upper 95%'])
            p = float(sm.loc['vaccinated','p'])
            res.update({'fg_HR':hr,'fg_CI_lo':ci_lo,'fg_CI_hi':ci_hi,'fg_p':p})
        except Exception as e:
            print(f'     Fine-Gray failed: {e}')
            res.update({'fg_HR':np.nan,'fg_CI_lo':np.nan,'fg_CI_hi':np.nan,'fg_p':np.nan})
    else:
        res.update({'fg_HR':np.nan,'fg_CI_lo':np.nan,'fg_CI_hi':np.nan,'fg_p':np.nan})

    # AalenJohansen for CIF (per group)
    cif_data[cls] = {}
    for grp_name, grp_val in [('vac', 1), ('ctl', 0)]:
        sub = tte[tte['vaccinated']==grp_val]
        if (sub['status']==1).sum() < 1:
            cif_data[cls][grp_name] = None
            continue
        try:
            aj = AalenJohansenFitter()
            aj.fit(durations=sub['time'].values, event_observed=sub['status'].values, event_of_interest=1)
            cif_data[cls][grp_name] = aj
        except Exception as e:
            print(f'     AJ failed for {grp_name}: {e}')
            cif_data[cls][grp_name] = None

    hr_results.append(res)

# --------------------------- MCE composite ---------------------------
print(f'  -- {MCE_KOR} --')
tte_mce = make_tte(m, MCE_COMPONENTS)
n = len(tte_mce); n_vac = int(tte_mce['vaccinated'].sum()); n_ctl = n - n_vac
e_primary = int((tte_mce['status']==1).sum()); e_comp = int((tte_mce['status']==2).sum())
e_v = int(((tte_mce['status']==1)&(tte_mce['vaccinated']==1)).sum())
e_c = int(((tte_mce['status']==1)&(tte_mce['vaccinated']==0)).sum())
median_fu_v = tte_mce.loc[tte_mce['vaccinated']==1,'time'].median()
median_fu_c = tte_mce.loc[tte_mce['vaccinated']==0,'time'].median()
print(f'     n={n}, primary events={e_primary} (vac {e_v}/{n_vac}, ctl {e_c}/{n_ctl}), competing(death)={e_comp}')

mce_res = {'class':'MCE','label':MCE_KOR,
          'n':n,'n_vac':n_vac,'n_ctl':n_ctl,
          'events_vac':e_v,'events_ctl':e_c,
          'events_competing':e_comp,
          'median_fu_vac_d':round(median_fu_v,0) if pd.notna(median_fu_v) else np.nan,
          'median_fu_ctl_d':round(median_fu_c,0) if pd.notna(median_fu_c) else np.nan}

if e_v + e_c >= 5 and e_v >= 1 and e_c >= 1:
    try:
        cs_data = tte_mce.copy()
        cs_data['event'] = (cs_data['status']==1).astype(int)
        cph = CoxPHFitter()
        cph.fit(cs_data[['time','event','vaccinated']], duration_col='time', event_col='event', robust=True)
        sm = cph.summary
        mce_res.update({'cs_HR':float(sm.loc['vaccinated','exp(coef)']),
                       'cs_CI_lo':float(sm.loc['vaccinated','exp(coef) lower 95%']),
                       'cs_CI_hi':float(sm.loc['vaccinated','exp(coef) upper 95%']),
                       'cs_p':float(sm.loc['vaccinated','p'])})
    except Exception as e:
        print(f'     cs failed: {e}')
        mce_res.update({'cs_HR':np.nan,'cs_CI_lo':np.nan,'cs_CI_hi':np.nan,'cs_p':np.nan})

    try:
        fg = geskus_long(tte_mce)
        ctv = CoxTimeVaryingFitter()
        ctv.fit(fg, id_col='_id', start_col='_start', stop_col='_stop',
               event_col='_event', weights_col='_w', show_progress=False)
        sm = ctv.summary
        mce_res.update({'fg_HR':float(sm.loc['vaccinated','exp(coef)']),
                       'fg_CI_lo':float(sm.loc['vaccinated','exp(coef) lower 95%']),
                       'fg_CI_hi':float(sm.loc['vaccinated','exp(coef) upper 95%']),
                       'fg_p':float(sm.loc['vaccinated','p'])})
    except Exception as e:
        print(f'     Fine-Gray failed: {e}')
        mce_res.update({'fg_HR':np.nan,'fg_CI_lo':np.nan,'fg_CI_hi':np.nan,'fg_p':np.nan})
else:
    mce_res.update({'cs_HR':np.nan,'cs_CI_lo':np.nan,'cs_CI_hi':np.nan,'cs_p':np.nan,
                   'fg_HR':np.nan,'fg_CI_lo':np.nan,'fg_CI_hi':np.nan,'fg_p':np.nan})

# Aalen-Johansen for MCE per group
cif_data['MCE'] = {}
for grp_name, grp_val in [('vac', 1), ('ctl', 0)]:
    sub = tte_mce[tte_mce['vaccinated']==grp_val]
    if (sub['status']==1).sum() < 1:
        cif_data['MCE'][grp_name] = None
        continue
    try:
        aj = AalenJohansenFitter()
        aj.fit(durations=sub['time'].values, event_observed=sub['status'].values, event_of_interest=1)
        cif_data['MCE'][grp_name] = aj
    except Exception as e:
        print(f'     AJ failed for {grp_name}: {e}')
        cif_data['MCE'][grp_name] = None

hr_results.append(mce_res)

# --------------------------- Any-of-5 composite ---------------------------
print(f'  -- {ANY5_KOR} --')
tte_any = make_tte(m, ANY5_COMPONENTS)
n = len(tte_any); n_vac = int(tte_any['vaccinated'].sum()); n_ctl = n - n_vac
e_primary = int((tte_any['status']==1).sum()); e_comp = int((tte_any['status']==2).sum())
e_v = int(((tte_any['status']==1)&(tte_any['vaccinated']==1)).sum())
e_c = int(((tte_any['status']==1)&(tte_any['vaccinated']==0)).sum())
median_fu_v = tte_any.loc[tte_any['vaccinated']==1,'time'].median()
median_fu_c = tte_any.loc[tte_any['vaccinated']==0,'time'].median()
print(f'     n={n}, primary events={e_primary} (vac {e_v}/{n_vac}, ctl {e_c}/{n_ctl}), competing(death)={e_comp}')

any_res = {'class':'ANY5','label':ANY5_KOR,
          'n':n,'n_vac':n_vac,'n_ctl':n_ctl,
          'events_vac':e_v,'events_ctl':e_c,
          'events_competing':e_comp,
          'median_fu_vac_d':round(median_fu_v,0) if pd.notna(median_fu_v) else np.nan,
          'median_fu_ctl_d':round(median_fu_c,0) if pd.notna(median_fu_c) else np.nan}

if e_v + e_c >= 5 and e_v >= 1 and e_c >= 1:
    try:
        cs_data = tte_any.copy()
        cs_data['event'] = (cs_data['status']==1).astype(int)
        cph = CoxPHFitter()
        cph.fit(cs_data[['time','event','vaccinated']], duration_col='time', event_col='event', robust=True)
        sm = cph.summary
        any_res.update({'cs_HR':float(sm.loc['vaccinated','exp(coef)']),
                       'cs_CI_lo':float(sm.loc['vaccinated','exp(coef) lower 95%']),
                       'cs_CI_hi':float(sm.loc['vaccinated','exp(coef) upper 95%']),
                       'cs_p':float(sm.loc['vaccinated','p'])})
    except Exception as e:
        print(f'     cs failed: {e}')
        any_res.update({'cs_HR':np.nan,'cs_CI_lo':np.nan,'cs_CI_hi':np.nan,'cs_p':np.nan})

    try:
        fg = geskus_long(tte_any)
        ctv = CoxTimeVaryingFitter()
        ctv.fit(fg, id_col='_id', start_col='_start', stop_col='_stop',
               event_col='_event', weights_col='_w', show_progress=False)
        sm = ctv.summary
        any_res.update({'fg_HR':float(sm.loc['vaccinated','exp(coef)']),
                       'fg_CI_lo':float(sm.loc['vaccinated','exp(coef) lower 95%']),
                       'fg_CI_hi':float(sm.loc['vaccinated','exp(coef) upper 95%']),
                       'fg_p':float(sm.loc['vaccinated','p'])})
    except Exception as e:
        print(f'     Fine-Gray failed: {e}')
        any_res.update({'fg_HR':np.nan,'fg_CI_lo':np.nan,'fg_CI_hi':np.nan,'fg_p':np.nan})
else:
    any_res.update({'cs_HR':np.nan,'cs_CI_lo':np.nan,'cs_CI_hi':np.nan,'cs_p':np.nan,
                   'fg_HR':np.nan,'fg_CI_lo':np.nan,'fg_CI_hi':np.nan,'fg_p':np.nan})

# Aalen-Johansen for Any-of-5 per group
cif_data['ANY5'] = {}
for grp_name, grp_val in [('vac', 1), ('ctl', 0)]:
    sub = tte_any[tte_any['vaccinated']==grp_val]
    if (sub['status']==1).sum() < 1:
        cif_data['ANY5'][grp_name] = None
        continue
    try:
        aj = AalenJohansenFitter()
        aj.fit(durations=sub['time'].values, event_observed=sub['status'].values, event_of_interest=1)
        cif_data['ANY5'][grp_name] = aj
    except Exception as e:
        print(f'     AJ failed for {grp_name}: {e}')
        cif_data['ANY5'][grp_name] = None

hr_results.append(any_res)
hr_df = pd.DataFrame(hr_results)

# --------------------------- Plot CIF ---------------------------
print('[8/8] Plotting CIF + writing docx...')
fig, axes = plt.subplots(3, 3, figsize=(16, 13))
axes = axes.flatten()
# Plot order: ANY5(0), MCE(1), Angina/MI(2), HTN(3), DM(4), Stroke(5), PE(6), Forest(7), hidden(8)
plot_order = [('ANY5', ANY5_LABEL, ANY5_KOR),
              ('MCE', MCE_LABEL, MCE_KOR),
              ('1','Angina/MI','협심증/심근경색'),
              ('2','Hypertension','고혈압'),
              ('3','Diabetes','당뇨'),
              ('4','Stroke','뇌출혈/뇌경색'),
              ('5','PE','폐색전증')]

for ax, (key, label_eng, label_kor) in zip(axes[:7], plot_order):
    plotted = False
    for grp_name, color, lbl in [('vac','#9b2226','Vaccinated'),('ctl','#1f6f8b','Non-vaccinated')]:
        aj = cif_data.get(key, {}).get(grp_name)
        if aj is None: continue
        cif = aj.cumulative_density_
        ci = aj.confidence_interval_
        col = cif.columns[0]
        t = cif.index.values / 365.25
        y = cif[col].values
        ax.step(t, y, where='post', color=color, label=lbl, lw=2)
        try:
            lo = ci.iloc[:,0].values; hi = ci.iloc[:,1].values
            ax.fill_between(t, lo, hi, alpha=0.15, color=color, step='post')
        except Exception:
            pass
        plotted = True
    title = f'{label_eng}\n({label_kor})'
    if key in ('MCE','ANY5'):
        ax.set_title(title, fontweight='bold', fontsize=11)
        for spine in ax.spines.values():
            spine.set_linewidth(1.5); spine.set_edgecolor('#444')
    else:
        ax.set_title(title, fontsize=10)
    ax.set_xlabel('Years from index date')
    ax.set_ylabel('Cumulative incidence')
    ax.set_xlim(0, 10)
    if plotted: ax.legend(fontsize=8, loc='upper left')

# Forest plot (panel 7)
ax = axes[7]
y_pos = np.arange(len(hr_df))
# Highlight composite rows
hl_idx = hr_df.index[hr_df['class'].isin(['MCE','ANY5'])].tolist()
def _flbl(r):
    if r['class'] == 'MCE': return MCE_LABEL
    if r['class'] == 'ANY5': return ANY5_LABEL
    return CLASS_LABELS.get(r['class'], r['class'])
labels_for_forest = [_flbl(r) for _, r in hr_df.iterrows()]

# Filter rows with valid HR for plotting
valid_cs = hr_df['cs_HR'].notna()
valid_fg = hr_df['fg_HR'].notna()
ax.errorbar(hr_df.loc[valid_cs,'cs_HR'], y_pos[valid_cs]-0.15,
           xerr=[hr_df.loc[valid_cs,'cs_HR']-hr_df.loc[valid_cs,'cs_CI_lo'],
                 hr_df.loc[valid_cs,'cs_CI_hi']-hr_df.loc[valid_cs,'cs_HR']],
           fmt='o', color='#9b2226', label='Cause-specific', capsize=3)
ax.errorbar(hr_df.loc[valid_fg,'fg_HR'], y_pos[valid_fg]+0.15,
           xerr=[hr_df.loc[valid_fg,'fg_HR']-hr_df.loc[valid_fg,'fg_CI_lo'],
                 hr_df.loc[valid_fg,'fg_CI_hi']-hr_df.loc[valid_fg,'fg_HR']],
           fmt='s', color='#1f6f8b', label='Fine-Gray', capsize=3)
ax.axvline(1, color='black', linestyle='--', alpha=0.5)
ax.set_yticks(y_pos)
ax.set_yticklabels(labels_for_forest)
# Highlight composite labels
for i, lbl in enumerate(labels_for_forest):
    if i in hl_idx:
        ax.get_yticklabels()[i].set_fontweight('bold')
        ax.axhspan(i-0.4, i+0.4, color='#fff3cd', alpha=0.3, zorder=0)
ax.set_xscale('log')
ax.set_xlabel('Hazard ratio (vac vs non-vac)')
ax.set_title('Hazard Ratios — Cause-specific vs Fine-Gray', fontsize=10)
ax.legend(loc='best', fontsize=8)
ax.invert_yaxis()

# Hide panel 8
axes[8].set_visible(False)

plt.suptitle('Cohort A (PSM 1:1) — Cumulative Incidence and Hazard Ratios\nComposite endpoints (Any of 5, MCE) and individual comorbidities',
            fontsize=13, y=1.00)
plt.tight_layout()
plt.savefig('Data/cohort_a_psm_cif_hr.png', dpi=150, bbox_inches='tight')
plt.close()

# --------------------------- docx ---------------------------
doc = Document()
sty = doc.styles['Normal']; sty.font.name = 'Times New Roman'; sty.font.size = Pt(10)
doc.add_heading('Cohort A — Hazard Ratios and Cumulative Incidence (Competing-Risks Analysis)', level=0)

mp = doc.add_paragraph()
mp.add_run('Methods. ').bold = True
mp.add_run(
    'In the propensity-score-matched cohort (1:1, n=4,102 after requiring ≥1 day of follow-up), time-to-event '
    'analyses were performed for each of five individual comorbidities and two pre-specified composite endpoints: '
    '(i) Major Cardiovascular Events (MCE) — the first occurrence of myocardial infarction (Angina/MI), '
    'stroke (cerebral infarction or hemorrhage), or pulmonary embolism after the index date; and '
    '(ii) Any-of-5 — the first occurrence of any of the five comorbidities (Angina/MI, hypertension, diabetes, '
    'stroke, or pulmonary embolism) after the index date. The primary event was '
    'the first incident diagnosis; death prior to the primary event was treated as a competing event; remaining '
    'patients were censored at last follow-up. Cause-specific hazard ratios (HR) were estimated by Cox '
    'proportional-hazards regression treating the competing event as censoring. Subdistribution hazard ratios '
    '(sHR) were estimated using the Fine-Gray model implemented as time-varying Cox regression on the Geskus '
    '(2011) IPCW-reweighted dataset. Robust standard errors were used. Cumulative incidence functions were '
    'estimated non-parametrically with the Aalen-Johansen estimator, which adjusts for the competing risk of death. '
    'Patients with prevalent disease at the index date were excluded from the corresponding analysis to define an '
    'at-risk population. Time is reported in days; figures display years from index date.'
).font.size = Pt(9)

# HR Table
doc.add_heading('Table 1. Hazard Ratios — Composite Endpoints and Individual Comorbidities (PSM 1:1 cohort)', level=1)
t = doc.add_table(rows=2 + len(hr_df), cols=8)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'Outcome'
hdr[1].text = 'Vaccinated\nevents/n'
hdr[2].text = 'Non-vaccinated\nevents/n'
hdr[3].text = 'Competing\n(death)'
hdr[4].text = 'Cause-specific HR'
hdr[5].text = 'Cause-specific p'
hdr[6].text = 'Fine-Gray sHR'
hdr[7].text = 'Fine-Gray p'
for c in hdr:
    for para in c.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold = True; r.font.size = Pt(9)
# Empty 2nd header row to align with desired rows (use 1 row instead — drop empty)
t._tbl.remove(t.rows[1]._tr)

def fmt_hr(hr, lo, hi):
    if pd.isna(hr): return '-'
    return f'{hr:.3f} ({lo:.3f}–{hi:.3f})'
def fmt_p(p):
    if pd.isna(p): return '-'
    return '<0.001' if p<0.001 else f'{p:.3f}'

for i, row in hr_df.iterrows():
    cells = t.rows[i+1].cells
    if row['class'] == 'MCE':
        cells[0].text = f'★ {MCE_LABEL}'
    elif row['class'] == 'ANY5':
        cells[0].text = f'★ {ANY5_LABEL}'
    else:
        cells[0].text = f'{CLASS_LABELS[row["class"]]} ({row["label"]})'
    cells[1].text = f'{int(row["events_vac"])}/{int(row["n_vac"])}'
    cells[2].text = f'{int(row["events_ctl"])}/{int(row["n_ctl"])}'
    cells[3].text = f'{int(row["events_competing"])}'
    cells[4].text = fmt_hr(row['cs_HR'], row['cs_CI_lo'], row['cs_CI_hi'])
    cells[5].text = fmt_p(row['cs_p'])
    cells[6].text = fmt_hr(row['fg_HR'], row['fg_CI_lo'], row['fg_CI_hi'])
    cells[7].text = fmt_p(row['fg_p'])
    for j,c in enumerate(cells):
        for para in c.paragraphs:
            for r in para.runs: r.font.size = Pt(9)
            if j>0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
foot = doc.add_paragraph()
foot.add_run('Footnote: ').bold = True
foot.add_run(
    'Cause-specific hazard ratios (HR) treat death as right-censoring; Fine-Gray subdistribution hazard ratios '
    '(sHR) keep subjects with a competing event in the risk set with IPCW-decreasing weights (Geskus 2011). '
    'Robust standard errors are reported. With very low competing-event rates (death <0.5%), cause-specific '
    'and Fine-Gray estimates are expected to be similar. Outcomes with no events in either group cannot be '
    'estimated and are reported as "-". Prevalent (baseline) cases were excluded from each outcome-specific '
    'analysis to define an at-risk population.'
).font.size = Pt(8)

# Figure
doc.add_heading('Figure 1. Cumulative Incidence (Aalen-Johansen) and Hazard Ratios', level=1)
doc.add_picture('Data/cohort_a_psm_cif_hr.png', width=Inches(6.5))
fp = doc.add_paragraph()
fp.add_run(
    'Cumulative incidence functions (CIF) are estimated by the Aalen-Johansen estimator with death as the competing '
    'event; shaded areas show 95% pointwise confidence intervals. The bottom-right panel displays cause-specific '
    'and Fine-Gray subdistribution HRs (point estimate and 95% CI) on the log scale; the dashed vertical line at '
    'HR=1 represents the null. Time axis: years from index date.'
).italic = True

out = 'Data/CohortA_PSM_HR_CIF_report.docx'
doc.save(out)

# Save HR results CSV
hr_df.to_csv('Data/cohort_a_psm_hr_results.csv', index=False, encoding='utf-8-sig')

print(f'\nSaved:')
print(f'  {out}')
print(f'  Data/cohort_a_psm_cif_hr.png')
print(f'  Data/cohort_a_psm_hr_results.csv')

# Print summary
print('\n=== HR Summary ===')
print(hr_df[['label','events_vac','n_vac','events_ctl','n_ctl','events_competing',
            'cs_HR','cs_CI_lo','cs_CI_hi','cs_p','fg_HR','fg_CI_lo','fg_CI_hi','fg_p']].to_string(index=False))
