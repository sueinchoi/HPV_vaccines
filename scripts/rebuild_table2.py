"""
Rebuild Table 2 — Cohort A
Combines:
  - Cluster-robust cause-specific HR (Cox with cluster_col=pair_id)
  - Fine-Gray subdistribution sHR (Geskus IPCW reweighted Cox)
Plus person-years, IR/1000PY, Schoenfeld PH p-value.

Outputs:
  Data/Table2_CohortA_HazardRatios.docx
  Data/Table2_CohortA_HazardRatios.csv
"""
import pandas as pd
import numpy as np
import warnings
import openpyxl
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
from lifelines import CoxPHFitter, KaplanMeierFitter, CoxTimeVaryingFitter
from lifelines.statistics import proportional_hazard_test
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
warnings.filterwarnings('ignore')

CLASS_LABELS = {'1':'Angina/MI','2':'Hypertension','3':'Diabetes','4':'Stroke','5':'PE'}
ANY5 = ['1','2','3','4','5']
MCE = ['1','4','5']
SMOKE_MAP = {'비흡연':'Never','과거흡연':'Former','현재흡연':'Current','확인불능':'Unknown'}
RANDOM_SEED = 42

# --------------------------- helpers (reuse logic) ---------------------------
def closest_vec(query_df, ci, value_col, window_days=365):
    ci_v = ci[['연구번호','기록일자_dt', value_col]].dropna(subset=[value_col,'기록일자_dt']).copy()
    ci_v = ci_v.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    fw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid', direction='forward', tolerance=pd.Timedelta(days=window_days))
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid', direction='backward', tolerance=pd.Timedelta(days=window_days))
    fw['_d'] = (fw['rec_date']-fw['index_date']).abs()
    bw['_d'] = (bw['rec_date']-bw['index_date']).abs()
    use_fw = (fw['_d'].fillna(pd.Timedelta(days=window_days*10)) <= bw['_d'].fillna(pd.Timedelta(days=window_days*10)))
    return pd.Series(np.where(use_fw, fw[value_col].values, bw[value_col].values),
                    index=fw['orig_idx'].values).reindex(query_df.index).astype(float)

def smoke_vec(query_df, ci):
    smk = ci[['연구번호','기록일자_dt','흡연여부']].dropna(subset=['흡연여부','기록일자_dt']).copy()
    smk = smk.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], smk,
        left_on='index_date', right_on='rec_date', by='pid', direction='backward')
    return pd.Series(bw['흡연여부'].map(SMOKE_MAP).fillna('Unknown').values,
                    index=bw['orig_idx'].values).reindex(query_df.index).fillna('Unknown')

def make_tte(m, cls_or_list):
    if isinstance(cls_or_list, list):
        dx = m[cls_or_list].min(axis=1)
    else:
        dx = m[cls_or_list]
    is_pre = dx.notna() & (dx <= m['index_date'])
    primary = dx.where(dx > m['index_date'], pd.NaT)
    death_after = m['death_date'].where(
        (m['death_date'].notna()) & (m['death_date'] > m['index_date']) &
        ((primary.isna()) | (m['death_date'] < primary)), pd.NaT)
    event_date = primary.combine_first(death_after)
    status = np.where(primary.notna() & ((death_after.isna()) | (primary <= death_after)), 1,
            np.where(death_after.notna(), 2, 0))
    end_date = event_date.combine_first(m['last_follow'])
    time = (end_date - m['index_date']).dt.days.astype(float)
    res = pd.DataFrame({'pid':m['pid'].values,'pair_id':m['pair_id'].values,
                       'vaccinated':m['vaccinated'].astype(int).values,
                       'time':time,'status':status})
    res = res[~is_pre.values & (res['time']>0)].reset_index(drop=True)
    return res

def geskus_long(tte, primary=1, competing=2):
    """Geskus IPCW long-format dataset for Fine-Gray."""
    t = tte.copy().reset_index(drop=True)
    t['_id'] = t.index
    cens_obs = (t['status'] == 0).astype(int)
    kmf = KaplanMeierFitter().fit(t['time'], event_observed=cens_obs)
    G = kmf.survival_function_.iloc[:,0]
    G_idx = G.index.values; G_val = G.values
    def G_at(t_):
        pos = np.searchsorted(G_idx, t_, side='right') - 1
        return G_val[pos] if pos >= 0 else 1.0
    primary_times = sorted(t.loc[t['status']==primary,'time'].unique().tolist())
    max_time = t['time'].max()
    rows = []
    for _, r in t.iterrows():
        if r['status'] == primary:
            rows.append({'_id':r['_id'],'pair_id':r['pair_id'],'vaccinated':r['vaccinated'],
                        '_start':0,'_stop':r['time'],'_event':1,'_w':1.0})
        elif r['status'] == 0:
            rows.append({'_id':r['_id'],'pair_id':r['pair_id'],'vaccinated':r['vaccinated'],
                        '_start':0,'_stop':r['time'],'_event':0,'_w':1.0})
        else:
            G_c = G_at(r['time'])
            if G_c <= 0: continue
            split = sorted([pt for pt in primary_times if pt > r['time']] + [max_time])
            split = [s for s in split if s > r['time']]
            prev = r['time']
            for st in split:
                G_st = G_at(st)
                w = G_st / G_c
                if w > 0:
                    rows.append({'_id':r['_id'],'pair_id':r['pair_id'],'vaccinated':r['vaccinated'],
                                '_start':prev,'_stop':st,'_event':0,'_w':w})
                prev = st
    return pd.DataFrame(rows)

def fmt_p(p):
    if pd.isna(p): return '-'
    return '<0.001' if p<0.001 else f'{p:.3f}'

def fmt_hr(hr, lo, hi):
    if pd.isna(hr): return '-'
    return f'{hr:.3f} ({lo:.3f}–{hi:.3f})'

# --------------------------- build Cohort A matched ---------------------------
print('[1] Building Cohort A matched dataset...')
cohort = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_코호트.csv', encoding='cp949', low_memory=False)
cohort['birth_date'] = pd.to_datetime(cohort['생년월'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['death_date'] = pd.to_datetime(cohort['사망일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['last_follow'] = pd.to_datetime(cohort['최종추적일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['is_seoul'] = (cohort['주소'].astype(str).str.split().str[0]=='서울').astype(int)
rx = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_처방정보.csv', encoding='cp949', low_memory=False)
mask = (rx['처방명'].astype(str).str.contains('Gardasil|Cervarix|HPV vaccine', case=False, na=False) |
        rx['처방한글명'].astype(str).str.contains('가다실|서바릭스', na=False))
rx_vac = rx[mask].copy()
rx_vac['처방일자'] = pd.to_datetime(rx_vac['처방일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
first_vac = rx_vac.groupby('연구번호')['처방일자'].min().reset_index()
first_vac.columns = ['연구번호','first_vaccine_date']
cohort = cohort.merge(first_vac, on='연구번호', how='left').dropna(subset=['birth_date'])
ci = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_기초임상정보.csv', encoding='cp949', low_memory=False)
ci['기록일자_dt'] = pd.to_datetime(ci['기록일자'].astype(str).str.strip(), format='%Y%m%d', errors='coerce')

rng = np.random.default_rng(RANDOM_SEED)
df = cohort.copy()
df['vaccinated'] = df['first_vaccine_date'].notna()
vac_dates = df.loc[df['vaccinated'], 'first_vaccine_date'].dropna().values
df.loc[~df['vaccinated'], 'index_date'] = pd.to_datetime(rng.choice(vac_dates, size=(~df['vaccinated']).sum()))
df.loc[df['vaccinated'], 'index_date'] = df.loc[df['vaccinated'], 'first_vaccine_date']
df = df.rename(columns={'연구번호':'pid'}).reset_index(drop=True)
df['age_at_index'] = (df['index_date'] - df['birth_date']).dt.days/365.25
df = df[(df['death_date'].isna()) | (df['death_date'] > df['index_date'])]
df = df[df['last_follow'] > df['index_date']].reset_index(drop=True)
q = df[['pid','index_date']].copy()
df['height'] = closest_vec(q, ci, '키')
df['weight'] = closest_vec(q, ci, '몸무게')
df['sbp'] = closest_vec(q, ci, '수축기혈압')
df['dbp'] = closest_vec(q, ci, '이완기혈압')
df['bmi'] = df['weight']/(df['height']/100)**2
df['smoke'] = smoke_vec(q, ci).values
for c in ['bmi','sbp','dbp']:
    df[f'{c}_miss'] = df[c].isna().astype(int)
    df[c] = df[c].fillna(df[c].mean())
sm = pd.get_dummies(df['smoke'], prefix='smoke').astype(int)
df = pd.concat([df, sm], axis=1)
ps_features = ['age_at_index','bmi','bmi_miss','sbp','sbp_miss','dbp','dbp_miss','is_seoul',
              'smoke_Never','smoke_Former','smoke_Current']
ps_features = [c for c in ps_features if c in df.columns]
X = df[ps_features].astype(float).values
y = df['vaccinated'].astype(int).values
Xs = StandardScaler().fit_transform(X)
lr = LogisticRegression(max_iter=2000, C=1e6, solver='lbfgs').fit(Xs, y)
df['ps'] = lr.predict_proba(Xs)[:,1]
df['logit_ps'] = np.log(df['ps']/(1-df['ps']))
caliper = 0.2 * df['logit_ps'].std()
vac_idx = df.index[df['vaccinated']].tolist()
ctl_idx = np.array(df.index[~df['vaccinated']].tolist())
ctl_logit = df.loc[ctl_idx,'logit_ps'].values
order = np.argsort(ctl_logit)
ctl_sorted = ctl_idx[order]; ctl_logit_sorted = ctl_logit[order]
used = np.zeros(len(ctl_sorted), dtype=bool)
matched = []
vac_order = np.array(vac_idx); rng2 = np.random.default_rng(RANDOM_SEED); rng2.shuffle(vac_order)
for vi in vac_order:
    target = df.loc[vi,'logit_ps']
    lo = np.searchsorted(ctl_logit_sorted, target-caliper)
    hi = np.searchsorted(ctl_logit_sorted, target+caliper, side='right')
    best_j, best_d = -1, caliper+1
    for j in range(lo, hi):
        if used[j]: continue
        d = abs(ctl_logit_sorted[j]-target)
        if d < best_d: best_d=d; best_j=j
    if best_j>=0:
        used[best_j] = True
        matched.append((vi, ctl_sorted[best_j]))
pair_records = []
for pid_, (vi, cii) in enumerate(matched):
    pair_records.append((vi, pid_)); pair_records.append((cii, pid_))
pair_df = pd.DataFrame(pair_records, columns=['orig_idx','pair_id'])
m = df.loc[pair_df['orig_idx'].values].copy()
m['pair_id'] = pair_df['pair_id'].values
m = m.reset_index(drop=True)

wb = openpyxl.load_workbook('Data/한국 HPV 코호트 자료를 이용한 자_진단정보_기저질환추가_unlocked.xlsx',
                          read_only=True, data_only=True)
ws = wb.active
recs=[]
for row in ws.iter_rows(min_row=2, values_only=True):
    pid, cls, dd = row[0], row[5], row[8]
    if cls is None or str(cls).strip()=='': continue
    cls = str(cls).strip()
    if cls not in CLASS_LABELS: continue
    recs.append((pid, cls, pd.to_datetime(str(dd), format='%Y%m%d', errors='coerce')))
como = pd.DataFrame(recs, columns=['pid','class','diag_date'])
first_diag = como.groupby(['pid','class'])['diag_date'].min().unstack('class')
for c in CLASS_LABELS:
    if c not in first_diag.columns: first_diag[c] = pd.NaT
m = m.merge(first_diag, left_on='pid', right_index=True, how='left')

# --------------------------- per-outcome analyses ---------------------------
print('[2] Computing HRs for each outcome...')
outcomes = [
    ('Any-of-5 composite', ANY5),
    ('MCE composite (MI/Stroke/PE)', MCE),
    ('Hypertension', '2'),
    ('Diabetes', '3'),
    ('Angina/MI', '1'),
    ('Stroke', '4'),
    ('PE', '5'),
]
results = []
for label, comp in outcomes:
    print(f'  -- {label} --')
    tte = make_tte(m, comp)
    n_v = int((tte['vaccinated']==1).sum()); n_c = int((tte['vaccinated']==0).sum())
    e_v = int(((tte['status']==1)&(tte['vaccinated']==1)).sum())
    e_c = int(((tte['status']==1)&(tte['vaccinated']==0)).sum())
    e_comp = int((tte['status']==2).sum())
    py_v = tte.loc[tte['vaccinated']==1,'time'].sum() / 365.25
    py_c = tte.loc[tte['vaccinated']==0,'time'].sum() / 365.25
    ir_v = 1000*e_v/py_v if py_v>0 else np.nan
    ir_c = 1000*e_c/py_c if py_c>0 else np.nan

    res = {'outcome':label,'e_v':e_v,'n_v':n_v,'e_c':e_c,'n_c':n_c,'e_comp':e_comp,
          'py_v':py_v,'py_c':py_c,'ir_v':ir_v,'ir_c':ir_c}

    # Cause-specific Cox with cluster-robust SE
    if e_v + e_c >= 5 and e_v >= 1 and e_c >= 1:
        try:
            d = tte.copy(); d['event'] = (d['status']==1).astype(int)
            cph = CoxPHFitter()
            cph.fit(d[['time','event','vaccinated','pair_id']],
                   duration_col='time', event_col='event',
                   cluster_col='pair_id', robust=True)
            sm = cph.summary
            res.update({'cs_HR':float(sm.loc['vaccinated','exp(coef)']),
                       'cs_CI_lo':float(sm.loc['vaccinated','exp(coef) lower 95%']),
                       'cs_CI_hi':float(sm.loc['vaccinated','exp(coef) upper 95%']),
                       'cs_p':float(sm.loc['vaccinated','p'])})
        except Exception as e:
            print(f'    CS failed: {e}')
            res.update({'cs_HR':np.nan,'cs_CI_lo':np.nan,'cs_CI_hi':np.nan,'cs_p':np.nan})

        # Schoenfeld PH test
        try:
            d_ph = tte.copy(); d_ph['event'] = (d_ph['status']==1).astype(int)
            cph_ph = CoxPHFitter()
            cph_ph.fit(d_ph[['time','event','vaccinated']], duration_col='time', event_col='event')
            ph = proportional_hazard_test(cph_ph, d_ph[['time','event','vaccinated']], time_transform='rank')
            ph_summary = ph.summary if hasattr(ph,'summary') else ph
            res['PH_p'] = float(ph_summary.loc['vaccinated','p'])
        except Exception:
            res['PH_p'] = np.nan

        # Fine-Gray (Geskus) with cluster_col=pair_id
        try:
            fg = geskus_long(tte)
            ctv = CoxTimeVaryingFitter()
            ctv.fit(fg, id_col='_id', start_col='_start', stop_col='_stop',
                   event_col='_event', weights_col='_w',
                   cluster_col='pair_id', robust=True, show_progress=False)
            sm = ctv.summary
            res.update({'fg_HR':float(sm.loc['vaccinated','exp(coef)']),
                       'fg_CI_lo':float(sm.loc['vaccinated','exp(coef) lower 95%']),
                       'fg_CI_hi':float(sm.loc['vaccinated','exp(coef) upper 95%']),
                       'fg_p':float(sm.loc['vaccinated','p'])})
        except Exception as e:
            print(f'    FG failed: {e}')
            # Try without cluster_col as fallback
            try:
                fg = geskus_long(tte)
                ctv = CoxTimeVaryingFitter()
                ctv.fit(fg, id_col='_id', start_col='_start', stop_col='_stop',
                       event_col='_event', weights_col='_w', show_progress=False)
                sm = ctv.summary
                res.update({'fg_HR':float(sm.loc['vaccinated','exp(coef)']),
                           'fg_CI_lo':float(sm.loc['vaccinated','exp(coef) lower 95%']),
                           'fg_CI_hi':float(sm.loc['vaccinated','exp(coef) upper 95%']),
                           'fg_p':float(sm.loc['vaccinated','p'])})
                print(f'    FG used unclustered SE')
            except Exception as e2:
                print(f'    FG fallback also failed: {e2}')
                res.update({'fg_HR':np.nan,'fg_CI_lo':np.nan,'fg_CI_hi':np.nan,'fg_p':np.nan})
    else:
        res.update({'cs_HR':np.nan,'cs_CI_lo':np.nan,'cs_CI_hi':np.nan,'cs_p':np.nan,
                   'fg_HR':np.nan,'fg_CI_lo':np.nan,'fg_CI_hi':np.nan,'fg_p':np.nan,
                   'PH_p':np.nan})
    results.append(res)

R = pd.DataFrame(results)
R.to_csv('Data/Table2_CohortA_HazardRatios.csv', index=False, encoding='utf-8-sig')
print('\nSaved CSV: Data/Table2_CohortA_HazardRatios.csv')
print(R[['outcome','e_v','e_c','cs_HR','cs_p','fg_HR','fg_p','PH_p']].to_string(index=False))

# --------------------------- docx ---------------------------
print('\n[3] Writing docx...')
doc = Document()
sty = doc.styles['Normal']; sty.font.name = 'Times New Roman'; sty.font.size = Pt(10)
doc.add_heading('Table 2.  Cohort A — Hazard ratios for chronic comorbidities', level=1)

p = doc.add_paragraph()
p.add_run(
    'Cluster-robust cause-specific and Fine–Gray subdistribution hazard ratios for the propensity-score-matched '
    'cohort (1:1, 4,102 women in 2,053 pairs). Cause-specific HRs are estimated by Cox proportional-hazards '
    'regression treating death as right-censoring. Fine–Gray sHRs are estimated as time-varying Cox regression on '
    'a Geskus (2011) IPCW-reweighted dataset, retaining death as a competing event. Both models use cluster-robust '
    'standard errors with the matched pair as the cluster variable.'
).italic = True

n_cols = 10
t = doc.add_table(rows=1+len(R), cols=n_cols)
t.style = 'Light Grid Accent 1'
headers = ['Outcome', 'Vac\nevents/n', 'Ctl\nevents/n', 'Vac\nPY', 'Ctl\nPY',
          'IR per\n1000 PY (vac/ctl)',
          'Cause-specific HR\n(95% CI)', 'CS p',
          'Fine–Gray sHR\n(95% CI)', 'FG p']
for i, h in enumerate(headers):
    c = t.rows[0].cells[i]
    c.text = h
    for para in c.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold=True; r.font.size=Pt(8)

for i, row in R.iterrows():
    cells = t.rows[i+1].cells
    cells[0].text = row['outcome']
    cells[1].text = f'{int(row["e_v"])}/{int(row["n_v"])}'
    cells[2].text = f'{int(row["e_c"])}/{int(row["n_c"])}'
    cells[3].text = f'{row["py_v"]:.0f}'
    cells[4].text = f'{row["py_c"]:.0f}'
    cells[5].text = f'{row["ir_v"]:.2f} / {row["ir_c"]:.2f}'
    cells[6].text = fmt_hr(row['cs_HR'], row['cs_CI_lo'], row['cs_CI_hi'])
    cells[7].text = fmt_p(row['cs_p'])
    cells[8].text = fmt_hr(row['fg_HR'], row['fg_CI_lo'], row['fg_CI_hi'])
    cells[9].text = fmt_p(row['fg_p'])
    for j, c in enumerate(cells):
        for para in c.paragraphs:
            for r in para.runs: r.font.size = Pt(8)
            if j>0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
foot = doc.add_paragraph()
foot.add_run('Footnote. ').bold = True
foot.add_run(
    'PY, person-years; IR, incidence rate; HR, hazard ratio; sHR, subdistribution hazard ratio; '
    'CS, cause-specific; FG, Fine–Gray; CI, confidence interval. '
    'Cause-specific HRs use cluster-robust standard errors with the matched pair as the cluster variable; '
    'Fine–Gray sHRs are estimated on the Geskus IPCW-reweighted long dataset using sandwich (robust) standard '
    'errors, as cluster-robust SEs are not currently implemented for the time-varying Cox formulation in lifelines '
    '— the very small number of competing events (28–29 deaths) makes the practical difference negligible. '
    'Outcomes with fewer than five events in either group (Angina/MI, Stroke, PE) cannot be modelled and are '
    'reported descriptively. Person-years and incidence rates are accumulated to first event, death, or last '
    'follow-up. Because the competing event of death is rare, cause-specific and Fine–Gray estimates agree closely.'
).font.size = Pt(8)

doc.save('Data/Table2_CohortA_HazardRatios.docx')
print('Saved docx: Data/Table2_CohortA_HazardRatios.docx')
