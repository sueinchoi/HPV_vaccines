"""
Second pass: handle remaining reinfection mentions where the cell text is not an
exact match (e.g., "B – HPV reinfection", "HPV reinfection: Gardasil 9") and
prose paragraphs in csi files.

Strategy:
  - Replace "HPV reinfection" / "HPV-reinfection" as a SUBSTRING with
    "post-index hr-HPV detection (sensitivity)" inside the run that contains it.
  - This keeps surrounding text intact.

Note: The csi (author-tracked) files are slightly older snapshots; we only
update the labels for consistency, not the numerical results.
"""
from docx import Document

FILES = [
    'docs/HPV_manuscript_csi.docx',
    'docs/HPV_supplementary_csi.docx',
    'Data/Methodology_Revisions.docx',
]

SUBS = [
    # short replacements (run-level substring)
    ('HPV-reinfection', 'post-index hr-HPV detection'),
    ('HPV reinfection', 'post-index hr-HPV detection (sensitivity)'),
]

def replace_in_runs(paragraph):
    n = 0
    for run in paragraph.runs:
        for old, new in SUBS:
            if old in run.text:
                run.text = run.text.replace(old, new)
                n += 1
    return n

def replace_in_paragraph_fallback(paragraph):
    """If subs span runs, join, replace, and put back into first run."""
    if not paragraph.runs:
        return 0
    full = paragraph.text
    new = full
    for old, nw in SUBS:
        new = new.replace(old, nw)
    if new != full:
        for r in list(paragraph.runs):
            r.text = ''
        paragraph.runs[0].text = new
        return 1
    return 0

for f in FILES:
    doc = Document(f)
    n_par = 0
    for p in doc.paragraphs:
        # first try run-level (preserves formatting)
        n_par += replace_in_runs(p)
        # if substring still present (cross-run), fallback
        if any(s in p.text for s,_ in SUBS):
            n_par += replace_in_paragraph_fallback(p)
    n_cell = 0
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for para in c.paragraphs:
                    n_cell += replace_in_runs(para)
                    if any(s in para.text for s,_ in SUBS):
                        n_cell += replace_in_paragraph_fallback(para)
    doc.save(f)
    print(f'{f}: paragraph subs={n_par}, cell subs={n_cell}')
