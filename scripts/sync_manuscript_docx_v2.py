"""
Second-pass sync of HPV_manuscript.docx to match the slimmed-down supplementary
(now 12 tables). Specifically:
  - P36: 'Supplementary Table S7' -> 'Supplementary Table S6'
  - P48: 'Supplementary Table S8' -> 'Supplementary Table S7'
  - P53: rewrite sensitivity narrative to match Manuscript_Draft.md
  - P54 (landmark for post-index detection): replace with a brief
    transparency mention pointing to Sup Table S6 (cluster-robust + sensitivity rows)
  - P55: drop entire paragraph (S3 dropped)
  - Replace the supplementary materials list with the new 12-table structure
"""
from docx import Document
from copy import deepcopy

SRC = 'docs/HPV_manuscript.docx'

doc = Document(SRC)

# Targeted text replacements (run-level first, paragraph fallback)
TARGETED = [
    # P36 (Schoenfeld ref)
    ('reported in Supplementary Table S7 and Supplementary Figures S5–S6',
     'reported in Supplementary Table S6 and Supplementary Figures S5–S6'),
    # P48 (pseudo-index ref)
    ('three reasonable strategies (random sampling, calendar-year-matched, and risk-set sampling) yielded concordant Any-of-five hazard ratios (Supplementary Table S8)',
     'three reasonable strategies (random sampling, calendar-year-matched, and risk-set sampling) yielded concordant Any-of-five hazard ratios (Supplementary Table S7)'),
    ('Supplementary Table S8)',
     'Supplementary Table S7)'),
    # P53 rewrites — replace mentions of generic dropped supplementary tables
    ('(Supplementary Table on dose threshold)',
     '(Supplementary Table S8)'),
    ('(Supplementary Table on strict matching)',
     '(Supplementary Table S9)'),
    ('Supplementary Table on time-stratified clearance',
     'Supplementary Table S10'),
    ('Supplementary Table on single-negative clearance sensitivity',
     'Supplementary Table S11'),
    ('Supplementary Table on prescription-code validation',
     ''),  # this whole phrase is dropped; surrounding cleanup handled below
    ('(; supporting',
     ' supporting'),
    ('(concordance 100%; )',
     '(concordance 100%)'),
]

def apply_text_subs(p, subs):
    n = 0
    full = p.text
    new = full
    for old, nw in subs:
        if old in new:
            new = new.replace(old, nw)
    if new != full:
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = new
        else:
            p.add_run(new)
        n += 1
    return n

n_text = 0
for p in doc.paragraphs:
    n_text += apply_text_subs(p, TARGETED)

# Rewrite P53 sensitivity narrative to match Manuscript_Draft.md
new_p53 = (
    'Five pre-specified sensitivity analyses (Sens-A–E) defended the principal Cohort B inferences. '
    'For the clearance co-primary, the more permissive single-negative event definition produced an '
    'attenuated hazard ratio of 1.23 (95% CI 0.89–1.72, p = 0.21; Supplementary Table S11) compared with '
    '1.40 under the two-consecutive-negative primary. The pre-specified piecewise time-stratified clearance '
    'model — fit because the Schoenfeld test indicated PH violation (p = 0.028) — decomposed the average '
    'hazard ratio into 0.73 (0.37–1.41) at 0–6 months, 3.19 (1.44–7.09, p = 0.004) at 6–12 months, '
    '0.77 (0.22–2.72) at 12–24 months, and 4.20 (0.73–24.14) at ≥24 months (Supplementary Table S10), '
    'identifying a delayed-response window consistent with the timing of HPV-vaccine antibody maturation. '
    'Restricting the vaccinated arm to ≥2 and ≥3 (complete schedule) recorded doses with matched-set '
    'integrity preserved produced lesion-recurrence hazard ratios of 0.74 (0.39–1.39) and 0.61 (0.30–1.24), '
    'respectively (Supplementary Table S8). A strict 1:4 fine-matched alternative yielded a lesion-recurrence '
    'hazard ratio of 0.72 (0.36–1.41), close to the variable-ratio primary (Supplementary Table S9). '
    'Applying minimum disease-free intervals of 3, 6, and 12 months to lesion recurrence yielded hazard '
    'ratios of 0.88, 0.86, and 0.82, respectively (Supplementary Table S12). All five sensitivity analyses '
    'preserved the null inference for lesion recurrence and the directionally favourable but non-significant '
    'overall inference for clearance.'
)
for r in list(doc.paragraphs[53].runs):
    r.text = ''
doc.paragraphs[53].runs[0].text = new_p53

# Drop P54 (legacy landmark sensitivity for post-index detection — sup table S13 dropped)
p54 = doc.paragraphs[54]
p54._element.getparent().remove(p54._element)

# Drop P55 (sup table S3 dropped)
# Note: after removing P54, paragraph indices shift. The previous P55 is now P54.
p55 = doc.paragraphs[54]
p55._element.getparent().remove(p55._element)

# Rebuild the supplementary materials list (former P61 onwards)
NEW_SUPPLIST = [
    '- Supplementary Figure S1. Love plot — covariate balance before and after 1:1 propensity score matching (Cohort A). File: Data/SupFigS1_loveplot_cohortA.png.',
    '- Supplementary Figure S2. Love plot — covariate balance before and after fine variable-ratio (1:up-to-4) matching (Cohort B). File: Data/SupFigS2_loveplot_cohortB.png.',
    '- Supplementary Figure S3. Propensity-score density distributions, before and after matching. File: Data/SupFigS3_ps_density.png.',
    '- Supplementary Figure S4. Schoenfeld residual plots for Cohort A primary models (Any-of-5, Diabetes). Files: Data/PH_check_A_*.png.',
    '- Supplementary Figure S5. Schoenfeld residual plots for Cohort B co-primary models — lesion recurrence (Data/PH_check_B_has_recurrence.png) and hr-HPV clearance (Data/PH_check_B_clearance.png, fitted on the n = 292 pre-vaccine hr-HPV-positive subset).',
    '- Supplementary Figure S6. Pre-specified sensitivity analyses for Cohort B — five-panel summary forest plot (Sens-A through Sens-E). File: Data/SupFigS6_Sensitivity_Forest.png.',
    '- Supplementary Table S1. Pre-matching baseline characteristics. File: Data/Table1_BaselineCharacteristics_unified.docx (pre-matching blocks).',
    '- Supplementary Table S2. Propensity-score model coefficients (Cohort A). File: Data/SupTableS2_ps_coefficients.docx.',
    '- Supplementary Table S3. Age-stratified hazard ratios for Cohort B (lesion recurrence). File: Data/SupTableS4_revised_age_fu_forest.docx.',
    '- Supplementary Table S4. Number-at-risk tables. File: Data/SupTableS5_number_at_risk.docx.',
    '- Supplementary Table S5. Vaccine-type detailed results — pairwise + single-model interaction (LRT) with sensitivity rows. Files: Data/vaccine_type_analysis.csv, Data/CohortB_vaccine_interaction.csv.',
    '- Supplementary Table S6. Cluster-robust hazard ratios with PY, IR, and Schoenfeld p-values for both cohorts. File: Data/Methodology_Revisions.docx.',
    '- Supplementary Table S7. Pseudo-index assignment sensitivity (Cohort A). File: Data/CohortA_pseudoindex_sensitivity.csv.',
    '- Supplementary Table S8 (Sens-C). Dose-threshold sensitivity for both cohorts. File: Data/Sensitivity_DoseThreshold_HR.csv.',
    '- Supplementary Table S9 (Sens-D). Strict 1:4 fine-matching sensitivity for Cohort B lesion recurrence. File: Data/Sensitivity_StrictMatching.csv.',
    '- Supplementary Table S10 (Sens-B). Time-stratified hr-HPV clearance hazard ratios. File: Data/Sensitivity_HPV_Clearance_TimeStratified.csv.',
    '- Supplementary Table S11 (Sens-A). Single-negative HPV clearance sensitivity. File: Data/Sensitivity_HPV_Clearance_SingleNegative.csv.',
    '- Supplementary Table S12 (Sens-E). Disease-free-interval sensitivity for lesion recurrence. File: Data/Sensitivity_Recurrence_DFInterval.csv.',
]

# Find and replace the existing sup list section.
# The current section starts with a paragraph beginning "- Supplementary Figure S1." and runs through the last "Supplementary Table" line.
# Identify start and end indices.
start_idx = None
end_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if start_idx is None and t.startswith('- Supplementary Figure S1'):
        start_idx = i
    if t.startswith('- Supplementary Table S') or t.startswith('- Supplementary Figure S'):
        end_idx = i

print(f'Sup list region: paragraphs [{start_idx}, {end_idx}]')

# Remove the old list paragraphs in-place (reverse order so indices remain valid during deletion)
if start_idx is not None and end_idx is not None:
    paras = doc.paragraphs[start_idx:end_idx + 1]
    # rewrite first paragraph with the entire new list joined by newlines (single paragraph approach,
    # since this section was originally a single multi-line paragraph too per markdown formatting)
    first = paras[0]
    # remove all subsequent paragraphs of the old list
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    # rewrite first paragraph runs with joined new list (\n)
    for r in list(first.runs):
        r.text = ''
    first.runs[0].text = '\n'.join(NEW_SUPPLIST)

doc.save(SRC)
print(f'Sensitivity rewrites: {n_text}')
print(f'Saved: {SRC}')
