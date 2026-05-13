"""
Final relabel pass on supplementary docx — replace any cell text 'HPV reinfection'
with 'Post-index hr-HPV detection (sensitivity)' in tables S7 (Table 9) and S9
(Table 12, dose-threshold). These tables retain the legacy post-index detection
endpoint as a supplementary sensitivity, but the label must make clear that this
is the sensitivity, not the clearance co-primary.
"""
from docx import Document

SRC = 'docs/HPV_supplementary.docx'
doc = Document(SRC)

OLD = 'HPV reinfection'
NEW = 'Post-index hr-HPV detection (sensitivity)'

changed = 0
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            if c.text.strip() == OLD:
                # rewrite cell — clear and set first paragraph text
                # preserve cell style by editing the existing paragraph's runs
                for para in c.paragraphs:
                    for run in list(para.runs):
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = ''
                # Set new text in the first paragraph
                c.paragraphs[0].text = NEW
                changed += 1

doc.save(SRC)
print(f'Cells relabelled: {changed}')
