"""
Table 3 v3 — Cohort B co-primary results under ≥2-dose + 3-mo landmark.

Reads Data/CohortB_HR_v3.csv and Data/CohortB_SustainedClearance.csv (from
analyze_primary_v3.py) and assembles the manuscript-ready table.

Outputs:
  Data/Table3_CohortB_HR_v3.csv
  Data/Table3_CohortB_HR_v3.docx
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
HR_FILE = ROOT / 'Data' / 'CohortB_HR_v3.csv'
SC_FILE = ROOT / 'Data' / 'CohortB_SustainedClearance.csv'
OUT_CSV = ROOT / 'Data' / 'Table3_CohortB_HR_v3.csv'
OUT_DOCX = ROOT / 'Data' / 'Table3_CohortB_HR_v3.docx'

hr = pd.read_csv(HR_FILE, encoding='utf-8-sig')
sc = pd.read_csv(SC_FILE, encoding='utf-8-sig')


def hr_str(row):
    return f'{row["HR"]:.2f} ({row["CI_lower"]:.2f}–{row["CI_upper"]:.2f})'


def p_str(p):
    return '<0.001' if p < 0.001 else f'{p:.3f}'


rows = []

# Co-primary
for label, key in [
    ('Lesion recurrence (≥CIN2/HSIL+ or invasive carcinoma)',
     'P1 — Lesion recurrence (CIN2+); ≥2 dose + 3mo landmark'),
    ('hr-HPV clearance (two consecutive post-index negatives)',
     'P2 — hr-HPV clearance; ≥2 dose + 3mo landmark'),
]:
    r = hr[hr['analysis'] == key].iloc[0]
    rows.append(
        {
            'Outcome': label,
            'Vaccinated, n (events)': f'{r["n_vac"]} ({r["events_vac"]})',
            'Non-vaccinated, n (events)': f'{r["n_non"]} ({r["events_non"]})',
            'PY (vac / non)': f'{r["PY_vac"]:.0f} / {r["PY_non"]:.0f}',
            'IR per 1000 PY (vac / non)': f'{r["IR_vac_per1000PY"]:.1f} / {r["IR_non_per1000PY"]:.1f}',
            'HR (95% CI)': hr_str(r),
            'p': p_str(r['p']),
        }
    )

# Sustained clearance duration — KM-based, supplementary block at bottom of Table 3
rows.append(
    {
        'Outcome': '',
        'Vaccinated, n (events)': '',
        'Non-vaccinated, n (events)': '',
        'PY (vac / non)': '',
        'IR per 1000 PY (vac / non)': '',
        'HR (95% CI)': '',
        'p': '',
    }
)
rows.append(
    {
        'Outcome': 'Sustained clearance — KM analysis of reversion-free time '
                   'among patients with clearance event',
        'Vaccinated, n (events)': '',
        'Non-vaccinated, n (events)': '',
        'PY (vac / non)': '',
        'IR per 1000 PY (vac / non)': '',
        'HR (95% CI)': '',
        'p': '',
    }
)
for _, s in sc.iterrows():
    if s['group'].startswith('Log-rank'):
        rows.append(
            {
                'Outcome': '  Log-rank (vac vs non-vac, reversion-free)',
                'Vaccinated, n (events)': '',
                'Non-vaccinated, n (events)': '',
                'PY (vac / non)': '',
                'IR per 1000 PY (vac / non)': '',
                'HR (95% CI)': str(s['KM_median_sustained_years']),
                'p': str(s['KM_q75_years']).replace('p=', ''),
            }
        )
        continue
    rows.append(
        {
            'Outcome': f'  {s["group"]} — KM median sustained clearance (years)',
            'Vaccinated, n (events)': (
                f'{int(s["n_clearance_events"])} clearance / '
                f'{int(s["reversion_events"])} reversion / '
                f'{int(s["censored"])} censored' if s['group'] == 'Vaccinated' else ''
            ),
            'Non-vaccinated, n (events)': (
                f'{int(s["n_clearance_events"])} clearance / '
                f'{int(s["reversion_events"])} reversion / '
                f'{int(s["censored"])} censored' if s['group'] == 'Non-vaccinated' else ''
            ),
            'PY (vac / non)': '',
            'IR per 1000 PY (vac / non)': '',
            'HR (95% CI)': (
                f'{s["KM_median_sustained_years"]} '
                f'(Q25–Q75 {s["KM_q25_years"]}–{s["KM_q75_years"]} yr)'
            ),
            'p': '',
        }
    )

# Sensitivity block
rows.append(
    {
        'Outcome': '',
        'Vaccinated, n (events)': '',
        'Non-vaccinated, n (events)': '',
        'PY (vac / non)': '',
        'IR per 1000 PY (vac / non)': '',
        'HR (95% CI)': '',
        'p': '',
    }
)
rows.append(
    {
        'Outcome': 'Sensitivity (lesion recurrence): exposure-definition',
        'Vaccinated, n (events)': '',
        'Non-vaccinated, n (events)': '',
        'PY (vac / non)': '',
        'IR per 1000 PY (vac / non)': '',
        'HR (95% CI)': '',
        'p': '',
    }
)
for label, key in [
    ('  ≥1 dose, no landmark', 'Sens — Lesion recurrence; ≥1 dose, no landmark'),
    ('  ≥3 dose, no landmark', 'Sens — Lesion recurrence; ≥3 dose, no landmark'),
]:
    r = hr[hr['analysis'] == key].iloc[0]
    rows.append(
        {
            'Outcome': label,
            'Vaccinated, n (events)': f'{r["n_vac"]} ({r["events_vac"]})',
            'Non-vaccinated, n (events)': f'{r["n_non"]} ({r["events_non"]})',
            'PY (vac / non)': f'{r["PY_vac"]:.0f} / {r["PY_non"]:.0f}',
            'IR per 1000 PY (vac / non)': f'{r["IR_vac_per1000PY"]:.1f} / {r["IR_non_per1000PY"]:.1f}',
            'HR (95% CI)': hr_str(r),
            'p': p_str(r['p']),
        }
    )

tbl = pd.DataFrame(rows)
tbl.to_csv(OUT_CSV, index=False, encoding='utf-8-sig')
print(f'Wrote {OUT_CSV.relative_to(ROOT)}')

doc = Document()
doc.add_heading(
    'Table 3. Cohort B — co-primary outcomes under ≥2-dose + 3-mo landmark primary definition '
    '(age-adjusted Cox PH, cluster-robust SE on fine_match_id)',
    level=1,
)
t = doc.add_table(rows=len(tbl) + 1, cols=len(tbl.columns))
t.style = 'Light Grid Accent 1'
for j, h in enumerate(tbl.columns):
    cell = t.cell(0, j)
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(8.5)
for i, row in enumerate(tbl.itertuples(index=False), start=1):
    for j, val in enumerate(row):
        c = t.cell(i, j)
        c.text = str(val)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

note = doc.add_paragraph()
note.add_run(
    '\nPrimary exposure: ≥2 distinct HPV-vaccine prescription dates. '
    '3-month landmark applied symmetrically across arms: index date shifted to index + 90 days; '
    'patients with <90 days follow-up or with outcome event in first 90 days excluded; '
    'matched non-vaccinated controls of excluded vaccinated cases dropped (matched-set integrity). '
    'For lesion recurrence, HR < 1 favours vaccination; for hr-HPV clearance, HR > 1 favours vaccination. '
    'Sustained clearance duration measured from first clearance event (first of two consecutive negatives) '
    'to first subsequent HR-HPV+ test (reversion) or last follow-up, whichever first. '
    'Reversion = any HR-HPV+ molecular pathology after the clearance event.'
).font.size = Pt(8)

doc.save(OUT_DOCX)
print(f'Wrote {OUT_DOCX.relative_to(ROOT)}')
