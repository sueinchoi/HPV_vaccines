"""
Rebuild Table 3 — Cohort B co-primary outcomes and key sensitivity analyses.

The table reports for each outcome:
  events / N for vaccinated and non-vaccinated arms
  person-years and incidence rate per 1,000 person-years
  age-adjusted Cox HR (95% CI) with cluster-robust SE on fine_match_id
  p value
  favourable direction (HR < 1 or HR > 1)

Outputs:
  Data/Table3_CohortB_HR.csv
  Data/Table3_CohortB_HR.docx
"""
import warnings; warnings.filterwarnings('ignore')
import sys, pandas as pd, numpy as np
sys.path.insert(0, 'scripts')
from extract_pathology_outcomes import detect_high_risk_hpv
from lifelines import CoxPHFitter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH_FILE = 'Data/한국 HPV 코호트 자료를 이용한 자_병리검사 (복구됨).CSV'

# ---------- Load ----------
patho = pd.read_csv(PATH_FILE, encoding='cp949', low_memory=False)
hpv = patho[patho['병리검사구분'].isin(['분자병리','HPV'])].copy()
hpv['실시일자_dt'] = pd.to_datetime(
    hpv['실시일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')

B = pd.read_csv('Data/final_matched_cohort.csv', encoding='utf-8-sig')
Bo = pd.read_csv('Data/final_matched_outcomes.csv', encoding='utf-8-sig')
B = B.merge(Bo[['연구번호','has_recurrence','recurrence_date','days_to_recurrence',
                'has_hpv_infection','hpv_infection_date','days_to_hpv','hpv_types']],
            on='연구번호')
B['index_date']  = pd.to_datetime(B['index_date'])
B['최종추적일자'] = pd.to_datetime(B['최종추적일자'])
B['recurrence_date']    = pd.to_datetime(B['recurrence_date'], errors='coerce')
B['hpv_infection_date'] = pd.to_datetime(B['hpv_infection_date'], errors='coerce')
B['vac']        = B['접종여부'].astype(bool).astype(int)
B['index_age']  = pd.to_numeric(B['index_age'], errors='coerce')
B['follow_up_days']     = (B['최종추적일자'] - B['index_date']).dt.days
B['days_to_recurrence'] = pd.to_numeric(B['days_to_recurrence'], errors='coerce')
B['days_to_hpv']        = pd.to_numeric(B['days_to_hpv'], errors='coerce')

# Pre-index baseline type set (pre-vaccine baseline)
hpv_b = hpv[hpv['연구번호'].isin(B['연구번호'])].copy()
hpv_b = hpv_b.merge(B[['연구번호','index_date']], on='연구번호')
pre = hpv_b[hpv_b['실시일자_dt'] < hpv_b['index_date']].copy()
pre['detect'] = pre['판독결과'].apply(detect_high_risk_hpv)
pre_summary = pre.groupby('연구번호').apply(lambda g: pd.Series({
    'pre_pos_any': any(r['is_high_risk_hpv_positive'] for r in g['detect']),
    'pre_types':   set().union(*[set(t for t in r['detected_hpv_types']
                                     if isinstance(t, int)) for r in g['detect']]),
})).reset_index()
pre_summary['pre_16_pos'] = pre_summary['pre_types'].apply(lambda s: int(16 in s))
pre_summary['pre_18_pos'] = pre_summary['pre_types'].apply(lambda s: int(18 in s))
B = B.merge(pre_summary, on='연구번호', how='left')

# Post-index molecular records (for clearance and novel-type)
post = hpv_b[hpv_b['실시일자_dt'] > hpv_b['index_date']].copy()
post['detect']     = post['판독결과'].apply(detect_high_risk_hpv)
post['post_pos']   = post['detect'].apply(lambda r: r['is_high_risk_hpv_positive'])
post['post_types'] = post['detect'].apply(
    lambda r: set(t for t in r['detected_hpv_types'] if isinstance(t, int)))
post = post.sort_values(['연구번호','실시일자_dt'])

# PRIMARY clearance event = first of two consecutive negative records
def first_two_consecutive(g, neg_predicate):
    g = g.sort_values('실시일자_dt').reset_index(drop=True)
    flags = neg_predicate(g)
    for i in range(len(g) - 1):
        if flags.iloc[i] and flags.iloc[i+1]:
            return g.loc[i, '실시일자_dt']
    return None

print('Computing two-consecutive-negative dates...')
two_neg = post.groupby('연구번호').apply(
    lambda g: first_two_consecutive(g, lambda gg: ~gg['post_pos'])
).dropna().rename('first_neg_date').reset_index()
B = B.merge(two_neg, on='연구번호', how='left')

# Type-specific (16, 18) two-consecutive negatives
two_neg16 = post.groupby('연구번호').apply(
    lambda g: first_two_consecutive(g, lambda gg: ~gg['post_types'].apply(lambda s: 16 in s))
).dropna().rename('first_neg16_date').reset_index()
B = B.merge(two_neg16, on='연구번호', how='left')

two_neg18 = post.groupby('연구번호').apply(
    lambda g: first_two_consecutive(g, lambda gg: ~gg['post_types'].apply(lambda s: 18 in s))
).dropna().rename('first_neg18_date').reset_index()
B = B.merge(two_neg18, on='연구번호', how='left')

# Novel-type
post_by_pid = {pid: g for pid, g in post.groupby('연구번호')}
def first_novel(row):
    if row['연구번호'] not in post_by_pid: return None
    pre_set = row['pre_types'] if isinstance(row['pre_types'], set) else set()
    for _, r in post_by_pid[row['연구번호']].iterrows():
        if r['post_types'] - pre_set:
            return r['실시일자_dt']
    return None
B['first_novel_date'] = B.apply(first_novel, axis=1)
B['first_novel_date'] = pd.to_datetime(B['first_novel_date'])

# ---------- Helpers ----------
def restrict(B_full, vac_cond):
    keep_ids = set(B_full.loc[(B_full['vac']==1) & vac_cond(B_full), 'fine_match_id'])
    sub = B_full[B_full['fine_match_id'].isin(keep_ids)].copy()
    sub = sub[(sub['vac']==1) | vac_cond(sub)]
    return sub


def fit_outcome(df, event_date_col, label, favourable):
    """Compute events, person-years, incidence rate, HR (95% CI), p."""
    d = df.copy()
    d['has_event']  = d[event_date_col].notna()
    d['days_to_ev'] = (d[event_date_col] - d['index_date']).dt.days
    d['time']  = np.where(d['has_event'], d['days_to_ev'], d['follow_up_days'])
    d['event'] = d['has_event'].astype(int)
    d = d[d['time'] > 0].dropna(subset=['index_age','fine_match_id'])
    out = {'outcome': label, 'favourable': favourable}
    for grp_val, key in [(1,'vac'), (0,'ctl')]:
        sub = d[d['vac']==grp_val]
        out[f'n_{key}']   = int(len(sub))
        out[f'ev_{key}']  = int(sub['event'].sum())
        out[f'py_{key}']  = float(sub['time'].sum() / 365.25)
        out[f'ir_{key}']  = (out[f'ev_{key}'] / out[f'py_{key}'] * 1000) if out[f'py_{key}'] > 0 else np.nan
    out.update(HR=np.nan, CIlo=np.nan, CIhi=np.nan, p=np.nan)
    if d['event'].sum() < 5: return out
    try:
        cph = CoxPHFitter().fit(
            d[['time','event','vac','index_age','fine_match_id']],
            duration_col='time', event_col='event',
            cluster_col='fine_match_id', robust=True)
        r = cph.summary.loc['vac']
        out.update(HR=float(r['exp(coef)']),
                   CIlo=float(r['exp(coef) lower 95%']),
                   CIhi=float(r['exp(coef) upper 95%']),
                   p=float(r['p']))
    except Exception as e:
        print(f'  fit failed for {label}: {e}')
    return out


# ---------- Build the Table 3 rows ----------
print('Building Table 3 rows...')
rows = []

# === CO-PRIMARY ===
# 1. Lesion recurrence (full Cohort B; HR < 1 favourable)
rows.append(fit_outcome(B, 'recurrence_date',
    'Lesion recurrence (≥CIN2 / HSIL+ or invasive carcinoma)', 'HR<1'))

# 2. Any hr-HPV clearance among pre-vaccine HPV+
sub_clear = restrict(B, lambda d: d['pre_pos_any'] == True)
rows.append(fit_outcome(sub_clear, 'first_neg_date',
    'hr-HPV clearance / regression (pre-vaccine HPV+ baseline)', 'HR>1'))

# === KEY SENSITIVITY ===
# 3. Post-index hr-HPV detection (full cohort, persistence + acquisition)
B_full_hpv = B.copy()
rows.append(fit_outcome(B_full_hpv, 'hpv_infection_date',
    'Post-index hr-HPV detection (sensitivity, full cohort)', 'HR<1'))

# 4. Novel-type acquisition
sub_novel = restrict(B, lambda d: d['pre_pos_any'].notna())
rows.append(fit_outcome(sub_novel, 'first_novel_date',
    'Novel-type acquisition (post-index type not in pre-vaccine baseline)', 'HR<1'))

# 5. HPV-16 clearance (pre-vaccine 16+ baseline)
sub16 = restrict(B, lambda d: d['pre_16_pos'] == 1)
rows.append(fit_outcome(sub16, 'first_neg16_date',
    'HPV-16 clearance (pre-vaccine 16+ baseline)', 'HR>1'))

# 6. HPV-18 clearance (pre-vaccine 18+ baseline)
sub18 = restrict(B, lambda d: d['pre_18_pos'] == 1)
rows.append(fit_outcome(sub18, 'first_neg18_date',
    'HPV-18 clearance (pre-vaccine 18+ baseline)', 'HR>1'))

# Format and save CSV
def fmt_hr(r):
    if np.isnan(r['HR']): return '—'
    return f"{r['HR']:.2f} ({r['CIlo']:.2f}–{r['CIhi']:.2f})"

def fmt_p(r):
    return '—' if np.isnan(r['p']) else f'{r["p"]:.3f}'

out_rows = []
for r in rows:
    out_rows.append({
        'Outcome': r['outcome'],
        'Favourable direction': r['favourable'],
        'Vaccinated events / N': f"{r['ev_vac']} / {r['n_vac']}",
        'Vaccinated PY (rate /1000 PY)': f"{r['py_vac']:.0f} ({r['ir_vac']:.1f})",
        'Non-vaccinated events / N': f"{r['ev_ctl']} / {r['n_ctl']}",
        'Non-vaccinated PY (rate /1000 PY)': f"{r['py_ctl']:.0f} ({r['ir_ctl']:.1f})",
        'HR (95% CI)': fmt_hr(r),
        'p value': fmt_p(r),
    })
table3 = pd.DataFrame(out_rows)
table3.to_csv('Data/Table3_CohortB_HR.csv', index=False, encoding='utf-8-sig')
print('Saved: Data/Table3_CohortB_HR.csv')

# ---------- DOCX ----------
print('Building Table 3 docx...')
doc = Document()
section = doc.sections[0]
section.page_width  = Cm(29.7); section.page_height = Cm(21.0)  # landscape A4
section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.5); section.right_margin = Cm(1.5)

title = doc.add_paragraph()
run = title.add_run('Table 3. Cohort B — Co-primary outcomes and key sensitivity analyses')
run.bold = True; run.font.size = Pt(11)

subtitle = doc.add_paragraph()
sr = subtitle.add_run(
    'Age-adjusted Cox proportional-hazards models with cluster-robust standard errors '
    'using the fine-matching identifier. PY = person-years; rate is per 1,000 person-years.')
sr.italic = True; sr.font.size = Pt(9)

cols = ['Outcome', 'Vaccinated events / N', 'Vaccinated PY (rate)',
        'Non-vaccinated events / N', 'Non-vaccinated PY (rate)',
        'HR (95% CI)', 'Favourable', 'p']
table = doc.add_table(rows=1, cols=len(cols))
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, c in enumerate(cols):
    p = hdr[i].paragraphs[0]; r = p.add_run(c); r.bold = True
    r.font.size = Pt(9)

co_primary_count = 2  # first two rows are co-primary
for idx, r in enumerate(rows):
    row = table.add_row().cells
    label = r['outcome']
    if idx < co_primary_count:
        label = '[Co-primary] ' + label
    else:
        label = '[Sensitivity] ' + label
    cells = [label,
             f"{r['ev_vac']} / {r['n_vac']}",
             f"{r['py_vac']:.0f} ({r['ir_vac']:.1f})",
             f"{r['ev_ctl']} / {r['n_ctl']}",
             f"{r['py_ctl']:.0f} ({r['ir_ctl']:.1f})",
             fmt_hr(r),
             r['favourable'],
             fmt_p(r)]
    for j, txt in enumerate(cells):
        para = row[j].paragraphs[0]
        run_ = para.add_run(str(txt))
        run_.font.size = Pt(9)
        if idx < co_primary_count and j == 0:
            run_.bold = True

# Footnote
fn = doc.add_paragraph()
fr = fn.add_run(
    'Notes: Co-primary outcomes are reported first; sensitivity analyses follow. '
    'For lesion recurrence and post-index detection / novel-type acquisition, '
    'a hazard ratio less than 1 is the favourable direction (vaccine reduces event rate); '
    'for clearance / regression outcomes, a hazard ratio greater than 1 is favourable '
    '(vaccine accelerates clearance). The clearance and type-specific clearance subsets '
    'are restricted to women with documented pre-vaccine high-risk HPV positivity '
    '(or, for type-specific variants, pre-vaccine type-specific positivity); for these '
    'analyses, matched-set integrity was preserved by also dropping non-vaccinated '
    'controls whose matched vaccinated case lacked the qualifying pre-vaccine status. '
    'The lesion recurrence outcome includes any post-index biopsy diagnosis of CIN2, '
    'CIN3, HSIL, carcinoma in situ, or invasive cervical cancer (i.e. CIN2 or worse / '
    'high-grade squamous intraepithelial lesion under the 2014 LAST nomenclature).'
)
fr.italic = True; fr.font.size = Pt(8)

doc.save('Data/Table3_CohortB_HR.docx')
print('Saved: Data/Table3_CohortB_HR.docx')
