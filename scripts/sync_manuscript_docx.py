"""
Sync HPV_manuscript.docx with Manuscript_Draft.md edits:
1) P37 (Methods, multiple-comparison para): replace 'reinfection' framing with 'clearance' framing.
2) P61: update Supplementary Figure S5 description (clearance instead of reinfection).
3) Drop "Recommended graphical abstract" paragraphs (P62 onwards through the empty line / next heading).
4) Discussion paragraphs: replace remaining 'HPV reinfection' / 'Gardasil-specific signal for HPV reinfection' with the clearance-framed text.
"""
from docx import Document
from copy import deepcopy
import re

SRC = 'docs/HPV_manuscript.docx'
OUT = 'docs/HPV_manuscript.docx'

doc = Document(SRC)

# Read updated Manuscript_Draft.md to get target text for the methods/discussion paragraphs
with open('docs/Manuscript_Draft.md', 'r') as f:
    md = f.read()

# Extract target paragraphs from markdown by anchor strings
def extract_between(md, start, end):
    s = md.find(start)
    if s < 0:
        return None
    e = md.find(end, s + len(start))
    return md[s:e] if e > 0 else md[s:]

# Target text 1: multiple-comparison paragraph (Methods)
mc_para = extract_between(
    md,
    'In line with the prevailing convention in observational epidemiology,',
    '\n\nTo assess robustness'
)
if mc_para:
    mc_para = mc_para.strip()

# Target text 2: Sup Fig S5 description
sup_s5 = ('Supplementary Figure S5. Schoenfeld residual plots for Cohort B co-primary models — '
          'lesion recurrence (Data/PH_check_B_has_recurrence.png) and hr-HPV clearance '
          '(Data/PH_check_B_clearance.png, fitted on the n = 292 pre-vaccine hr-HPV-positive '
          'subset with the two-consecutive-negative event).')

# Target text 3: Discussion vaccine-type interaction first sentence
vt_intro = extract_between(
    md,
    'The vaccine-type interaction analysis on the hr-HPV clearance co-primary',
    '. We report this finding transparently but interpret it as most plausibly artefactual'
)

# Target text 4: Limitations - replace whole Gardasil-specific signal sentence + reinfection mention
limits_replacement = extract_between(
    md,
    'Second, the original variable-ratio (1:up-to-4) fine-matched Cohort B subgroup analysis by vaccine type',
    '. The subgroup and interaction findings'
)

# Walk paragraphs and edit in place
changed = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if not t.strip():
        continue
    new_t = t
    # Methods multi-comparison para
    if 'no formal multiple-comparison adjustment' in t and 'reinfection' in t:
        if mc_para:
            new_t = mc_para
    # Sup Fig S5 line
    if 'Supplementary Figure S5.' in t and 'HPV reinfection' in t:
        new_t = re.sub(r'Supplementary Figure S5\..*?(?=- Supplementary Figure S6|$)',
                       '- ' + sup_s5 + '\n', t, flags=re.S)
    # Generic reinfection -> clearance fallback (only for label-level mentions)
    new_t = new_t.replace('high-risk HPV reinfection', 'hr-HPV clearance')
    new_t = new_t.replace('HPV reinfection', 'hr-HPV clearance')
    new_t = new_t.replace('Gardasil-specific signal for hr-HPV clearance',
                          'Gardasil-quadrivalent signal in the supplementary post-index hr-HPV '
                          'detection sensitivity (not the co-primary clearance outcome)')
    if new_t != t:
        # rewrite paragraph runs
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = new_t
        else:
            p.add_run(new_t)
        changed += 1

# Remove "Recommended graphical abstract" heading + the two body paragraphs that follow
# Stop at: References / a heading-styled paragraph / a numbered reference line
to_remove = []
in_abstract_block = False
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    style_name = getattr(p.style, 'name', '') or ''
    if t.lower().startswith('recommended graphical abstract'):
        in_abstract_block = True
        to_remove.append(p)
        continue
    if in_abstract_block:
        # STOP at: heading-styled paragraph, "References", or a numbered reference (e.g., "1. ...")
        if ('Heading' in style_name
            or t.lower().startswith('references')
            or re.match(r'^\d+\.\s+[A-Z]', t)
            or t.lower().startswith('suggested discussion outline')):
            in_abstract_block = False
            break
        to_remove.append(p)

for p in to_remove:
    p._element.getparent().remove(p._element)

print(f'Paragraphs edited: {changed}')
print(f'Paragraphs removed (graphical abstract): {len(to_remove)}')

doc.save(OUT)
print(f'Saved: {OUT}')
