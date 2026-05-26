"""Rebuild the three docx deliverables with embedded figures and real tables.

Outputs:
  docs/HPV_manuscript.docx       — pandoc-generated narrative (Abstract..References)
  docs/HPV_tables_figures.docx   — Figure 1-4 (images) + Table 1-3 (data tables)
  docs/HPV_supplementary.docx    — Sup Fig S1-S6 (images) + Sup Table S1-S12 (data)
"""

from __future__ import annotations

import csv
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

ROOT = Path(__file__).resolve().parent.parent
SRC_MD = ROOT / "docs" / "Manuscript_Draft.md"
DATA = ROOT / "Data"
DOCS = ROOT / "docs"

# ---------------------------------------------------------------------------
# Markdown narrative (pandoc) → HPV_manuscript.docx
# ---------------------------------------------------------------------------


def split_sections(text: str) -> dict[str, str]:
    lines = text.splitlines(keepends=True)

    def h2(needle: str) -> int:
        for i, ln in enumerate(lines):
            if ln.startswith("## ") and needle in ln:
                return i
        raise SystemExit(f"could not find heading: {needle}")

    title_end = next(i for i, ln in enumerate(lines) if ln.startswith("## "))
    tab_fig_h2 = h2("Tables and Figures")
    discussion_h2 = h2("Suggested Discussion outline")

    return {
        "manuscript": (
            "".join(lines[:title_end])
            + "".join(lines[title_end:tab_fig_h2])
            + "".join(lines[discussion_h2:])
        )
    }


def normalise(md: str) -> str:
    md = re.sub(r"\s*File:\s*`[^`]+`\.", ".", md)
    md = re.sub(r"\s*Files:\s*`[^`]+`(,\s*`[^`]+`)*\.", ".", md)
    md = re.sub(r"\s*Sources?:\s*`[^`]+`(,\s*`[^`]+`)*[^.]*\.", ".", md)
    return md


def run_pandoc(md: str, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    proc = subprocess.run(
        ["pandoc", "--from", "markdown", "--to", "docx", "--standalone",
         "--output", str(out_path)],
        input=md, text=True, capture_output=True, check=False,
    )
    if proc.returncode != 0:
        sys.stderr.write(proc.stderr)
        raise SystemExit(f"pandoc failed for {out_path.name}")
    print(f"wrote {out_path.relative_to(ROOT)}")


def build_manuscript_docx(suffix: str = "") -> None:
    text = SRC_MD.read_text(encoding="utf-8")
    parts = split_sections(text)
    run_pandoc(normalise(parts["manuscript"]),
               DOCS / f"HPV_manuscript{suffix}.docx")


# ---------------------------------------------------------------------------
# python-docx helpers for figures + tables
# ---------------------------------------------------------------------------


def read_csv(path: Path) -> tuple[list[str], list[list[str]]]:
    with path.open(encoding="utf-8-sig") as f:
        rdr = csv.reader(f)
        rows = list(rdr)
    if not rows:
        return [], []
    return rows[0], rows[1:]


# ---------------------------------------------------------------------------
# Translation + formatting helpers
# ---------------------------------------------------------------------------

KOR2ENG = {
    "협심증/심근경색": "Angina / Myocardial infarction",
    "고혈압": "Hypertension",
    "당뇨": "Diabetes",
    "뇌출혈/뇌경색": "Stroke",
    "폐색전증": "Pulmonary embolism",
    "원추절제술": "Conization",
    "자궁절제술": "Hysterectomy",
}


def translate_korean(s: str) -> str:
    out = s
    # Match Korean span (trim leading spaces but preserve them in output)
    leading_ws = len(out) - len(out.lstrip())
    body = out.strip()
    if body in KOR2ENG:
        return " " * leading_ws + KOR2ENG[body]
    return out


def fmt_num(v: str, decimals: int = 2) -> str:
    """Round a numeric string to `decimals` places; pass through non-numerics."""
    if v in ("", None):
        return ""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return v
    return f"{f:.{decimals}f}"


def fmt_p(v: str) -> str:
    if v in ("", None):
        return ""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return v
    if f < 0.001:
        return "<0.001"
    return f"{f:.3f}"


def fmt_hr_triplet(hr: str, lo: str, hi: str) -> str:
    try:
        return f"{float(hr):.2f} ({float(lo):.2f}–{float(hi):.2f})"
    except (TypeError, ValueError):
        return "—"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_caption(doc: Document, label: str, caption: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(". " + caption)


def add_image(doc: Document, image_path: Path, width_in: float = 6.5) -> None:
    if not image_path.exists():
        doc.add_paragraph(f"[Missing image: {image_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_in))


def add_table(doc: Document, header: list[str], body: list[list[str]],
              col_widths_in: list[float] | None = None) -> None:
    """Render a single docx table with bold header, thin borders."""
    if not header:
        return
    tbl = doc.add_table(rows=1 + len(body), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False

    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            val = row[j] if j < len(row) else ""
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    if col_widths_in:
        for row in tbl.rows:
            for cell, w in zip(row.cells, col_widths_in):
                cell.width = Inches(w)


def add_spacer(doc: Document) -> None:
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# CSV-derived table loaders (returns header, body)
# ---------------------------------------------------------------------------


def table1_rows() -> tuple[list[str], list[list[str]]]:
    """Reshape Table1_BaselineCharacteristics_unified.csv (long → wide-ish)."""
    header, body = read_csv(DATA / "Table1_BaselineCharacteristics_unified.csv")
    keep = {"CohortA_post", "CohortB_post", "CohortB_post_v3", "CohortB_clearance"}
    out_header = ["Block", "Variable", "Vaccinated", "Non-vaccinated", "p value", "|SMD|"]
    out = []
    for r in body:
        if len(r) < 6:
            continue
        block, var, vac, ctl, p, smd = r[0], r[1], r[2], r[3], r[4], r[5]
        if block not in keep:
            continue
        block_label = {
            "CohortA_post": "Cohort A (post-PSM)",
            "CohortB_post": "Cohort B (post fine match — ≥1 dose, no landmark; legacy)",
            "CohortB_post_v3": "Cohort B (≥2 dose + 3-mo landmark; PRIMARY)",
            "CohortB_clearance": "Cohort B — clearance subset (≥1 dose, no landmark)",
        }[block]
        out.append([block_label, translate_korean(var), vac, ctl, p, smd])
    return out_header, out


def table1_split_rows() -> dict[str, list[list[str]]]:
    """Return {block_id: rows_without_block_column} for the four Table 1 splits."""
    _h, body = read_csv(DATA / "Table1_BaselineCharacteristics_unified.csv")
    out = {"CohortA_post": [], "CohortB_post": [], "CohortB_post_v3": [], "CohortB_clearance": []}
    for r in body:
        if len(r) < 6 or r[0] not in out:
            continue
        out[r[0]].append([translate_korean(r[1]), r[2], r[3], r[4], r[5]])
    return out


def table2_rows() -> tuple[list[str], list[list[str]]]:
    """Cohort A HR results: cause-specific + Fine-Gray columns."""
    header, body = read_csv(DATA / "Table2_CohortA_HazardRatios.csv")
    out_header = ["Outcome", "Events Vac / N", "Events Ctl / N",
                  "IR Vac (per 1,000 PY)", "IR Ctl",
                  "Cause-specific HR (95% CI)", "p", "Fine–Gray HR (95% CI)", "p"]

    def fmt_hr(hr: str, lo: str, hi: str) -> str:
        try:
            return f"{float(hr):.2f} ({float(lo):.2f}–{float(hi):.2f})"
        except ValueError:
            return "—"

    def fmt_p(p: str) -> str:
        try:
            v = float(p)
            return f"{v:.3f}" if v >= 0.001 else "<0.001"
        except ValueError:
            return "—"

    def fmt_rate(r: str) -> str:
        try:
            return f"{float(r):.2f}"
        except ValueError:
            return "—"

    out = []
    for r in body:
        if len(r) < 19:
            continue
        (outcome, e_v, n_v, e_c, n_c, _ec, _pyv, _pyc, ir_v, ir_c,
         cs_hr, cs_lo, cs_hi, cs_p, _ph,
         fg_hr, fg_lo, fg_hi, fg_p) = r[:19]
        out.append([
            outcome,
            f"{e_v} / {n_v}",
            f"{e_c} / {n_c}",
            fmt_rate(ir_v),
            fmt_rate(ir_c),
            fmt_hr(cs_hr, cs_lo, cs_hi),
            fmt_p(cs_p),
            fmt_hr(fg_hr, fg_lo, fg_hi),
            fmt_p(fg_p),
        ])
    return out_header, out


def table3_rows() -> tuple[list[str], list[list[str]]]:
    header, body = read_csv(DATA / "Table3_CohortB_HR.csv")
    return header, body


# ---------- Supplementary table loaders ----------


def csv_pretty(path: Path, columns: list[tuple[str, str]] | None = None,
               num_cols: list[str] | None = None,
               decimals: int = 2) -> tuple[list[str], list[list[str]]]:
    """Generic CSV reader that optionally renames+filters columns and formats
    numeric columns to `decimals` places."""
    header, body = read_csv(path)
    if columns is None:
        col_idx = list(range(len(header)))
        out_header = header
    else:
        wanted = [c[0] for c in columns]
        col_idx = []
        for w in wanted:
            if w in header:
                col_idx.append(header.index(w))
            else:
                col_idx.append(None)
        out_header = [c[1] for c in columns]

    out = []
    for r in body:
        row = []
        for ci, w in zip(col_idx, (columns or [(h, h) for h in header])):
            if ci is None or ci >= len(r):
                row.append("")
                continue
            val = r[ci]
            if num_cols and w[0] in num_cols and val:
                try:
                    row.append(f"{float(val):.{decimals}f}")
                except ValueError:
                    row.append(val)
            else:
                row.append(val)
        out.append(row)
    return out_header, out


def hr_table(path: Path, label_col: str, extra_cols: list[tuple[str, str]]
             ) -> tuple[list[str], list[list[str]]]:
    """Build a compact HR table: label + event counts + HR (95% CI) + p."""
    header, body = read_csv(path)

    def col(name: str, default: str | None = None) -> int | None:
        if name in header:
            return header.index(name)
        return None

    li = col(label_col)
    n_v = col("n_v") or col("N_vac")
    n_c = col("n_c") or col("N_ctl")
    ev_v = col("ev_v") or col("events_vac")
    ev_c = col("ev_c") or col("events_ctl")
    hr = col("HR") or col("hr") or col("cs_HR")
    lo = col("CIlo") or col("CI_lo") or col("cs_CI_lo")
    hi = col("CIhi") or col("CI_hi") or col("cs_CI_hi")
    p = col("p") or col("cs_p")

    out_header = [label_col]
    for src, dst in extra_cols:
        out_header.append(dst)
    out_header += ["Vac events / N", "Ctl events / N", "HR (95% CI)", "p"]

    out = []
    for r in body:
        row = [r[li] if li is not None else ""]
        for src, _dst in extra_cols:
            row.append(r[header.index(src)] if src in header else "")
        try:
            row.append(f"{r[ev_v]} / {r[n_v]}" if ev_v is not None else "")
            row.append(f"{r[ev_c]} / {r[n_c]}" if ev_c is not None else "")
            row.append(
                f"{float(r[hr]):.2f} ({float(r[lo]):.2f}–{float(r[hi]):.2f})"
                if hr is not None and r[hr] else "—"
            )
            try:
                pv = float(r[p])
                row.append(f"{pv:.3f}" if pv >= 0.001 else "<0.001")
            except (ValueError, TypeError):
                row.append("—")
        except (IndexError, KeyError, TypeError):
            pass
        out.append(row)
    return out_header, out


# ---------------------------------------------------------------------------
# Build HPV_tables_figures.docx
# ---------------------------------------------------------------------------


def build_tables_figures_docx(suffix: str = "") -> None:
    doc = Document()
    # Body width ~6.5"
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    add_heading(doc, "Main tables and figures", level=1)

    # ---- Figure 1 ----
    add_image(doc, DATA / "Figure1_CohortSelection.png", width_in=6.2)
    add_caption(
        doc, "Figure 1",
        "CONSORT-style cohort selection flow diagram. From the single source "
        "population (N = 32,969) the diagram traces the derivation of Cohort A "
        "(whole-cohort safety analysis) through eligibility filtering, "
        "pseudo-index assignment for unvaccinated controls, and 1:1 propensity-"
        "score matching to a final n = 4,102 (2,051 vaccinated / 2,051 "
        "unvaccinated); and Cohort B (post-surgical efficacy analysis) through "
        "the cervical-surgery filter (n = 6,890), 1:up-to-5 initial matching "
        "(411 / 1,815), index-date eligibility, 1:up-to-4 fine matching to "
        "1,108 (241 / 867; legacy intermediate), and finally the ≥2-dose + "
        "3-month landmark filter with matched-set integrity yielding the v3 "
        "primary cohort n = 934 (204 vaccinated / 730 unvaccinated).",
    )
    doc.add_page_break()

    # ---- Figure 2 ----
    add_image(doc, DATA / "Figure2_CohortA_CIF_HR.png", width_in=6.5)
    add_caption(
        doc, "Figure 2",
        "Cohort A — Aalen–Johansen cumulative incidence functions and cluster-"
        "robust hazard ratios. Panels a–e: Any-of-5 composite (a), MCE "
        "composite (b), Diabetes (c), Hypertension (d), Angina/MI (e); "
        "vaccinated red, non-vaccinated blue with 95% pointwise CI bands. "
        "Panel f: forest plot of cluster-robust cause-specific hazard ratios; "
        "endpoints with fewer than five total events are marked insufficient.",
    )
    doc.add_page_break()

    # ---- Figure 3 ----
    add_image(doc, DATA / "Figure3_CohortB_CIF.png", width_in=6.5)
    add_caption(
        doc, "Figure 3",
        "Cohort B co-primary outcomes under the v3 primary (≥2-dose + 3-month "
        "landmark) — cumulative incidence curves with number-at-risk tables. "
        "(a) lesion recurrence (≥CIN2) in the v3 Cohort B (n = 934; 204 "
        "vaccinated / 730 controls); (b) cumulative clearance of pre-vaccine "
        "hr-HPV (two-consecutive-negative event) in the v3 pre-vaccine HPV+ "
        "subset (n = 235; 92 / 143). X-axis is years from landmark (index + 90 "
        "days). For clearance, HR > 1 favours vaccinated.",
    )
    doc.add_page_break()

    # ---- Figure 4 ----
    add_image(doc, DATA / "Figure4_CohortB_Subgroup.png", width_in=6.8)
    add_caption(
        doc, "Figure 4",
        "Cohort B subgroup analyses — combined table-with-forest plot for both "
        "co-primary outcomes. (a) lesion recurrence (HR < 1 favours "
        "vaccination); (b) hr-HPV clearance (HR > 1 favours vaccination). "
        "Strata: Overall, age at index (<40 / 40–49 / ≥50), vaccine type "
        "(Gardasil 9 / Cervarix / quadrivalent Gardasil). Likelihood-ratio P "
        "values for age × vaccination and vaccine-type × vaccination "
        "interaction terms are reported beneath each subgroup section.",
    )
    doc.add_page_break()

    # ---- Table 1 — split by cohort (A / B), N in caption ----
    splits = table1_split_rows()
    t1_header = ["Variable", "Vaccinated", "Non-vaccinated", "p value", "|SMD|"]
    t1_widths = [2.4, 1.1, 1.2, 0.6, 0.6]

    # Append a single "Pre-vaccine hr-HPV+" row to the Cohort B block in lieu
    # of a separate clearance-subset table. Counts/p/SMD precomputed against
    # CohortB_Clearance_Analytic.csv (n_v=110/241, n_c=182/867).
    cohortB_rows = splits["CohortB_post"] + [
        ["HPV history", "", "", "", ""],
        ["  Pre-vaccine hr-HPV+ (clearance analytic subset)",
         "110 (45.6%)", "182 (21.0%)", "<0.001", "0.542"],
    ]

    add_caption(
        doc, "Table 1A",
        "Baseline characteristics after 1:1 propensity-score matching — "
        "Cohort A (long-term safety analysis). N = 4,102 "
        "(2,051 vaccinated / 2,051 unvaccinated). Absolute standardised "
        "mean differences (|SMD|) < 0.10 indicate adequate balance.",
    )
    add_spacer(doc)
    add_table(doc, t1_header, splits["CohortA_post"], col_widths_in=t1_widths)
    doc.add_page_break()

    add_caption(
        doc, "Table 1B",
        "Baseline characteristics after 1:up-to-4 fine matching — Cohort B "
        "(post-surgical efficacy analysis; ≥1-dose, no-landmark legacy "
        "exposure definition retained for reference). N = 1,108 "
        "(241 vaccinated / 867 unvaccinated; mean realised ratio 3.60). "
        "Absolute standardised mean differences (|SMD|) < 0.10 indicate "
        "adequate balance. The final row (Pre-vaccine hr-HPV+) defines the "
        "clearance co-primary analytic subset (n = 292: 110 vaccinated / "
        "182 unvaccinated); the large |SMD| on this row reflects differential "
        "pre-vaccine molecular-pathology test frequency rather than a "
        "balance failure of the fine matching on the other covariates.",
    )
    add_spacer(doc)
    add_table(doc, t1_header, cohortB_rows, col_widths_in=t1_widths)
    doc.add_page_break()

    add_caption(
        doc, "Table 1C",
        "Baseline characteristics under the ≥2-dose + 3-month landmark PRIMARY "
        "exposure definition — Cohort B (post-surgical efficacy analysis). "
        "N = 934 (204 vaccinated / 730 fine-matched controls). Matched-set "
        "integrity preserved: vaccinated cases failing the ≥2-dose or "
        "landmark filter had their full fine-matched set removed. Additional "
        "variables — pre-surgery / post-surgery HPV test status (any-time and "
        "pre-vaccine windowed) and surgical-pathology severity (HSIL/CIN3 vs "
        "invasive cancer vs lower-grade) — are added to support cohort "
        "characterisation requested at revision.",
    )
    add_spacer(doc)
    add_table(doc, t1_header, splits["CohortB_post_v3"], col_widths_in=t1_widths)
    doc.add_page_break()

    # ---- Table 2 ----
    add_caption(
        doc, "Table 2",
        "Cohort A — events, person-time, and cluster-robust cause-specific and "
        "Fine–Gray subdistribution hazard ratios for each endpoint.",
    )
    add_spacer(doc)
    h, b = table2_rows()
    add_table(doc, h, b, col_widths_in=[1.4, 0.85, 0.85, 0.7, 0.5, 1.05, 0.4, 1.05, 0.4])
    doc.add_page_break()

    # ---- Table 3 ----
    add_caption(
        doc, "Table 3",
        "Cohort B — events, age-adjusted Cox hazard ratios (cluster-robust on "
        "the fine-matching identifier) and p-values for the two co-primary "
        "outcomes and the legacy post-index hr-HPV detection sensitivity row.",
    )
    add_spacer(doc)
    h, b = table3_rows()
    add_table(doc, h, b, col_widths_in=[2.2, 0.7, 1.1, 0.5, 1.1, 0.5, 1.1, 0.4])

    out = DOCS / f"HPV_tables_figures{suffix}.docx"
    doc.save(out)
    print(f"wrote {out.relative_to(ROOT)}")


# ---------------------------------------------------------------------------
# Build HPV_supplementary.docx
# ---------------------------------------------------------------------------


def build_supplementary_docx(suffix: str = "") -> None:
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    add_heading(doc, "Supplementary materials", level=1)

    # ---- Supplementary Figures S1–S6 ----
    add_heading(doc, "Supplementary Figures", level=2)

    sup_figs = [
        ("S1", "SupFigS1_loveplot_cohortA.png",
         "Love plot — covariate balance before and after 1:1 propensity-score "
         "matching (Cohort A)."),
        ("S2", "SupFigS2_loveplot_cohortB.png",
         "Love plot — covariate balance before and after fine variable-ratio "
         "(1:up-to-4) matching (Cohort B)."),
        ("S3", "SupFigS3_ps_density.png",
         "Propensity-score density distributions, before and after matching."),
        ("S4_a", "PH_check_A_0.png",
         "Schoenfeld residual plot — Cohort A, Any-of-5 composite."),
        ("S4_b", "PH_check_A_1.png",
         "Schoenfeld residual plot — Cohort A, Diabetes."),
        ("S5_a", "PH_check_B_has_recurrence.png",
         "Schoenfeld residual plot — Cohort B lesion recurrence."),
        ("S5_b", "PH_check_B_clearance.png",
         "Schoenfeld residual plot — Cohort B hr-HPV clearance "
         "(two-consecutive-negative event, pre-vaccine HPV-positive subset)."),
        ("S6", "SupFigS6_Sensitivity_Forest.png",
         "Pre-specified Cohort B sensitivity analyses — five-panel summary "
         "forest plot (Sens-A event-definition, Sens-B time-stratified, "
         "Sens-C dose threshold, Sens-D strict matching, Sens-E disease-free "
         "interval)."),
    ]
    for label, fname, caption in sup_figs:
        add_image(doc, DATA / fname, width_in=6.0)
        add_caption(doc, f"Supplementary Figure {label.replace('_', ' panel ')}", caption)
        add_spacer(doc)

    doc.add_page_break()

    # ---- Supplementary Tables ----
    add_heading(doc, "Supplementary Tables", level=2)

    # S1: pre-matching baseline — split by cohort
    h, body = read_csv(DATA / "Table1_BaselineCharacteristics_unified.csv")
    rows_a, rows_b = [], []
    for r in body:
        if len(r) < 6:
            continue
        if r[0] == "CohortA_pre":
            rows_a.append([translate_korean(r[1]), r[2], r[3], r[4], r[5]])
        elif r[0] == "CohortB_pre":
            rows_b.append([translate_korean(r[1]), r[2], r[3], r[4], r[5]])

    add_caption(
        doc, "Supplementary Table S1A",
        "Pre-matching baseline characteristics — Cohort A (full source "
        "population, before 1:1 propensity-score matching).",
    )
    add_spacer(doc)
    add_table(doc,
              ["Variable", "Vaccinated", "Non-vaccinated", "p", "|SMD|"],
              rows_a, col_widths_in=[2.5, 1.1, 1.2, 0.6, 0.6])
    doc.add_page_break()

    add_caption(
        doc, "Supplementary Table S1B",
        "Pre-matching baseline characteristics — Cohort B (post-surgical "
        "population, before 1:up-to-4 fine matching).",
    )
    add_spacer(doc)
    add_table(doc,
              ["Variable", "Vaccinated", "Non-vaccinated", "p", "|SMD|"],
              rows_b, col_widths_in=[2.5, 1.1, 1.2, 0.6, 0.6])
    doc.add_page_break()

    # S2: PS coefficients
    add_caption(
        doc, "Supplementary Table S2",
        "Propensity-score model coefficients (logistic regression) for Cohort A.",
    )
    add_spacer(doc)
    h, b = read_csv(DATA / "SupTableS2_ps_coefficients.csv")
    # Round means to 2 dp, coefficients to 3 dp, ORs to 2 dp.
    rows = []
    for r in b:
        if len(r) < 6:
            continue
        rows.append([
            r[0],
            fmt_num(r[1], 2),
            fmt_num(r[2], 2),
            fmt_num(r[3], 3),
            fmt_num(r[4], 3),
            fmt_num(r[5], 2),
        ])
    add_table(doc, h, rows)
    doc.add_page_break()

    # S3: Age × FU forest data
    add_caption(
        doc, "Supplementary Table S3",
        "Age-stratified hazard ratios for Cohort B lesion recurrence across "
        "follow-up windows.",
    )
    add_spacer(doc)
    h_src, b_src = read_csv(DATA / "CohortB_age_fu_forest.csv")
    s3_header = ["Stratum", "Follow-up window", "n total",
                 "Vac events / N", "Ctl events / N", "HR (95% CI)", "p"]
    s3_rows = []
    for r in b_src:
        try:
            (n, n_vac, n_ctl, ev_vac, ev_ctl, hr, lo, hi, p,
             stratum, fu_label, _fu_yr) = r[:12]
            s3_rows.append([
                stratum, fu_label, n,
                f"{ev_vac} / {n_vac}", f"{ev_ctl} / {n_ctl}",
                fmt_hr_triplet(hr, lo, hi), fmt_p(p),
            ])
        except (ValueError, IndexError):
            continue
    add_table(doc, s3_header, s3_rows,
              col_widths_in=[1.1, 1.3, 0.6, 1.0, 1.0, 1.2, 0.5])
    doc.add_page_break()

    # S4: Number at risk — generate inline from Figure 3 source if available
    add_caption(
        doc, "Supplementary Table S4",
        "Number-at-risk tables for Cohort B Kaplan–Meier and Aalen–Johansen "
        "cumulative-incidence curves (years 0, 2, 4, 6, 8, 10).",
    )
    add_spacer(doc)
    # static rows lifted from Figure 3 panel a + b
    nar_header = ["Outcome", "Group", "0", "2", "4", "6", "8", "10"]
    nar_body = [
        ["Lesion recurrence",  "Vaccinated",     "241", "209", "149", "89",  "60",  "43"],
        ["Lesion recurrence",  "Non-vaccinated", "867", "730", "547", "338", "233", "152"],
        ["hr-HPV clearance",   "Vaccinated",     "110", "65",  "44",  "22",  "13",  "8"],
        ["hr-HPV clearance",   "Non-vaccinated", "182", "117", "87",  "50",  "28",  "14"],
    ]
    add_table(doc, nar_header, nar_body)
    doc.add_page_break()

    # S5: Vaccine-type interaction
    add_caption(
        doc, "Supplementary Table S5",
        "Per-vaccine-type detailed results — single-model interaction-derived "
        "hazard ratios with the likelihood-ratio test for vaccine-type "
        "heterogeneity. Rows include both co-primary outcomes plus the legacy "
        "post-index hr-HPV detection sensitivity.",
    )
    add_spacer(doc)
    h, b = read_csv(DATA / "CohortB_vaccine_interaction.csv")
    # CSV column order: outcome, LRT_chi2, df, LRT_p, then HR/CI/CI triplets.
    fmt = []
    for r in b:
        try:
            outcome, lrt_chi2, df_, lrt_p = r[0], r[1], r[2], r[3]
        except IndexError:
            continue
        new = [outcome, fmt_num(lrt_chi2, 2), df_, fmt_p(lrt_p)]
        for i in (4, 7, 10):
            try:
                new.append(fmt_hr_triplet(r[i], r[i+1], r[i+2]))
            except IndexError:
                new.append("—")
        fmt.append(new)
    add_table(doc,
              ["Outcome", "LRT χ²", "df", "LRT p",
               "Gardasil 9 HR (95% CI)", "Cervarix HR (95% CI)",
               "Quadrivalent Gardasil HR (95% CI)"],
              fmt)
    doc.add_page_break()

    # S6: Cluster-robust HR with PH p — Cohort A + Cohort B with PY/IR/Schoenfeld
    add_caption(
        doc, "Supplementary Table S6",
        "Cluster-robust hazard ratios with person-years, incidence rates per "
        "1,000 person-years, and Schoenfeld residual p-values for both "
        "cohorts. Cohort B rows include both co-primary outcomes (lesion "
        "recurrence, hr-HPV clearance) and the legacy post-index hr-HPV "
        "detection sensitivity row.",
    )
    add_spacer(doc)
    h, b = read_csv(DATA / "Table2_CohortA_HazardRatios.csv")
    rows_a = []
    for r in b:
        try:
            outcome = r[0]; e_v=r[1]; n_v=r[2]; e_c=r[3]; n_c=r[4]
            ir_v = f"{float(r[8]):.2f}"; ir_c = f"{float(r[9]):.2f}"
            hr = f"{float(r[10]):.2f} ({float(r[11]):.2f}–{float(r[12]):.2f})" if r[10] else "—"
            pv = f"{float(r[13]):.3f}" if r[13] else "—"
            ph = f"{float(r[14]):.3f}" if r[14] else "—"
            rows_a.append(["Cohort A", outcome, f"{e_v}/{n_v}", f"{e_c}/{n_c}",
                           ir_v, ir_c, hr, pv, ph])
        except (ValueError, IndexError):
            continue

    # Cohort B from Table3 — parse the embedded "PY (rate /1000 PY)" strings
    h_b, b_b = read_csv(DATA / "Table3_CohortB_HR.csv")
    # Hard-coded Schoenfeld p values for Cohort B primary models (Methods §Stat).
    SCH_B = {
        "Lesion recurrence (≥CIN2 / HSIL+ or invasive carcinoma)": "0.820",
        "hr-HPV clearance / regression (pre-vaccine HPV+ baseline)": "0.028",
    }

    def split_py(s: str) -> tuple[str, str]:
        """Parse '1405 (9.2)' → ('1405', '9.2')."""
        m = re.match(r"\s*(\S+)\s*\((\S+)\)\s*", s)
        if m:
            return m.group(1), m.group(2)
        return s, "—"

    rows_b = []
    for r in b_b:
        if len(r) < 8:
            continue
        outcome = r[0]
        py_v, ir_v = split_py(r[3])
        py_c, ir_c = split_py(r[5])
        ph = SCH_B.get(outcome, "—")
        rows_b.append(["Cohort B", outcome, r[2], r[4], ir_v, ir_c, r[6], r[7], ph])

    add_table(doc,
              ["Cohort", "Outcome", "Vac events / N", "Ctl events / N",
               "IR Vac", "IR Ctl", "HR (95% CI)", "p", "Schoenfeld p"],
              rows_a + rows_b,
              col_widths_in=[0.6, 2.1, 0.8, 0.8, 0.5, 0.5, 1.1, 0.4, 0.5])
    doc.add_page_break()

    # S7: pseudo-index sensitivity
    add_caption(
        doc, "Supplementary Table S7",
        "Pseudo-index assignment sensitivity analysis for Cohort A (Any-of-five) — "
        "random sampling, calendar-year-matched, and risk-set sampling strategies.",
    )
    add_spacer(doc)
    h_src, b_src = read_csv(DATA / "CohortA_pseudoindex_sensitivity.csv")
    s7_header = ["Strategy", "Vac events / N", "Ctl events / N", "HR (95% CI)", "p"]
    s7_rows = []
    for r in b_src:
        try:
            strategy, n_vac, n_ctl, ev_vac, ev_ctl, hr, lo, hi, p = r[:9]
            s7_rows.append([
                strategy,
                f"{ev_vac} / {n_vac}", f"{ev_ctl} / {n_ctl}",
                fmt_hr_triplet(hr, lo, hi), fmt_p(p),
            ])
        except (ValueError, IndexError):
            continue
    add_table(doc, s7_header, s7_rows)
    doc.add_page_break()

    # S8A: Dose threshold (Sens-C, primary)
    add_caption(
        doc, "Supplementary Table S8A",
        "(Sens-C, primary) Dose-threshold sensitivity for both cohorts — "
        "re-fitted hazard ratios under ≥1, ≥2, and ≥3 (complete schedule) "
        "dose definitions; matched-set integrity preserved by dropping the "
        "entire matched set when the vaccinated case fails the threshold.",
    )
    add_spacer(doc)
    h, b = read_csv(DATA / "Sensitivity_DoseThreshold_HR.csv")
    rows = []
    for r in b:
        try:
            cohort, outcome, definition, threshold, n_v, n_c, ev_v, ev_c, hr, lo, hi, p = r[:12]
            hr_txt = (f"{float(hr):.2f} ({float(lo):.2f}–{float(hi):.2f})" if hr else "—")
            pv = (f"{float(p):.3f}" if p else "—")
            rows.append([cohort, outcome, definition,
                         f"{ev_v}/{n_v}", f"{ev_c}/{n_c}", hr_txt, pv])
        except (ValueError, IndexError):
            continue
    add_table(doc,
              ["Cohort", "Outcome", "Dose threshold",
               "Vac events / N", "Ctl events / N", "HR (95% CI)", "p"],
              rows,
              col_widths_in=[0.5, 1.6, 1.5, 0.8, 0.8, 1.1, 0.4])
    doc.add_page_break()

    # S8B: Dose threshold with landmark (immortal-time-corrected)
    add_caption(
        doc, "Supplementary Table S8B",
        "(Sens-C, landmark) Immortal-time-bias-corrected dose-threshold "
        "sensitivity for Cohort B. Landmarks reflect the standard HPV-vaccine "
        "0–2–6 month schedule plus grace: ≥1 dose at 30 days, ≥2 doses at "
        "90 days, ≥3 doses at 240 days. Patients (both arms) must be alive "
        "and event-free at the landmark to enter the analysis; for vaccinated "
        "cases, the k-th dose must additionally have been received by the "
        "landmark. Time is left-truncated at the landmark.",
    )
    add_spacer(doc)
    h, b = read_csv(DATA / "Sensitivity_DoseThreshold_Landmark.csv")
    rows = []
    for r in b:
        try:
            outcome, definition, threshold, lm, n_v, n_c, ev_v, ev_c, hr, lo, hi, p = r[:12]
            hr_txt = (f"{float(hr):.2f} ({float(lo):.2f}–{float(hi):.2f})" if hr else "—")
            pv = fmt_p(p)
            rows.append([outcome, definition,
                         f"{ev_v}/{n_v}", f"{ev_c}/{n_c}", hr_txt, pv])
        except (ValueError, IndexError):
            continue
    add_table(doc,
              ["Outcome", "Definition (landmark)",
               "Vac events / N", "Ctl events / N", "HR (95% CI)", "p"],
              rows,
              col_widths_in=[1.5, 2.0, 0.9, 0.9, 1.1, 0.4])
    doc.add_page_break()

    # S9: Strict matching (Sens-D)
    add_caption(
        doc, "Supplementary Table S9",
        "(Sens-D) Strict 1:4 fine-matching sensitivity for Cohort B — "
        "variable-ratio (primary) versus strict (sensitivity) specifications.",
    )
    add_spacer(doc)
    h, b = read_csv(DATA / "Sensitivity_StrictMatching.csv")
    rows = []
    for r in b:
        try:
            outcome, design, n_v, n_c, ev_v, ev_c, hr, lo, hi, p, _det = r[:11]
            hr_txt = f"{float(hr):.2f} ({float(lo):.2f}–{float(hi):.2f})"
            pv = f"{float(p):.3f}"
            rows.append([outcome, design,
                         f"{ev_v}/{n_v}", f"{ev_c}/{n_c}", hr_txt, pv])
        except (ValueError, IndexError):
            continue
    add_table(doc,
              ["Outcome", "Design", "Vac events / N", "Ctl events / N",
               "HR (95% CI)", "p"],
              rows)
    doc.add_page_break()

    # S10: Time-stratified clearance (Sens-B)
    add_caption(
        doc, "Supplementary Table S10",
        "(Sens-B) Time-stratified hr-HPV clearance hazard ratios decomposed into "
        "0–6, 6–12, 12–24, and ≥24-month windows post-index. Left-truncation at "
        "each window's lower bound.",
    )
    add_spacer(doc)
    h, b = read_csv(DATA / "Sensitivity_HPV_Clearance_TimeStratified_v3.csv")
    rows = []
    for r in b:
        try:
            # v3 schema: period, n_v, n_c, ev_v, ev_c, HR, CIlo, CIhi, p
            period, n_v, n_c, ev_v, ev_c, hr, lo, hi, p = r[:9]
            hr_txt = f"{float(hr):.2f} ({float(lo):.2f}–{float(hi):.2f})"
            pv = f"{float(p):.3f}"
            rows.append([period, f"{ev_v}/{n_v}", f"{ev_c}/{n_c}", hr_txt, pv])
        except (ValueError, IndexError):
            continue
    add_table(doc,
              ["Window", "Vac events / N", "Ctl events / N", "HR (95% CI)", "p"],
              rows)
    doc.add_page_break()

    # S11: Single-negative clearance (Sens-A)
    add_caption(
        doc, "Supplementary Table S11",
        "(Sens-A) Single-negative HPV clearance sensitivity — alternative event "
        "definition using the FIRST single post-index hr-HPV-negative record, "
        "contrasted with the two-consecutive-negative primary.",
    )
    add_spacer(doc)
    h, b = read_csv(DATA / "Sensitivity_HPV_Clearance_SingleNegative_v3.csv")
    rows = []
    for r in b:
        try:
            definition, n_v, n_c, ev_v, ev_c, hr, lo, hi, p = r[:9]
            hr_txt = f"{float(hr):.2f} ({float(lo):.2f}–{float(hi):.2f})"
            pv = f"{float(p):.3f}"
            rows.append([definition, f"{ev_v}/{n_v}", f"{ev_c}/{n_c}", hr_txt, pv])
        except (ValueError, IndexError):
            continue
    add_table(doc,
              ["Clearance definition", "Vac events / N", "Ctl events / N",
               "HR (95% CI)", "p"],
              rows)
    doc.add_page_break()

    # S12: Disease-free interval (Sens-E)
    add_caption(
        doc, "Supplementary Table S12",
        "(Sens-E) Disease-free-interval sensitivity for lesion recurrence — "
        "minimum 3-, 6-, and 12-month buffer from the index date before counting "
        "a recurrence event.",
    )
    add_spacer(doc)
    h, b = read_csv(DATA / "Sensitivity_Recurrence_DFInterval_v3.csv")
    rows = []
    for r in b:
        try:
            definition, n_v, n_c, ev_v, ev_c, hr, lo, hi, p = r[:9]
            hr_txt = f"{float(hr):.2f} ({float(lo):.2f}–{float(hi):.2f})"
            pv = f"{float(p):.3f}"
            rows.append([definition, f"{ev_v}/{n_v}", f"{ev_c}/{n_c}", hr_txt, pv])
        except (ValueError, IndexError):
            continue
    add_table(doc,
              ["Disease-free interval", "Vac events / N", "Ctl events / N",
               "HR (95% CI)", "p"],
              rows)

    out = DOCS / f"HPV_supplementary{suffix}.docx"
    doc.save(out)
    print(f"wrote {out.relative_to(ROOT)}")


def main() -> None:
    # Emit both the canonical filenames and the _v2 set so legacy paths
    # keep working while the new cohort-split Table 1 ships under v2.
    for suffix in ("", "_v2"):
        build_manuscript_docx(suffix)
        build_tables_figures_docx(suffix)
        build_supplementary_docx(suffix)


if __name__ == "__main__":
    main()
