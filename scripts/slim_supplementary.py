"""
Slim down the supplementary docx by dropping 7 tables and renumbering the rest.

DROP:
  - S3  (Sensitivity analyses for Cohort B — restricted FU + adjusted/unadjusted)
  - S11 (Prescription-code vs drug-name cross-validation)
  - S12 (Vaccine-type recipient patterns)
  - S13 (Landmark sensitivity for post-index hr-HPV detection — superseded by S17)
  - S14 (Vaccine-type interaction by calendar period — Discussion narrative only)
  - S15 (Novel-type acquisition sensitivity)
  - S16 (HPV clearance sensitivity — any-clearance + HPV-16/18 type-specific)

RENUMBER (final):
  S1 -> S1
  S2 -> S2
  S4 -> S3   (age × FU subgroup)
  S5 -> S4   (number-at-risk)
  S6 -> S5   (vaccine-type detailed)
  S7 -> S6   (cluster-robust HR + Schoenfeld)
  S8 -> S7   (pseudo-index sensitivity, Cohort A)
  S9 -> S8   (dose threshold — Sens-C)
  S10 -> S9  (strict matching — Sens-D)
  S17 -> S10 (time-stratified clearance — Sens-B)
  S18 -> S11 (single-negative clearance — Sens-A)
  S19 -> S12 (disease-free interval — Sens-E)

Process:
  1) Locate each (header paragraph, table, footnote paragraph) triplet by header text.
  2) Drop the OOXML elements between the header and the next non-table-related element.
  3) Pass over remaining text and apply the renumber mapping.
"""
from docx import Document
from docx.oxml.ns import qn
import re

SRC = 'docs/HPV_supplementary.docx'

# Old -> new label mapping
RENUM = {
    'Table S4': 'Table S3',
    'Table S5': 'Table S4',
    'Table S6': 'Table S5',
    'Table S7': 'Table S6',
    'Table S8': 'Table S7',
    'Table S9': 'Table S8',
    'Table S10': 'Table S9',
    'Table S17': 'Table S10',
    'Table S18': 'Table S11',
    'Table S19': 'Table S12',
}

# Header signatures of tables to drop (prefix match on stripped paragraph text)
DROP_HEADERS = [
    'Table S3. Sensitivity analyses for Cohort B',
    'Table S11. Prescription-code vs',
    'Table S12. Vaccine-type recipient patterns',
    'Table S13. Landmark sensitivity analysis for the post-index hr-HPV detection',
    'Table S14. Vaccine-type interaction analysis stratified by calendar period',
    'Table S15. Novel-type acquisition sensitivity',
    'Table S16. HPV clearance sensitivity analysis among women',
]

# TOC lines (in the first section P5-P23) that reference the dropped tables.
TOC_DROP_PREFIXES = [
    'Table S3. Sensitivity analyses for Cohort B',
    'Table S11. Prescription-code',
    'Table S12. Vaccine-type recipient patterns',
    'Table S13. Landmark sensitivity analysis',
    'Table S14. Vaccine-type interaction by calendar period',
    'Table S15. Novel-type acquisition sensitivity analysis',
    'Table S16. HPV clearance sensitivity analysis',
]


def starts_with_any(text, prefixes):
    t = text.strip()
    return any(t.startswith(p) for p in prefixes)


def drop_section_after_header(body, header_p):
    """Remove the header paragraph and every following element until we hit the
    next paragraph that begins with 'Table S' or 'Figure S' or the doc end.
    Within the block we may encounter:
      - sub-heading paragraphs (e.g., 'Lesion recurrence', 'A. Pairwise...')
      - one or more tables
      - empty paragraphs
      - a footnote paragraph
    All of these belong to the dropped table.
    """
    hdr_el = header_p._element
    siblings = list(body.iterchildren())
    try:
        start = siblings.index(hdr_el)
    except ValueError:
        return 0
    to_remove = [hdr_el]
    i = start + 1
    while i < len(siblings):
        el = siblings[i]
        if el.tag == qn('w:p'):
            t = el.xpath('string(.)').strip()
            if t.startswith('Table S') or t.startswith('Figure S'):
                break
        to_remove.append(el)
        i += 1
    for el in to_remove:
        body.remove(el)
    return len(to_remove)


def remove_toc_lines(doc, prefixes):
    """Remove TOC paragraphs ONLY in the TOC region — defined as everything
    BEFORE the first paragraph that introduces a body table (i.e. the first
    paragraph whose text starts with a Table S1 body description like
    'Table S1. Pre-matching baseline characteristics of the source population').

    The TOC is short, single-line entries; the body has a longer descriptive
    title that runs to ~100+ chars. Easier: find the index of the FIRST 'Figure
    S5.' TOC entry (which is always present and is the last TOC line in this
    doc), then anything after it is body content.
    """
    removed = 0
    # find the body-region boundary: first paragraph whose text starts with
    # 'Cohort A (pre-matching)' — that paragraph is inside the S1 body table.
    boundary_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'Cohort A (pre-matching)':
            boundary_idx = i
            break
    if boundary_idx is None:
        # fallback: after Figure S5 TOC line
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith('Figure S5. Schoenfeld'):
                boundary_idx = i + 1
                break

    to_remove = []
    for i, p in enumerate(doc.paragraphs[:boundary_idx]):
        t = p.text.strip()
        if not t:
            continue
        if starts_with_any(t, prefixes) and len(t) < 200:
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)
        removed += 1
    return removed


def renumber_labels(doc, mapping):
    """Two-phase substitution to avoid cascading:
      Phase 1: old label -> sentinel token (unique, won't match any other rule)
      Phase 2: sentinel token -> new label
    This way, after Phase 1 no original-form labels remain, so Phase 2 only
    sees sentinels and can't accidentally collide.
    """
    n = 0
    phase1 = []  # (old_label, sentinel)
    phase2 = []  # (sentinel, new_label)
    for old, new in mapping.items():
        sentinel = f'§§TBL{old.split("S")[-1]}§§'  # e.g. §§TBL17§§
        phase1.append((old, sentinel))
        phase2.append((sentinel, new))

    # apply phase1 over longer old labels first to prevent prefix overlap
    phase1.sort(key=lambda kv: -len(kv[0]))

    def apply_subs(p, subs):
        nonlocal n
        for run in p.runs:
            for old, new in subs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    n += 1
        # cross-run fallback only if substring still present at paragraph level
        if any(old in p.text for old, _ in subs):
            full = p.text
            new_full = full
            for old, new in subs:
                new_full = new_full.replace(old, new)
            if new_full != full:
                for r in list(p.runs):
                    r.text = ''
                if p.runs:
                    p.runs[0].text = new_full
                else:
                    p.add_run(new_full)
                n += 1

    # Phase 1
    for p in doc.paragraphs:
        apply_subs(p, phase1)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    apply_subs(p, phase1)
    # Phase 2
    for p in doc.paragraphs:
        apply_subs(p, phase2)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    apply_subs(p, phase2)

    return n


# ============================ MAIN ============================
doc = Document(SRC)
body = doc.element.body

# 1) FIRST remove the TOC entries (single-line index at top of doc) so that
#    the next step's "find header by prefix" matches the body table heading,
#    not the TOC line.
n_toc = remove_toc_lines(doc, TOC_DROP_PREFIXES)
print(f'TOC lines removed: {n_toc}\n')

# 2) Drop body table blocks (header + subheadings + tables + footnote) by
#    locating the body table heading and removing everything up to the next
#    'Table S*' or 'Figure S*' heading.
total_removed = 0
for prefix in DROP_HEADERS:
    header_p = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(prefix):
            header_p = p
            break
    if header_p is None:
        print(f'  [warn] could not find header for: {prefix}')
        continue
    n = drop_section_after_header(body, header_p)
    print(f'  dropped {prefix[:60]}: {n} elements')
    total_removed += n

# 3) Renumber remaining labels (two-phase sentinel substitution; no cascade)
n_renum = renumber_labels(doc, RENUM)
print(f'\nLabel renumbers: {n_renum}')

doc.save(SRC)
print(f'\nSaved: {SRC}')
print(f'Total OOXML elements removed: {total_removed}')
