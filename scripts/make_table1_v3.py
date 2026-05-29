"""
Build Table 1 v3 — Cohort B baseline characteristics under ≥2-dose + 3-mo landmark
primary cohort (n=912; 203 vaccinated / 709 fine-matched controls). Adds:
  - Pre-surgery HPV (HR+ / non-HR+ / Negative / No test)
  - Post-surgery HPV, first test after surgery (two versions: any-time + pre-vaccine windowed)
  - Surgical pathology severity (Invasive cancer / HSIL/CIN3 / CIN2/moderate / lower / benign)

Reuses helpers from baseline_table1_unified.py.
Outputs:
  Data/Table1B_v3_CohortB.csv
  Data/Table1B_v3_CohortB.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from scipy.stats import fisher_exact, ttest_ind

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / 'scripts'))
from extract_pathology_outcomes import detect_high_risk_hpv  # noqa: E402


# ------------------------------ helpers ------------------------------
def smd_cont(a, b):
    a, b = pd.Series(a).dropna(), pd.Series(b).dropna()
    if len(a) < 2 or len(b) < 2:
        return np.nan
    pooled = np.sqrt((a.var(ddof=1) + b.var(ddof=1)) / 2)
    return (a.mean() - b.mean()) / pooled if pooled > 0 else np.nan


def smd_bin(p1, n1, p2, n2):
    if n1 == 0 or n2 == 0:
        return np.nan
    pp1, pp2 = p1 / n1, p2 / n2
    pooled = np.sqrt((pp1 * (1 - pp1) + pp2 * (1 - pp2)) / 2)
    return (pp1 - pp2) / pooled if pooled > 0 else np.nan


def p_cont(a, b):
    a, b = pd.Series(a).dropna(), pd.Series(b).dropna()
    if len(a) < 2 or len(b) < 2:
        return np.nan
    return ttest_ind(a, b, equal_var=False).pvalue


def p_bin(a_yes, a_n, b_yes, b_n):
    if a_n == 0 or b_n == 0:
        return np.nan
    return fisher_exact([[a_yes, a_n - a_yes], [b_yes, b_n - b_yes]])[1]


def fmt_p(p):
    if pd.isna(p):
        return '-'
    return '<0.001' if p < 0.001 else f'{p:.3f}'


def fmt_smd(s):
    return '-' if pd.isna(s) else f'{abs(s):.3f}'


def fmt_cont(s):
    s = pd.Series(s).dropna()
    return '-' if len(s) == 0 else f'{s.mean():.2f} ± {s.std():.2f}'


def fmt_pct(n_yes, n_total):
    return f'{int(n_yes)} ({100 * n_yes / n_total:.1f}%)' if n_total > 0 else '-'


# ------------------------------ HPV / severity classifiers ------------------------------
SEVERITY_ORDER = [
    'Invasive cancer',
    'HSIL/CIN3',
    'CIN2 / moderate dysplasia',
    'CIN1 / LSIL / mild dysplasia',
    'Benign / other',
]

INVASIVE_RE = re.compile(
    r'(invasive\s+carcinoma|squamous\s+cell\s+carcinoma|\bSCC\b(?!\s*type)|adenocarcinoma|'
    r'cervical\s+cancer|침윤(성|암)|편평세포암|선암)',
    re.I,
)
CIS_RE = re.compile(r'(carcinoma\s*in\s*situ|\bCIS\b)', re.I)
HSIL_RE = re.compile(
    r'(HSIL|H[-\s]?SIL|CIN\s*[-]?\s*3|CIN\s*III|'
    r'high[-\s]?grade(\s+squamous)?\s+(intraepithelial\s+)?lesion|severe\s+dysplasia)',
    re.I,
)
CIN2_RE = re.compile(r'(CIN\s*[-]?\s*2|CIN\s*II|moderate\s+dysplasia)', re.I)
LSIL_RE = re.compile(
    r'(LSIL|L[-\s]?SIL|CIN\s*[-]?\s*1|CIN\s*I\b|low[-\s]?grade|mild\s+dysplasia)',
    re.I,
)


def classify_severity(text: str) -> str:
    if INVASIVE_RE.search(text):
        return 'Invasive cancer'
    if CIS_RE.search(text) or HSIL_RE.search(text):
        return 'HSIL/CIN3'
    if CIN2_RE.search(text):
        return 'CIN2 / moderate dysplasia'
    if LSIL_RE.search(text):
        return 'CIN1 / LSIL / mild dysplasia'
    return 'Benign / other'


def hpv_category(text: str, hr_pos: bool):
    if hr_pos:
        return 'HR-HPV+'
    t = str(text).lower()
    if re.search(r'\bnegative\b|음성|not detected|none detected', t):
        return 'HPV negative'
    if re.search(r'\bpositive\b|양성|detected', t):
        return 'non-HR HPV+'
    return None


# ------------------------------ build cohort ------------------------------
print('Loading...')
df = pd.read_csv(ROOT / 'Data' / 'primary_cohort_v3.csv', encoding='utf-8-sig')
df['index_date'] = pd.to_datetime(df['index_date'])
df['최종추적일자'] = pd.to_datetime(df['최종추적일자'])

mc = pd.read_csv(ROOT / 'Data' / 'final_matched_cohort.csv', encoding='utf-8-sig')
mc['첫수술일자'] = pd.to_datetime(mc['첫수술일자'])
surg_map = mc.set_index('연구번호')['첫수술일자']
df['첫수술일자'] = df['연구번호'].map(surg_map)

cohort = pd.read_csv(
    ROOT / 'Data' / '한국 HPV 코호트 자료를 이용한 자_코호트.csv',
    encoding='cp949',
    low_memory=False,
)
cohort['region_top'] = cohort['주소'].astype(str).str.split().str[0]
cohort['is_seoul'] = (cohort['region_top'] == '서울').astype(int)
df = df.merge(
    cohort[['연구번호', 'is_seoul']].drop_duplicates('연구번호'), on='연구번호', how='left'
)

# Smoking, BP, comorbidities — pull from clinical info (closest record before index)
clin = pd.read_csv(
    ROOT / 'Data' / '한국 HPV 코호트 자료를 이용한 자_기초임상정보.csv',
    encoding='cp949',
    low_memory=False,
)
clin['기록일자'] = pd.to_datetime(
    clin['기록일자'].astype(str).str.strip(), format='%Y%m%d', errors='coerce'
)
clin['SBP'] = pd.to_numeric(clin.get('수축기혈압'), errors='coerce')
clin['DBP'] = pd.to_numeric(clin.get('이완기혈압'), errors='coerce')
SMOKE_MAP = {
    '비흡연': 'Never',
    '과거흡연': 'Former',
    '현재흡연': 'Current',
    '확인불능': 'Unknown',
}
clin['smoke_norm'] = clin.get('흡연여부').map(SMOKE_MAP).fillna('Unknown')


def closest_record(pid, idx_dt, col):
    sub = clin[(clin['연구번호'] == pid) & (clin['기록일자'].notna())].copy()
    sub = sub[sub[col].notna()]
    if len(sub) == 0:
        return np.nan
    sub['delta'] = (sub['기록일자'] - idx_dt).abs()
    sub = sub[sub['delta'] <= pd.Timedelta(days=365)]
    if len(sub) == 0:
        return np.nan
    return sub.sort_values('delta').iloc[0][col]


def latest_smoke(pid, idx_dt):
    sub = clin[(clin['연구번호'] == pid) & (clin['기록일자'] <= idx_dt)]
    if len(sub) == 0:
        return 'Unknown'
    return sub.sort_values('기록일자').iloc[-1]['smoke_norm']


print('Computing SBP/DBP/smoking…')
df['SBP'] = df.apply(lambda r: closest_record(r['연구번호'], r['index_date'], 'SBP'), axis=1)
df['DBP'] = df.apply(lambda r: closest_record(r['연구번호'], r['index_date'], 'DBP'), axis=1)
df['smoke'] = df.apply(lambda r: latest_smoke(r['연구번호'], r['index_date']), axis=1)

# Comorbidities — load pre-classified diag file
diag = pd.read_excel(
    ROOT / 'Data' / '한국 HPV 코호트 자료를 이용한 자_진단정보_기저질환추가_unlocked.xlsx'
)
diag['진단일자'] = pd.to_datetime(diag['진단일자'].astype(str), format='%Y%m%d', errors='coerce')
CLASS_LABELS = {'1': 'Angina/MI', '2': 'Hypertension', '3': 'Diabetes', '4': 'Stroke', '5': 'PE'}
COMOR_COL = [c for c in diag.columns if c.startswith('기저질환분류')][0]


def has_comor(pid, idx_dt, cls):
    sub = diag[(diag['연구번호'] == pid) & (diag[COMOR_COL].astype(str) == cls)]
    return int(((sub['진단일자'] <= idx_dt)).any())


for cls, lbl in CLASS_LABELS.items():
    df[f'comor_{lbl}'] = df.apply(
        lambda r: has_comor(r['연구번호'], r['index_date'], cls), axis=1
    )
df['comor_any'] = df[[f'comor_{lbl}' for lbl in CLASS_LABELS.values()]].max(axis=1)

# ------------------------------ HPV pre/post-surgery + severity ------------------------------
print('Loading pathology…')
path = pd.read_csv(
    ROOT / 'Data' / '한국 HPV 코호트 자료를 이용한 자_병리검사 (복구됨).CSV',
    encoding='cp949',
    low_memory=False,
)
path['실시일자'] = pd.to_datetime(path['실시일자'], format='%Y%m%d', errors='coerce')

mol = path[path['병리검사구분'].isin(['분자병리', 'HPV'])].dropna(subset=['실시일자', '판독결과']).copy()
res = mol['판독결과'].apply(detect_high_risk_hpv)
mol['hpv_pos'] = res.apply(lambda d: d['is_high_risk_hpv_positive'])
mol['hpv_cat'] = mol.apply(lambda r: hpv_category(r['판독결과'], r['hpv_pos']), axis=1)
mol = mol.dropna(subset=['hpv_cat'])
mol_by_pid = {pid: g.sort_values('실시일자') for pid, g in mol.groupby('연구번호')}

PRIORITY = {'HR-HPV+': 3, 'non-HR HPV+': 2, 'HPV negative': 1}


def worst_hpv(cats):
    return max(cats, key=lambda c: PRIORITY[c])


def pre_surg_hpv(pid, sd):
    sub = mol_by_pid.get(pid)
    if sub is None:
        return 'No test'
    pre = sub[sub['실시일자'] < sd]
    return worst_hpv(pre['hpv_cat'].tolist()) if len(pre) > 0 else 'No test'


def post_surg_first(pid, sd, idx_dt=None):
    sub = mol_by_pid.get(pid)
    if sub is None:
        return 'No test'
    post = sub[sub['실시일자'] > sd]
    if idx_dt is not None:
        post = post[post['실시일자'] < idx_dt]
    return post.iloc[0]['hpv_cat'] if len(post) > 0 else 'No test'


tissue = path[
    path['병리검사구분'].isin(['조직병리', '조직', '위탁조직병리', '조직-동결절편'])
].dropna(subset=['실시일자', '판독결과'])
tissue_by_pid = {pid: g.sort_values('실시일자') for pid, g in tissue.groupby('연구번호')}


def worst_severity(pid, sd):
    sub = tissue_by_pid.get(pid)
    if sub is None or pd.isna(sd):
        return 'Benign / other'
    sub = sub[sub['실시일자'] <= sd + pd.Timedelta(days=30)]
    if len(sub) == 0:
        return 'Benign / other'
    sevs = sub['판독결과'].apply(classify_severity)
    for s in SEVERITY_ORDER:
        if (sevs == s).any():
            return s
    return 'Benign / other'


print('Computing HPV/severity…')
df['pre_surg_hpv'] = df.apply(lambda r: pre_surg_hpv(r['연구번호'], r['첫수술일자']), axis=1)
df['post_surg_anytime'] = df.apply(
    lambda r: post_surg_first(r['연구번호'], r['첫수술일자'], None), axis=1
)
df['post_surg_prevac'] = df.apply(
    lambda r: post_surg_first(r['연구번호'], r['첫수술일자'], r['index_date']), axis=1
)
df['surg_severity'] = df.apply(
    lambda r: worst_severity(r['연구번호'], r['첫수술일자']), axis=1
)

# ------------------------------ assemble Table 1 rows ------------------------------
v = df[df['접종여부'] == True]
nv = df[df['접종여부'] == False]
n_v, n_nv = len(v), len(nv)


def row_cont(name, col):
    a, b = v[col], nv[col]
    return {
        'Variable': name,
        f'Vaccinated (n={n_v})': fmt_cont(a),
        f'Non-vaccinated (n={n_nv})': fmt_cont(b),
        'p': fmt_p(p_cont(a, b)),
        '|SMD|': fmt_smd(smd_cont(a, b)),
    }


def row_bin(name, col, positive=True):
    a_yes = int(v[col].sum()) if positive else int((~v[col].astype(bool)).sum())
    b_yes = int(nv[col].sum()) if positive else int((~nv[col].astype(bool)).sum())
    return {
        'Variable': name,
        f'Vaccinated (n={n_v})': fmt_pct(a_yes, n_v),
        f'Non-vaccinated (n={n_nv})': fmt_pct(b_yes, n_nv),
        'p': fmt_p(p_bin(a_yes, n_v, b_yes, n_nv)),
        '|SMD|': fmt_smd(smd_bin(a_yes, n_v, b_yes, n_nv)),
    }


def row_cat(label_prefix, col, levels):
    rows = []
    for lvl in levels:
        a_yes = (v[col] == lvl).sum()
        b_yes = (nv[col] == lvl).sum()
        rows.append(
            {
                'Variable': f'{label_prefix}: {lvl}',
                f'Vaccinated (n={n_v})': fmt_pct(a_yes, n_v),
                f'Non-vaccinated (n={n_nv})': fmt_pct(b_yes, n_nv),
                'p': fmt_p(p_bin(a_yes, n_v, b_yes, n_nv)),
                '|SMD|': fmt_smd(smd_bin(a_yes, n_v, b_yes, n_nv)),
            }
        )
    return rows


rows = []
rows.append(row_cont('Age at index, years', 'index_age'))
rows.append(row_cont('BMI, kg/m²', 'closest_bmi'))
rows.append(row_cont('Systolic BP, mmHg', 'SBP'))
rows.append(row_cont('Diastolic BP, mmHg', 'DBP'))

rows.extend(row_cat('Smoking', 'smoke', ['Never', 'Former', 'Current', 'Unknown']))

rows.append(row_bin('Seoul residence', 'is_seoul'))

# Surgery
df['fu_yrs'] = (df['최종추적일자'] - df['index_date']).dt.days / 365.25
v = df[df['접종여부'] == True]
nv = df[df['접종여부'] == False]
rows.append(row_cont('Follow-up, years (from index)', 'fu_yrs'))
rows.append(row_cont('Surgery year', '수술연도'))

# Surgery method (stored as Korean labels in primary_cohort_v3)
for lvl, label in [('원추절제술', 'Conization'), ('자궁절제술', 'Hysterectomy')]:
    a_yes = (v['수술방법'].astype(str) == lvl).sum()
    b_yes = (nv['수술방법'].astype(str) == lvl).sum()
    rows.append(
        {
            'Variable': f'Surgery type: {label}',
            f'Vaccinated (n={n_v})': fmt_pct(a_yes, n_v),
            f'Non-vaccinated (n={n_nv})': fmt_pct(b_yes, n_nv),
            'p': fmt_p(p_bin(a_yes, n_v, b_yes, n_nv)),
            '|SMD|': fmt_smd(smd_bin(a_yes, n_v, b_yes, n_nv)),
        }
    )

# Comorbidities
for lbl in CLASS_LABELS.values():
    rows.append(row_bin(f'Comorbidity (pre-index): {lbl}', f'comor_{lbl}'))
rows.append(row_bin('Comorbidity (pre-index): Any of 5', 'comor_any'))

# === New variables ===
rows.append({'Variable': '— HPV testing history —', f'Vaccinated (n={n_v})': '', f'Non-vaccinated (n={n_nv})': '', 'p': '', '|SMD|': ''})
rows.extend(
    row_cat('Pre-surgery HPV status', 'pre_surg_hpv',
            ['HR-HPV+', 'non-HR HPV+', 'HPV negative', 'No test'])
)
rows.extend(
    row_cat('Post-surgery HPV (first test, any time)', 'post_surg_anytime',
            ['HR-HPV+', 'non-HR HPV+', 'HPV negative', 'No test'])
)
rows.extend(
    row_cat('Post-surgery HPV (first test, before vaccine/index)', 'post_surg_prevac',
            ['HR-HPV+', 'non-HR HPV+', 'HPV negative', 'No test'])
)

rows.append({'Variable': '— Surgical pathology severity —', f'Vaccinated (n={n_v})': '', f'Non-vaccinated (n={n_nv})': '', 'p': '', '|SMD|': ''})
rows.extend(
    row_cat('Surgical pathology', 'surg_severity', SEVERITY_ORDER)
)

# Pre-vaccine hr-HPV+ (clearance subset eligibility) — derived from pre-vaccine baseline
# pre-vaccine HR via molecular records before index
def has_prevac_hr(pid, idx_dt):
    sub = mol_by_pid.get(pid)
    if sub is None:
        return False
    return bool((sub[sub['실시일자'] < idx_dt]['hpv_pos'] == True).any())


df['prevac_hr'] = df.apply(lambda r: has_prevac_hr(r['연구번호'], r['index_date']), axis=1)
v = df[df['접종여부'] == True]
nv = df[df['접종여부'] == False]
rows.append(row_bin('Pre-vaccine hr-HPV+ (clearance subset eligibility)', 'prevac_hr'))

# ------------------------------ output ------------------------------
table = pd.DataFrame(rows)
out_csv = ROOT / 'Data' / 'Table1B_v3_CohortB.csv'
table.to_csv(out_csv, index=False, encoding='utf-8-sig')
print(f'Wrote {out_csv.relative_to(ROOT)} ({len(table)} rows)')

# docx
doc = Document()
doc.add_heading(
    f'Table 1B. Cohort B baseline characteristics (≥2-dose + 3-mo landmark primary cohort; n={n_v + n_nv})',
    level=1,
)
t = doc.add_table(rows=len(table) + 1, cols=5)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
for i, h in enumerate(table.columns):
    hdr[i].text = str(h)
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9)
for i, row in enumerate(table.itertuples(index=False), start=1):
    for j, val in enumerate(row):
        t.cell(i, j).text = str(val)
        for p in t.cell(i, j).paragraphs:
            for r in p.runs:
                r.font.size = Pt(8.5)

note = doc.add_paragraph()
note.add_run(
    f'\nN={n_v + n_nv} ({n_v} vaccinated, {n_nv} non-vaccinated). '
    f'≥2 distinct vaccine dose dates required for vaccinated case eligibility; '
    'matched non-vaccinated controls of excluded cases dropped to preserve matched-set integrity. '
    '3-month landmark applied symmetrically: index date shifted to index + 90 days; '
    'patients with <90 days follow-up or with outcome event in first 90 days excluded. '
    'Pre-surgery HPV uses worst-grade molecular pathology before first surgery; '
    'post-surgery HPV uses first molecular result after surgery either ever or restricted to before vaccine/index. '
    'Surgical pathology severity uses highest-grade tissue diagnosis within ±30 days of first surgery.'
).font.size = Pt(8)

out_docx = ROOT / 'Data' / 'Table1B_v3_CohortB.docx'
doc.save(out_docx)
print(f'Wrote {out_docx.relative_to(ROOT)}')
