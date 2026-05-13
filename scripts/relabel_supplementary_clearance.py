"""
Relabel supplementary docx text/labels so all references to the legacy
"HPV reinfection" / "post-index hr-HPV detection (primary)" framing are
either:
  (a) renamed to "hr-HPV clearance" (where the table/figure now reports
      the n=292 two-consecutive-negative clearance co-primary), OR
  (b) renamed to "Post-index hr-HPV detection (sensitivity)" (where the
      table is the supplementary post-index detection sensitivity that
      conflates persistence and new acquisition — App-1/2/3).

Targets in docs/HPV_supplementary.docx:
  - P39 "HPV reinfection"                  -> "Post-index hr-HPV detection (sensitivity)"
  - Table 13 (S10 strict matching) R3      -> remove "HPV reinfection" row (Sens-D defends P1 only)
  - P102 Figure S5(b) caption              -> "hr-HPV clearance (vaccination covariate)"
  - P99 Figure S5 header sentence (PH p)   -> updated p value for clearance
  - any other "HPV reinfection" string remnant in body paragraphs
"""

from docx import Document
from copy import deepcopy

SRC = 'docs/HPV_supplementary.docx'
OUT = 'docs/HPV_supplementary.docx'   # in-place

doc = Document(SRC)

# ---------- 1) Paragraph relabels ----------
replace_map = [
    # Table S3 subheading
    ('HPV reinfection',
     'Post-index hr-HPV detection (sensitivity)'),
    # Figure S5(b) caption
    ('Post-index hr-HPV detection (vaccination covariate)',
     'hr-HPV clearance, n = 292 (vaccination covariate)'),
]

# Schoenfeld result line update (P104): keep recurrence p, replace
# "0.12 (hr-HPV detection)" with clearance p "0.028 (hr-HPV clearance, n = 292)".
sentence_map = [
    ('The Schoenfeld rank test p-values were 0.59 (recurrence) and 0.12 (hr-HPV detection)',
     'The Schoenfeld rank test p-values were 0.82 (recurrence) and 0.028 (hr-HPV clearance, n = 292; '
     'see Sens-B time-stratified decomposition in Supplementary Table S17 for the period-specific '
     'hazard ratios that explain this non-proportional pattern)'),
]

changed_paragraphs = 0
for p in doc.paragraphs:
    txt = p.text
    new_txt = txt
    for old, new in replace_map + sentence_map:
        if old in new_txt:
            new_txt = new_txt.replace(old, new)
    if new_txt != txt:
        # rewrite paragraph: clear runs then add one run preserving style
        for r in list(p.runs):
            r.text = ''
        # set first run text (simple approach since these are short labels)
        if p.runs:
            p.runs[0].text = new_txt
        else:
            p.add_run(new_txt)
        changed_paragraphs += 1

# ---------- 2) Table 13 (S10): drop "HPV reinfection" rows ----------
# Strict matching Sens-D defends P1 (recurrence) only.
# Find the table by header.
def find_table_with_header(doc, header_tokens):
    for ti, t in enumerate(doc.tables):
        hdr = [c.text.strip() for c in t.rows[0].cells]
        if all(any(tok.lower() in h.lower() for h in hdr) for tok in header_tokens):
            return ti, t
    return None, None

ti, strict_t = find_table_with_header(doc, ['Outcome', 'Design', 'HR'])
removed_rows = 0
if strict_t is not None:
    rows_to_remove = []
    for r in strict_t.rows:
        first = r.cells[0].text.strip().lower()
        if 'reinfection' in first or 'post-index' in first or 'hpv detection' in first:
            rows_to_remove.append(r)
    for r in rows_to_remove:
        r._element.getparent().remove(r._element)
        removed_rows += 1

# ---------- 3) Table S6 (vaccine-type interaction) ----------
# Make sure clearance is listed BEFORE the post-index detection sensitivity row.
# (already correct per current docx, just verify and reorder if needed)
ti2, vt_t = find_table_with_header(doc, ['Outcome', 'LRT', 'Gardasil 9 HR'])
if vt_t is not None:
    rows = list(vt_t.rows)
    headers = rows[0]
    data_rows = rows[1:]
    def row_key(r):
        s = r.cells[0].text.lower()
        if 'recurrence' in s:
            return 0
        if 'clearance' in s:
            return 1
        if 'detection' in s or 'reinfection' in s:
            return 2
        return 9
    sorted_data = sorted(data_rows, key=row_key)
    needs_reorder = any(a is not b for a, b in zip(data_rows, sorted_data))
    if needs_reorder:
        tbl = vt_t._element
        # remove existing data rows
        for r in data_rows:
            tbl.remove(r._element)
        # append in sorted order
        for r in sorted_data:
            tbl.append(r._element)

# Save
doc.save(OUT)
print(f'Updated paragraphs: {changed_paragraphs}')
print(f'Removed rows from strict matching table: {removed_rows}')
print(f'Vaccine-type interaction reordered: {needs_reorder if vt_t is not None else "no table"}')
print(f'Saved: {OUT}')
