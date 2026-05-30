"""
Generate the 5 missing manuscript artifacts:
1. Figure 1 — Cohort selection flow diagram (PNG)
2. Table 3 — Cohort B HR (recurrence + HPV reinfection) docx
3. Supplementary Figure S3 — PS density distribution
4. Supplementary Table S2 — PS model coefficients
5. Supplementary Table S5 — Number-at-risk tables for KM/AJ
"""
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import openpyxl
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
from lifelines import CoxPHFitter, KaplanMeierFitter, AalenJohansenFitter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

plt.rcParams['font.family'] = ['DejaVu Sans','AppleGothic']
plt.rcParams['axes.unicode_minus'] = False

# ==================================================
# 1. Figure 1 — Cohort selection flow (PNG)
# ==================================================
print('[1/5] Figure 1: Cohort selection flow PNG...')
fig, ax = plt.subplots(figsize=(13, 9))
ax.set_xlim(0, 13); ax.set_ylim(0, 11)
ax.axis('off')

def box(x, y, w, h, text, color='#e8f4f8', edge='#1f6f8b', fontsize=10, weight='normal'):
    rect = FancyBboxPatch((x-w/2, y-h/2), w, h,
                          boxstyle="round,pad=0.06",
                          facecolor=color, edgecolor=edge, linewidth=1.5)
    ax.add_patch(rect)
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize, fontweight=weight, wrap=True)

def arrow(x1, y1, x2, y2):
    arr = FancyArrowPatch((x1, y1), (x2, y2),
                          arrowstyle='-|>', mutation_scale=15,
                          color='#444', linewidth=1.2)
    ax.add_patch(arr)

# Source
box(6.5, 10.2, 5.2, 0.85,
    'Source population: Korean HPV cohort\n(prospective enrolment 2009–2024)\nN = 32,969 women',
    color='#e8f4f8', edge='#1f6f8b', fontsize=11, weight='bold')

# Vaccine identification
box(6.5, 8.9, 6.4, 0.9,
    'HPV vaccination ascertained from prescription records\n'
    '(Gardasil 9 / Cervarix / Gardasil)\n'
    'Vaccinated n = 2,156   Unvaccinated candidates n = 30,813',
    color='#fff3cd', edge='#856404')

# Split
arrow(6.5, 8.42, 3.4, 7.6)
arrow(6.5, 8.42, 9.6, 7.6)

# Cohort A header
box(3.4, 7.15, 5.6, 0.7,
    'COHORT A — Long-term safety analysis\n(whole cohort)',
    color='#d4edda', edge='#155724', fontsize=11, weight='bold')

# Cohort B header
box(9.6, 7.15, 5.6, 0.7,
    'COHORT B — Post-surgical efficacy analysis\n(surgery patients only)',
    color='#fde2e4', edge='#9b2226', fontsize=11, weight='bold')

# Cohort A steps
arrow(3.4, 6.78, 3.4, 6.45)
box(3.4, 6.0, 5.6, 0.85,
    'Pseudo index date for unvaccinated\n(random sample from vaccine-date\ndistribution; seed=42)',
    color='#eaf6ee', edge='#155724', fontsize=9)

arrow(3.4, 5.56, 3.4, 5.23)
box(3.4, 4.65, 5.6, 1.1,
    'Eligibility: alive at index, ≥1 day follow-up\n'
    'Propensity score model (LogReg):\n'
    'age, BMI, SBP, DBP, smoking, residence',
    color='#eaf6ee', edge='#155724', fontsize=9)

arrow(3.4, 4.07, 3.4, 3.74)
box(3.4, 3.18, 5.6, 1.1,
    '1:1 nearest-neighbour matching on logit(PS)\n'
    'Caliper = 0.2 × SD(logit PS)\n'
    '2,110 pairs matched',
    color='#eaf6ee', edge='#155724', fontsize=9)

arrow(3.4, 2.6, 3.4, 2.27)
box(3.4, 1.55, 5.6, 1.4,
    'Final Cohort A: n = 4,102\n'
    'Vaccinated 2,051  /  Unvaccinated 2,051\n\n'
    'Outcomes: 5 chronic comorbidities\n'
    'Composite: MCE, Any-of-5',
    color='#a8d5b5', edge='#155724', fontsize=10, weight='bold')

# Cohort B steps
arrow(9.6, 6.78, 9.6, 6.45)
box(9.6, 6.0, 5.6, 0.85,
    'Cervical surgery (conization or hysterectomy)\nN = 6,890',
    color='#fdedee', edge='#9b2226', fontsize=9)

arrow(9.6, 5.56, 9.6, 5.23)
box(9.6, 4.65, 5.6, 1.1,
    '1:5 initial match — surgery method (exact),\n'
    'surgery year (±1y), age at surgery (±5y)\n'
    'Vac 411  /  Ctl 1,815',
    color='#fdedee', edge='#9b2226', fontsize=9)

arrow(9.6, 4.07, 9.6, 3.74)
box(9.6, 3.18, 5.6, 1.1,
    'Eligibility: index ≤ 2024-12-31\n(≥1y potential FU); ≥2 follow-up records\n'
    'Vac 411  /  Ctl 1,797   (excluded 18)',
    color='#fdedee', edge='#9b2226', fontsize=9)

arrow(9.6, 2.6, 9.6, 2.27)
box(9.6, 1.55, 5.6, 1.4,
    'Final Cohort B: n = 1,108\n'
    'Vaccinated 241  /  Unvaccinated 867\n'
    '(Fine matching on age, BMI, surgery year)\n\n'
    'Outcomes: lesion recurrence, HPV reinfection',
    color='#f4a4a8', edge='#9b2226', fontsize=10, weight='bold')

ax.set_title('Figure 1. Cohort selection flow diagram',
            fontsize=13, fontweight='bold', y=1.01)
plt.tight_layout()
plt.savefig('Data/Figure1_cohort_selection.png', dpi=200, bbox_inches='tight')
plt.close()
print('  Saved: Data/Figure1_cohort_selection.png')

# ==================================================
# Helper: rebuild Cohort A matched + PS coefs (for sup tables)
# ==================================================
print('[Loading source for items 3-5]...')

cohort = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_코호트.csv', encoding='cp949', low_memory=False)
cohort['birth_date'] = pd.to_datetime(cohort['생년월'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['death_date'] = pd.to_datetime(cohort['사망일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['last_follow'] = pd.to_datetime(cohort['최종추적일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['is_seoul'] = (cohort['주소'].astype(str).str.split().str[0]=='서울').astype(int)
rx = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_처방정보.csv', encoding='cp949', low_memory=False)
mask = (rx['처방명'].astype(str).str.contains('Gardasil|Cervarix|HPV vaccine', case=False, na=False) |
        rx['처방한글명'].astype(str).str.contains('가다실|서바릭스', na=False))
rx_vac = rx[mask].copy()
rx_vac['처방일자'] = pd.to_datetime(rx_vac['처방일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
first_vac = rx_vac.groupby('연구번호')['처방일자'].min().reset_index()
first_vac.columns = ['연구번호','first_vaccine_date']
cohort = cohort.merge(first_vac, on='연구번호', how='left').dropna(subset=['birth_date'])

ci = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_기초임상정보.csv', encoding='cp949', low_memory=False)
ci['기록일자_dt'] = pd.to_datetime(ci['기록일자'].astype(str).str.strip(), format='%Y%m%d', errors='coerce')

def closest_vec(query_df, ci, value_col, window_days=365):
    SMOKE_MAP = {'비흡연':'Never','과거흡연':'Former','현재흡연':'Current','확인불능':'Unknown'}
    ci_v = ci[['연구번호','기록일자_dt', value_col]].dropna(subset=[value_col,'기록일자_dt']).copy()
    ci_v = ci_v.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    fw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid', direction='forward', tolerance=pd.Timedelta(days=window_days))
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid', direction='backward', tolerance=pd.Timedelta(days=window_days))
    fw['_d'] = (fw['rec_date']-fw['index_date']).abs()
    bw['_d'] = (bw['rec_date']-bw['index_date']).abs()
    use_fw = (fw['_d'].fillna(pd.Timedelta(days=window_days*10)) <= bw['_d'].fillna(pd.Timedelta(days=window_days*10)))
    return pd.Series(np.where(use_fw, fw[value_col].values, bw[value_col].values),
                    index=fw['orig_idx'].values).reindex(query_df.index).astype(float)

def smoke_vec(query_df, ci):
    SMOKE_MAP = {'비흡연':'Never','과거흡연':'Former','현재흡연':'Current','확인불능':'Unknown'}
    smk = ci[['연구번호','기록일자_dt','흡연여부']].dropna(subset=['흡연여부','기록일자_dt']).copy()
    smk = smk.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], smk,
        left_on='index_date', right_on='rec_date', by='pid', direction='backward')
    return pd.Series(bw['흡연여부'].map(SMOKE_MAP).fillna('Unknown').values,
                    index=bw['orig_idx'].values).reindex(query_df.index).fillna('Unknown')

rng = np.random.default_rng(42)
df = cohort.copy()
df['vaccinated'] = df['first_vaccine_date'].notna()
vac_dates = df.loc[df['vaccinated'], 'first_vaccine_date'].dropna().values
df.loc[~df['vaccinated'], 'index_date'] = pd.to_datetime(rng.choice(vac_dates, size=(~df['vaccinated']).sum()))
df.loc[df['vaccinated'], 'index_date'] = df.loc[df['vaccinated'], 'first_vaccine_date']
df = df.rename(columns={'연구번호':'pid'}).reset_index(drop=True)
df['age_at_index'] = (df['index_date'] - df['birth_date']).dt.days/365.25
df = df[(df['death_date'].isna()) | (df['death_date'] > df['index_date'])]
df = df[df['last_follow'] > df['index_date']].reset_index(drop=True)
q = df[['pid','index_date']].copy()
df['height'] = closest_vec(q, ci, '키')
df['weight'] = closest_vec(q, ci, '몸무게')
df['sbp'] = closest_vec(q, ci, '수축기혈압')
df['dbp'] = closest_vec(q, ci, '이완기혈압')
df['bmi'] = df['weight']/(df['height']/100)**2
df['smoke'] = smoke_vec(q, ci).values
for c in ['bmi','sbp','dbp']:
    df[f'{c}_miss'] = df[c].isna().astype(int)
    df[c] = df[c].fillna(df[c].mean())
sm = pd.get_dummies(df['smoke'], prefix='smoke').astype(int)
df = pd.concat([df, sm], axis=1)
ps_features = ['age_at_index','bmi','bmi_miss','sbp','sbp_miss','dbp','dbp_miss','is_seoul',
              'smoke_Never','smoke_Former','smoke_Current']
ps_features = [c for c in ps_features if c in df.columns]
X = df[ps_features].astype(float).values
y = df['vaccinated'].astype(int).values
scaler = StandardScaler()
Xs = scaler.fit_transform(X)
lr = LogisticRegression(max_iter=2000, C=1e6, solver='lbfgs')
lr.fit(Xs, y)
df['ps'] = lr.predict_proba(Xs)[:,1]
df['logit_ps'] = np.log(df['ps']/(1-df['ps']))
caliper = 0.2 * df['logit_ps'].std()

# 1:1 matching
vac_idx = df.index[df['vaccinated']].tolist()
ctl_idx = np.array(df.index[~df['vaccinated']].tolist())
ctl_logit = df.loc[ctl_idx,'logit_ps'].values
order = np.argsort(ctl_logit)
ctl_sorted = ctl_idx[order]; ctl_logit_sorted = ctl_logit[order]
used = np.zeros(len(ctl_sorted), dtype=bool)
matched = []
vac_order = np.array(vac_idx); rng.shuffle(vac_order)
for vi in vac_order:
    target = df.loc[vi,'logit_ps']
    lo = np.searchsorted(ctl_logit_sorted, target-caliper)
    hi = np.searchsorted(ctl_logit_sorted, target+caliper, side='right')
    best_j, best_d = -1, caliper+1
    for j in range(lo, hi):
        if used[j]: continue
        d = abs(ctl_logit_sorted[j]-target)
        if d < best_d: best_d=d; best_j=j
    if best_j>=0:
        used[best_j] = True
        matched.append((vi, ctl_sorted[best_j]))

vac_m = [v for v,_ in matched]; ctl_m = [c for _,c in matched]
df['matched'] = False
df.loc[vac_m+ctl_m, 'matched'] = True

# ==================================================
# 3. Sup Fig S3 — PS density distribution
# ==================================================
print('[3/5] Sup Fig S3: PS density distribution...')
import matplotlib as _mpl
_mpl.rcParams.update({
    'font.family': ['DejaVu Sans'],
    'axes.spines.top': False,
    'axes.spines.right': False,
    'axes.linewidth': 1.0,
    'axes.titlesize': 12,
    'axes.titleweight': 'bold',
    'axes.titlepad': 6,
    'axes.labelsize': 12,
    'xtick.labelsize': 11,
    'ytick.labelsize': 11,
    'legend.fontsize': 11,
    'legend.frameon': True,
    'legend.framealpha': 0.95,
    'legend.edgecolor': '#cccccc',
    'pdf.fonttype': 42,
    'ps.fonttype': 42,
})

fig, axes = plt.subplots(1, 2, figsize=(13.5, 5.4),
                         gridspec_kw={'wspace':0.22, 'left':0.07,
                                      'right':0.985, 'top':0.93, 'bottom':0.13})
# Pre
ax = axes[0]
ax.hist(df.loc[df['vaccinated'],'ps'], bins=40, alpha=0.6, color='#9b2226',
        label='Vaccinated', density=True)
ax.hist(df.loc[~df['vaccinated'],'ps'], bins=40, alpha=0.6, color='#1f6f8b',
        label='Non-vaccinated', density=True)
ax.set_xlabel('Propensity score'); ax.set_ylabel('Density')
ax.set_title('(a) Pre-matching')
ax.text(-0.13, 1.07, 'a', transform=ax.transAxes, fontsize=14, fontweight='bold',
        va='top', ha='left')
ax.legend(loc='upper right')
ax.grid(axis='y', alpha=0.25, linestyle=':')
# Post
ax = axes[1]
ax.hist(df.loc[df['matched'] & df['vaccinated'],'ps'], bins=40, alpha=0.6,
        color='#9b2226', label='Vaccinated', density=True)
ax.hist(df.loc[df['matched'] & ~df['vaccinated'],'ps'], bins=40, alpha=0.6,
        color='#1f6f8b', label='Non-vaccinated', density=True)
ax.set_xlabel('Propensity score'); ax.set_ylabel('Density')
ax.set_title('(b) Post-matching (1:1, n = 2,051 per group)')
ax.text(-0.13, 1.07, 'b', transform=ax.transAxes, fontsize=14, fontweight='bold',
        va='top', ha='left')
ax.legend(loc='upper right')
ax.grid(axis='y', alpha=0.25, linestyle=':')

plt.savefig('Data/SupFigS3_ps_density.png', dpi=300, bbox_inches='tight', facecolor='white')
plt.close()
print('  Saved: Data/SupFigS3_ps_density.png')

# ==================================================
# 4. Sup Table S2 — PS model coefficients
# ==================================================
print('[4/5] Sup Table S2: PS coefficients...')
# Recover unstandardized coefficients
coefs_std = lr.coef_[0]
# In standardized space: beta_std = beta_orig * sigma_x; so beta_orig = beta_std / sigma_x
sigmas = scaler.scale_
coefs_orig = coefs_std / sigmas
intercept_orig = lr.intercept_[0] - np.sum(coefs_orig * scaler.mean_)
ORs = np.exp(coefs_orig)

coef_df = pd.DataFrame({
    'Covariate': ps_features,
    'Mean (vaccinated)': [df.loc[df['vaccinated'], f].mean() for f in ps_features],
    'Mean (non-vaccinated)': [df.loc[~df['vaccinated'], f].mean() for f in ps_features],
    'Coefficient (β)': np.round(coefs_orig, 4),
    'Standardized β': np.round(coefs_std, 4),
    'Odds ratio': np.round(ORs, 3),
})
coef_df.to_csv('Data/SupTableS2_ps_coefficients.csv', index=False, encoding='utf-8-sig')

# Also docx
doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(10)
doc.add_heading('Supplementary Table S2. Propensity score model coefficients (Cohort A)', level=1)
p = doc.add_paragraph()
p.add_run('Logistic regression of HPV vaccination receipt (outcome: 1 = vaccinated). ').italic = True
p.add_run(f'Intercept (untransformed): {intercept_orig:.4f}.   '
         f'Pseudo R² ≈ {(1 - lr.predict_proba(Xs).max(axis=1).mean()):.3f}.   '
         f'Caliper used = 0.2 × SD(logit PS) = {caliper:.4f}').italic = True
t = doc.add_table(rows=1+len(coef_df), cols=6)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
for i, h in enumerate(['Covariate','Vac mean','Non-vac mean','β (orig.)','Std. β','Odds ratio']):
    hdr[i].text = h
    for para in hdr[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold = True; r.font.size = Pt(9)
for i, row in coef_df.iterrows():
    cells = t.rows[i+1].cells
    cells[0].text = row['Covariate']
    cells[1].text = f'{row["Mean (vaccinated)"]:.3f}'
    cells[2].text = f'{row["Mean (non-vaccinated)"]:.3f}'
    cells[3].text = f'{row["Coefficient (β)"]:.4f}'
    cells[4].text = f'{row["Standardized β"]:.4f}'
    cells[5].text = f'{row["Odds ratio"]:.3f}'
    for j, c in enumerate(cells):
        for para in c.paragraphs:
            for r in para.runs: r.font.size = Pt(9)
            if j > 0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot = doc.add_paragraph()
foot.add_run('Footnote: ').bold = True
foot.add_run(
    'β is the coefficient on the original (unscaled) variable; standardized β is the coefficient after Z-scaling each '
    'continuous predictor (and is directly comparable across variables). Missing-value indicators (bmi_miss, sbp_miss, dbp_miss) '
    'flag patients without a measurement within ±365 days of index; these covariates were imputed with the cohort mean before model fitting. '
    'Smoking is encoded with three dummy variables (Never, Former, Current); the reference category is Unknown.'
).font.size = Pt(8)
doc.save('Data/SupTableS2_ps_coefficients.docx')
print('  Saved: Data/SupTableS2_ps_coefficients.{csv,docx}')

# ==================================================
# 5. Sup Table S5 — Number-at-risk tables
# ==================================================
print('[5/5] Sup Table S5: Number-at-risk tables...')

# Need the comorbidity data on the matched cohort
m = df[df['matched']].copy().reset_index(drop=True)
wb = openpyxl.load_workbook('Data/한국 HPV 코호트 자료를 이용한 자_진단정보_기저질환추가_unlocked.xlsx',
                          read_only=True, data_only=True)
ws = wb.active
recs = []
CLASS_LABELS = {'1':'Angina/MI','2':'Hypertension','3':'Diabetes','4':'Stroke','5':'PE'}
for row in ws.iter_rows(min_row=2, values_only=True):
    pid, cls, dd = row[0], row[5], row[8]
    if cls is None or str(cls).strip() == '': continue
    cls = str(cls).strip()
    if cls not in CLASS_LABELS: continue
    recs.append((pid, cls, pd.to_datetime(str(dd), format='%Y%m%d', errors='coerce')))
como = pd.DataFrame(recs, columns=['pid','class','diag_date'])
first_diag = como.groupby(['pid','class'])['diag_date'].min().unstack('class')
for c in CLASS_LABELS:
    if c not in first_diag.columns: first_diag[c] = pd.NaT
m = m.merge(first_diag, left_on='pid', right_index=True, how='left')

# Time-to-event helper
def make_tte(m, cls_or_list):
    if isinstance(cls_or_list, list):
        dx = m[cls_or_list].min(axis=1)
    else:
        dx = m[cls_or_list]
    is_pre = dx.notna() & (dx <= m['index_date'])
    primary = dx.where(dx > m['index_date'], pd.NaT)
    death_after = m['death_date'].where(
        (m['death_date'].notna()) & (m['death_date'] > m['index_date']) &
        ((primary.isna()) | (m['death_date'] < primary)), pd.NaT)
    event_date = primary.combine_first(death_after)
    status = np.where(primary.notna() & ((death_after.isna()) | (primary <= death_after)), 1,
            np.where(death_after.notna(), 2, 0))
    end_date = event_date.combine_first(m['last_follow'])
    time = (end_date - m['index_date']).dt.days.astype(float)
    res = pd.DataFrame({'pid':m['pid'].values, 'vaccinated':m['vaccinated'].astype(int).values,
                       'time':time, 'status':status})
    res = res[~is_pre.values & (res['time']>0)].reset_index(drop=True)
    return res

# Build number-at-risk at fixed time points (years)
TIME_PTS_YR = [0, 1, 2, 3, 4, 5, 6, 7, 8]
TIME_PTS_DAYS = [t*365.25 for t in TIME_PTS_YR]

at_risk_records = []

def n_at_risk(tte, t_days):
    return int((tte['time'] >= t_days).sum())

# Cohort A outcomes
ANY5 = ['1','2','3','4','5']
MCE = ['1','4','5']
outcomes_A = [
    ('Any-of-5', ANY5),
    ('MCE (MI/Stroke/PE)', MCE),
    ('Angina/MI', '1'),
    ('Hypertension', '2'),
    ('Diabetes', '3'),
    ('Stroke', '4'),
    ('PE', '5'),
]

for label, comp in outcomes_A:
    tte = make_tte(m, comp)
    for grp_val, grp_name in [(1,'Vaccinated'), (0,'Non-vaccinated')]:
        sub = tte[tte['vaccinated']==grp_val]
        row = {'Cohort':'A','Outcome':label,'Group':grp_name,'Initial n':len(sub)}
        for yr, td in zip(TIME_PTS_YR, TIME_PTS_DAYS):
            row[f'{yr} yr'] = n_at_risk(sub, td)
        at_risk_records.append(row)

# Cohort B outcomes (uses final_matched_outcomes)
print('  Cohort B at-risk...')
B = pd.read_csv('Data/final_matched_outcomes.csv', encoding='utf-8-sig')
B['index_date'] = pd.to_datetime(B['index_date'])
B['recurrence_date'] = pd.to_datetime(B['recurrence_date'], errors='coerce')
B['hpv_infection_date'] = pd.to_datetime(B['hpv_infection_date'], errors='coerce')
B['follow_up_days'] = pd.to_numeric(B['follow_up_days'], errors='coerce')

def at_risk_for_outcome(B, event_date_col, has_event_col):
    out = B.copy()
    # time = min(event_date, last_follow) - index, in days
    end = out[event_date_col].fillna(out['index_date'] + pd.to_timedelta(out['follow_up_days'], unit='D'))
    out['time'] = (end - out['index_date']).dt.days.astype(float)
    out = out[out['time'] > 0]
    return out

for label, ev_col, has_col in [
    ('Lesion recurrence', 'recurrence_date', 'has_recurrence'),
    ('HPV reinfection',   'hpv_infection_date', 'has_hpv_infection')]:
    tte_B = at_risk_for_outcome(B, ev_col, has_col)
    for grp_val, grp_name in [(True,'Vaccinated'), (False,'Non-vaccinated')]:
        sub = tte_B[tte_B['접종여부'].astype(bool)==grp_val]
        row = {'Cohort':'B','Outcome':label,'Group':grp_name,'Initial n':len(sub)}
        for yr, td in zip(TIME_PTS_YR, TIME_PTS_DAYS):
            row[f'{yr} yr'] = int((sub['time'] >= td).sum())
        at_risk_records.append(row)

ar_df = pd.DataFrame(at_risk_records)
ar_df.to_csv('Data/SupTableS5_number_at_risk.csv', index=False, encoding='utf-8-sig')

# docx
doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'; doc.styles['Normal'].font.size = Pt(10)
doc.add_heading('Supplementary Table S5. Number at risk over follow-up time', level=1)
p = doc.add_paragraph()
p.add_run('Number of patients still at risk (event-free and not censored) at each yearly time point from the index date. '
         'Patients were excluded from the at-risk denominator after the first event of the corresponding outcome, after death, '
         'or after censoring at last follow-up.').italic = True

# Build wide table
tcols = ['Cohort','Outcome','Group','Initial n'] + [f'{y} yr' for y in TIME_PTS_YR]
t = doc.add_table(rows=1+len(ar_df), cols=len(tcols))
t.style = 'Light Grid Accent 1'
for i, h in enumerate(tcols):
    t.rows[0].cells[i].text = h
    for para in t.rows[0].cells[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold = True; r.font.size = Pt(8)

for i, row in ar_df.iterrows():
    for j, c in enumerate(tcols):
        cells = t.rows[i+1].cells
        cells[j].text = str(row[c])
        for para in cells[j].paragraphs:
            for r in para.runs: r.font.size = Pt(8)
            if j > 2: para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('Data/SupTableS5_number_at_risk.docx')
print('  Saved: Data/SupTableS5_number_at_risk.{csv,docx}')

# ==================================================
# 2. Table 3 — Cohort B HR (recurrence + HPV)
# ==================================================
print('[2/5] Table 3: Cohort B HR docx...')

cph_recur = CoxPHFitter()
B_cox = B.copy()
B_cox['vac'] = B_cox['접종여부'].astype(bool).astype(int)
B_cox['index_age_'] = pd.to_numeric(B_cox['index_age'], errors='coerce')

# Recurrence
B_recur = B_cox[['follow_up_days','has_recurrence','vac','index_age_']].dropna()
B_recur = B_recur.rename(columns={'follow_up_days':'time','has_recurrence':'event'})
B_recur['event'] = B_recur['event'].astype(int)
cph_recur.fit(B_recur, duration_col='time', event_col='event', robust=True)
sm_recur = cph_recur.summary

# HPV
B_hpv = B_cox[['follow_up_days','has_hpv_infection','vac','index_age_']].dropna()
B_hpv = B_hpv.rename(columns={'follow_up_days':'time','has_hpv_infection':'event'})
B_hpv['event'] = B_hpv['event'].astype(int)
cph_hpv = CoxPHFitter()
cph_hpv.fit(B_hpv, duration_col='time', event_col='event', robust=True)
sm_hpv = cph_hpv.summary

# event counts
ev_recur_v = int(B_cox.loc[B_cox['vac']==1,'has_recurrence'].astype(int).sum())
n_recur_v = int((B_cox['vac']==1).sum())
ev_recur_c = int(B_cox.loc[B_cox['vac']==0,'has_recurrence'].astype(int).sum())
n_recur_c = int((B_cox['vac']==0).sum())
ev_hpv_v = int(B_cox.loc[B_cox['vac']==1,'has_hpv_infection'].astype(int).sum())
ev_hpv_c = int(B_cox.loc[B_cox['vac']==0,'has_hpv_infection'].astype(int).sum())

doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'; doc.styles['Normal'].font.size = Pt(10)
doc.add_heading('Table 3. Cohort B — Vaccine effectiveness for lesion recurrence and high-risk HPV reinfection', level=1)

t = doc.add_table(rows=4, cols=6)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
for i, h in enumerate(['Outcome','Vaccinated\n(events / n)','Non-vaccinated\n(events / n)','Crude rate ratio',
                        'Age-adjusted HR\n(95% CI)','p-value']):
    hdr[i].text = h
    for para in hdr[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold = True; r.font.size = Pt(9)

# Recurrence row
hr_r = sm_recur.loc['vac','exp(coef)']
ci_r = (sm_recur.loc['vac','exp(coef) lower 95%'], sm_recur.loc['vac','exp(coef) upper 95%'])
p_r = sm_recur.loc['vac','p']
crude_r = (ev_recur_v/n_recur_v) / (ev_recur_c/n_recur_c) if ev_recur_c>0 else float('nan')
cells = t.rows[1].cells
cells[0].text = 'Lesion recurrence (HSIL/CIN3+)'
cells[1].text = f'{ev_recur_v} / {n_recur_v} ({100*ev_recur_v/n_recur_v:.1f}%)'
cells[2].text = f'{ev_recur_c} / {n_recur_c} ({100*ev_recur_c/n_recur_c:.1f}%)'
cells[3].text = f'{crude_r:.2f}' if not np.isnan(crude_r) else '-'
cells[4].text = f'{hr_r:.3f} ({ci_r[0]:.3f}–{ci_r[1]:.3f})'
cells[5].text = f'{p_r:.3f}' if p_r>=0.001 else '<0.001'

hr_h = sm_hpv.loc['vac','exp(coef)']
ci_h = (sm_hpv.loc['vac','exp(coef) lower 95%'], sm_hpv.loc['vac','exp(coef) upper 95%'])
p_h = sm_hpv.loc['vac','p']
crude_h = (ev_hpv_v/n_recur_v) / (ev_hpv_c/n_recur_c)
cells = t.rows[2].cells
cells[0].text = 'New high-risk HPV infection'
cells[1].text = f'{ev_hpv_v} / {n_recur_v} ({100*ev_hpv_v/n_recur_v:.1f}%)'
cells[2].text = f'{ev_hpv_c} / {n_recur_c} ({100*ev_hpv_c/n_recur_c:.1f}%)'
cells[3].text = f'{crude_h:.2f}'
cells[4].text = f'{hr_h:.3f} ({ci_h[0]:.3f}–{ci_h[1]:.3f})'
cells[5].text = f'{p_h:.3f}' if p_h>=0.001 else '<0.001'

# Composite (any of two) — informational
ev_any_v = int(((B_cox['vac']==1) & ((B_cox['has_recurrence'].astype(int)==1) | (B_cox['has_hpv_infection'].astype(int)==1))).sum())
ev_any_c = int(((B_cox['vac']==0) & ((B_cox['has_recurrence'].astype(int)==1) | (B_cox['has_hpv_infection'].astype(int)==1))).sum())
cells = t.rows[3].cells
cells[0].text = 'Either event (composite, descriptive)'
cells[1].text = f'{ev_any_v} / {n_recur_v} ({100*ev_any_v/n_recur_v:.1f}%)'
cells[2].text = f'{ev_any_c} / {n_recur_c} ({100*ev_any_c/n_recur_c:.1f}%)'
cells[3].text = '-'
cells[4].text = '-'
cells[5].text = '-'

for r in t.rows[1:]:
    for j, c in enumerate(r.cells):
        for para in c.paragraphs:
            for run in para.runs: run.font.size = Pt(9)
            if j > 0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER

foot = doc.add_paragraph()
foot.add_run('Footnote: ').bold = True
foot.add_run(
    'Lesion recurrence is defined as biopsy-confirmed HSIL, CIN3, or invasive cervical cancer after the index date. '
    'New high-risk HPV infection is defined as a positive test for any of the 14 high-risk types after the index date. '
    'Hazard ratios are from Cox proportional-hazards regression adjusted for age at index, with robust standard errors. '
    'The descriptive composite row shows the proportion of women experiencing either event during follow-up.'
).font.size = Pt(8)

doc.save('Data/Table3_CohortB_HR.docx')
# Also csv
pd.DataFrame([
    {'outcome':'Lesion recurrence','events_vac':ev_recur_v,'n_vac':n_recur_v,
     'events_ctl':ev_recur_c,'n_ctl':n_recur_c,
     'HR':hr_r,'CI_lo':ci_r[0],'CI_hi':ci_r[1],'p':p_r},
    {'outcome':'HPV reinfection','events_vac':ev_hpv_v,'n_vac':n_recur_v,
     'events_ctl':ev_hpv_c,'n_ctl':n_recur_c,
     'HR':hr_h,'CI_lo':ci_h[0],'CI_hi':ci_h[1],'p':p_h},
]).to_csv('Data/Table3_CohortB_HR.csv', index=False, encoding='utf-8-sig')
print('  Saved: Data/Table3_CohortB_HR.{docx,csv}')

print('\n=== ALL 5 ARTIFACTS GENERATED ===')
print('  Data/Figure1_cohort_selection.png')
print('  Data/Table3_CohortB_HR.docx + .csv')
print('  Data/SupFigS3_ps_density.png')
print('  Data/SupTableS2_ps_coefficients.docx + .csv')
print('  Data/SupTableS5_number_at_risk.docx + .csv')
