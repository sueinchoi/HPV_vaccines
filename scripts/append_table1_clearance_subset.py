"""
Extend Table 1 with a baseline-characteristics block for the
Cohort B clearance analytic subset (women with documented pre-vaccine
hr-HPV positivity, n = 292: vac 110 / non-vac 182). This block lets
reviewers verify that the matching balance is preserved after the
pre-vaccine HPV+ restriction that defines the clearance co-primary
outcome.

Reads the existing Data/Table1_BaselineCharacteristics_unified.{csv,docx}
and appends a new block "CohortB_clearance" to both, with the same
variable rows as the existing CohortB_post block.

Output: in-place update of
  Data/Table1_BaselineCharacteristics_unified.csv
  Data/Table1_BaselineCharacteristics_unified.docx
"""
import warnings; warnings.filterwarnings('ignore')
import sys, pandas as pd, numpy as np
from scipy.stats import fisher_exact, ttest_ind
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- Load matched cohort + clearance subset ----------
B  = pd.read_csv('Data/final_matched_cohort.csv', encoding='utf-8-sig')
BC = pd.read_csv('Data/CohortB_Clearance_Analytic.csv', encoding='utf-8-sig')

# Restrict to clearance subset
B = B[B['연구번호'].isin(BC['연구번호'])].copy()
B['vaccinated'] = B['접종여부'].astype(bool)
B['index_date']  = pd.to_datetime(B['index_date'])
B['최종추적일자']  = pd.to_datetime(B['최종추적일자'])
B['생년월일']      = pd.to_datetime(B['생년월일'])
B['follow_up_days'] = (B['최종추적일자'] - B['index_date']).dt.days
B['birth_year']  = B['생년월일'].dt.year
B['index_year']  = B['index_date'].dt.year
B['hysterectomy'] = (B['수술방법'] == '자궁절제술').astype(int)
B['conization']   = (B['수술방법'] == '원추절제술').astype(int)
B['index_age']    = pd.to_numeric(B['index_age'], errors='coerce')
B['closest_bmi']  = pd.to_numeric(B['closest_bmi'], errors='coerce')
B['수술시나이']    = pd.to_numeric(B['수술시나이'], errors='coerce')
B['수술연도']      = pd.to_numeric(B['수술연도'], errors='coerce')
B['수술_접종_간격일'] = pd.to_numeric(B['수술_접종_간격일'], errors='coerce')

vac = B[B['vaccinated']].copy()
ctl = B[~B['vaccinated']].copy()
print(f'Clearance subset: vac={len(vac)}, non-vac={len(ctl)}')

# ---------- Helpers ----------
def smd_cont(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a) < 2 or len(b) < 2: return np.nan
    pooled = np.sqrt((a.var(ddof=1) + b.var(ddof=1)) / 2)
    return (a.mean() - b.mean()) / pooled if pooled > 0 else np.nan

def smd_bin(a, b):
    a = pd.Series(a).dropna().astype(float); b = pd.Series(b).dropna().astype(float)
    if len(a) == 0 or len(b) == 0: return np.nan
    p1, p2 = a.mean(), b.mean()
    pooled = np.sqrt((p1*(1-p1) + p2*(1-p2)) / 2)
    return (p1 - p2) / pooled if pooled > 0 else np.nan

def fmt_cont(s):
    s = pd.Series(s).dropna()
    return '-' if len(s) == 0 else f'{s.mean():.2f} ± {s.std():.2f}'

def fmt_pct(n_yes, n_total):
    return f'{int(n_yes)} ({100*n_yes/n_total:.1f}%)' if n_total > 0 else '-'

def p_cont(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a) < 2 or len(b) < 2: return np.nan
    return ttest_ind(a, b, equal_var=False).pvalue

def p_bin(a_yes, a_n, b_yes, b_n):
    if a_n == 0 or b_n == 0: return np.nan
    return fisher_exact([[a_yes, a_n - a_yes], [b_yes, b_n - b_yes]])[1]

def fmt_p(p):
    return '-' if pd.isna(p) else ('<0.001' if p < 0.001 else f'{p:.3f}')

def fmt_smd(s):
    return '-' if pd.isna(s) else f'{abs(s):.3f}'

def row_cont(label, vac_s, ctl_s):
    return [label, fmt_cont(vac_s), fmt_cont(ctl_s),
            fmt_p(p_cont(vac_s, ctl_s)), fmt_smd(smd_cont(vac_s, ctl_s))]

def row_bin(label, vac_yes, vac_n, ctl_yes, ctl_n):
    a_yes, b_yes = int(vac_yes), int(ctl_yes)
    return [label, fmt_pct(a_yes, vac_n), fmt_pct(b_yes, ctl_n),
            fmt_p(p_bin(a_yes, vac_n, b_yes, ctl_n)),
            fmt_smd(smd_bin([1]*a_yes + [0]*(vac_n-a_yes), [1]*b_yes + [0]*(ctl_n-b_yes)))]

def section_header(label):
    return [label, '', '', '', '']

# ---------- Build rows ----------
n_v, n_c = len(vac), len(ctl)
rows = []
rows.append([f'N (analytic subset, pre-vaccine hr-HPV+)',
             str(n_v), str(n_c), '', ''])
rows.append(section_header('Demographics'))
rows.append(row_cont('  Age at index, years', vac['index_age'], ctl['index_age']))
rows.append(row_cont('  Birth year', vac['birth_year'], ctl['birth_year']))
rows.append(row_cont('  Index year', vac['index_year'], ctl['index_year']))
rows.append(section_header('Anthropometry'))
rows.append(row_cont('  BMI, kg/m² (closest to index)',
                      vac['closest_bmi'], ctl['closest_bmi']))
rows.append(row_bin('  BMI missing',
                     vac['closest_bmi'].isna().sum(), n_v,
                     ctl['closest_bmi'].isna().sum(), n_c))
rows.append(section_header('Surgery'))
rows.append(row_bin('  Hysterectomy', vac['hysterectomy'].sum(), n_v,
                                       ctl['hysterectomy'].sum(), n_c))
rows.append(row_bin('  Conization',  vac['conization'].sum(), n_v,
                                      ctl['conization'].sum(), n_c))
rows.append(row_cont('  Age at surgery, years', vac['수술시나이'], ctl['수술시나이']))
rows.append(row_cont('  Surgery year', vac['수술연도'], ctl['수술연도']))
rows.append(row_cont('  Days surgery → index', vac['수술_접종_간격일'],
                                                  ctl['수술_접종_간격일']))
rows.append(section_header('Follow-up'))
rows.append(row_cont('  Follow-up days', vac['follow_up_days'], ctl['follow_up_days']))
rows.append(row_bin('  Death during follow-up',
                     (vac['사망여부']=='Y').sum(), n_v,
                     (ctl['사망여부']=='Y').sum(), n_c))

# ---------- Update the unified Table 1 CSV ----------
print('Reading existing Table 1...')
csv_path = 'Data/Table1_BaselineCharacteristics_unified.csv'
existing = pd.read_csv(csv_path, encoding='utf-8-sig')
existing = existing[existing['block'] != 'CohortB_clearance']  # drop any prior version

block_label = 'CohortB_clearance'
new_rows = pd.DataFrame(
    [[block_label] + r for r in rows],
    columns=existing.columns)
out = pd.concat([existing, new_rows], ignore_index=True)
out.to_csv(csv_path, index=False, encoding='utf-8-sig')
print(f'Updated: {csv_path}')

# ---------- Rebuild the docx with the new block appended ----------
print('Rebuilding docx...')
docx_path = 'Data/Table1_BaselineCharacteristics_unified.docx'
doc = Document()
doc.sections[0].top_margin = doc.sections[0].bottom_margin = Pt(20*1.5)
title = doc.add_paragraph()
tr = title.add_run('Table 1. Baseline characteristics')
tr.bold = True; tr.font.size = Pt(11)

block_titles = {
    'CohortA_pre':  'Cohort A — pre-matching (full source population, pseudo-index for unvaccinated)',
    'CohortA_post': 'Cohort A — post-matching (1:1 propensity-score matched)',
    'CohortB_pre':  'Cohort B — pre-matching (all surgical patients, pseudo-index for unvaccinated)',
    'CohortB_post': 'Cohort B — post-matching (variable-ratio 1:up-to-4 fine-matched)',
    'CohortB_clearance':
        'Cohort B clearance subset — pre-vaccine hr-HPV+ analytic population for the clearance co-primary outcome (n = 292)',
}
for blk_key, blk_title in block_titles.items():
    blk = out[out['block'] == blk_key].copy()
    if blk.empty: continue
    h = doc.add_paragraph()
    hr = h.add_run(blk_title)
    hr.bold = True; hr.font.size = Pt(10)
    cols = ['Variable','Vaccinated','Non-vaccinated','p-value','|SMD|']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Light Grid Accent 1'
    for i, c in enumerate(cols):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(c); run.bold = True; run.font.size = Pt(9)
    for _, r in blk.iterrows():
        row_cells = table.add_row().cells
        for j, key in enumerate(['variable','vaccinated','non_vaccinated','p_value','abs_SMD']):
            run = row_cells[j].paragraphs[0].add_run(str(r[key]) if pd.notna(r[key]) else '')
            run.font.size = Pt(9)
    doc.add_paragraph('')

doc.save(docx_path)
print(f'Updated: {docx_path}')
