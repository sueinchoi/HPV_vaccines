"""
Table 1. Unified baseline characteristics
- Pre-matching vs Post-matching에 동일한 변수 행 사용
- Cohort A (전체) / Cohort B (수술 환자) 각각에 대해 4개 sub-table
- Pre-matching 비접종군은 pseudo index_date 부여 (재현성: seed=42)
  · Cohort A: 접종군의 first_vaccine_date 분포에서 random sample
  · Cohort B: 비접종군의 수술일 + 매칭된 접종군의 (수술-접종 간격) 분포에서 random sample

변수:
  Demographics: Age at index, Birth year, Index year, Region(서울)
  Vital signs/anthropometry (closest to index, ±1y): BMI, SBP, DBP
  Smoking (latest before index): Never/Former/Current/Unknown
  Comorbidities (pre-index): 5개 분류 + Composite
  Surgery (B only): Surgery type, Age at surgery, Surgery year, Time surgery→index
  Follow-up: Follow-up days, Mortality during follow-up
"""
import pandas as pd
import numpy as np
from scipy.stats import fisher_exact, ttest_ind
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl

RANDOM_SEED = 42
rng = np.random.default_rng(RANDOM_SEED)

CLASS_LABELS = {'1':'협심증/심근경색', '2':'고혈압', '3':'당뇨', '4':'뇌출혈/뇌경색', '5':'폐색전증'}
SMOKE_MAP = {'비흡연':'Never', '과거흡연':'Former', '현재흡연':'Current', '확인불능':'Unknown', None:'Unknown'}

# ------------------------------ helpers ------------------------------
def smd_cont(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a) < 2 or len(b) < 2: return np.nan
    pooled = np.sqrt((a.var(ddof=1) + b.var(ddof=1)) / 2)
    return (a.mean() - b.mean()) / pooled if pooled > 0 else np.nan

def smd_bin(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a) == 0 or len(b) == 0: return np.nan
    p1, p2 = a.mean(), b.mean()
    pooled = np.sqrt((p1*(1-p1) + p2*(1-p2)) / 2)
    return (p1 - p2) / pooled if pooled > 0 else np.nan

def fmt_cont(s):
    s = pd.Series(s).dropna()
    if len(s) == 0: return '-'
    return f'{s.mean():.2f} ± {s.std():.2f}'

def fmt_pct(n_yes, n_total):
    return f'{int(n_yes)} ({100*n_yes/n_total:.1f}%)' if n_total > 0 else '-'

def p_cont(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a) < 2 or len(b) < 2: return np.nan
    return ttest_ind(a, b, equal_var=False).pvalue

def p_bin(a_yes, a_n, b_yes, b_n):
    if a_n == 0 or b_n == 0: return np.nan
    return fisher_exact([[a_yes, a_n - a_yes], [b_yes, b_n - b_yes]])[1]

def fmt_p(p):
    if pd.isna(p): return '-'
    if p < 0.001: return '<0.001'
    return f'{p:.3f}'

def fmt_smd(s):
    if pd.isna(s): return '-'
    return f'{abs(s):.3f}'

# ------------------------------ load data ------------------------------
print('Loading...')
cohort = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_코호트.csv', encoding='cp949', low_memory=False)
cohort['birth_date'] = pd.to_datetime(cohort['생년월'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['death_date'] = pd.to_datetime(cohort['사망일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['last_follow'] = pd.to_datetime(cohort['최종추적일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['region_top'] = cohort['주소'].astype(str).str.split().str[0]
cohort['is_seoul'] = (cohort['region_top'] == '서울').astype(int)

# Vaccinated
rx = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_처방정보.csv', encoding='cp949', low_memory=False)
mask = (rx['처방명'].astype(str).str.contains('Gardasil|Cervarix|HPV vaccine', case=False, na=False) |
        rx['처방한글명'].astype(str).str.contains('가다실|서바릭스', na=False))
rx_vac = rx[mask].copy()
rx_vac['처방일자'] = pd.to_datetime(rx_vac['처방일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
first_vac = rx_vac.groupby('연구번호')['처방일자'].min().reset_index()
first_vac.columns = ['연구번호','first_vaccine_date']
cohort = cohort.merge(first_vac, on='연구번호', how='left')

# Surgery
surg = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_수술처방_수술종류구분완료.csv', encoding='cp949', low_memory=False)
surg['수술처방일자'] = pd.to_datetime(surg['수술처방일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
surg = surg[surg['수술 종류'].astype(str).isin(['1','3'])].copy()
SURG_TYPE = {'1':'원추절제술', '3':'자궁절제술'}
surg['수술방법'] = surg['수술 종류'].astype(str).map(SURG_TYPE)
first_surg = surg.sort_values('수술처방일자').groupby('연구번호').first()[['수술처방일자','수술방법']].reset_index()
first_surg.columns = ['연구번호','first_surg_date','first_surg_type']

# Comorbidities
print('Comorbidities...')
wb = openpyxl.load_workbook(
    'Data/한국 HPV 코호트 자료를 이용한 자_진단정보_기저질환추가_unlocked.xlsx',
    read_only=True, data_only=True)
ws = wb.active
recs = []
for row in ws.iter_rows(min_row=2, values_only=True):
    pid, cls, dd = row[0], row[5], row[8]
    if cls is None or str(cls).strip() == '': continue
    cls = str(cls).strip()
    if cls not in CLASS_LABELS: continue
    d = pd.to_datetime(str(dd), format='%Y%m%d', errors='coerce')
    recs.append((pid, cls, d))
como = pd.DataFrame(recs, columns=['pid','class','diag_date'])
first_diag = como.groupby(['pid','class'])['diag_date'].min().unstack('class')
for c in CLASS_LABELS:
    if c not in first_diag.columns:
        first_diag[c] = pd.NaT

# Clinical info (height/weight/BP/smoking) — closest to index_date
print('Clinical info...')
ci = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_기초임상정보.csv', encoding='cp949', low_memory=False)
ci['기록일자_str'] = ci['기록일자'].astype(str).str.strip()
ci['기록일자_dt'] = pd.to_datetime(ci['기록일자_str'], format='%Y%m%d', errors='coerce')

def closest_clinical_vec(query_df, ci, value_col, window_days=365):
    """Vectorized closest-by-date merge.
    query_df: ['pid','index_date'] — must be sorted by index_date
    Returns Series indexed like query_df with closest value within ±window_days, else NaN."""
    ci_v = ci[['연구번호','기록일자_dt', value_col]].dropna(subset=[value_col,'기록일자_dt']).copy()
    ci_v = ci_v.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    # forward (next record after index)
    fw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
                      left_on='index_date', right_on='rec_date', by='pid',
                      direction='forward', tolerance=pd.Timedelta(days=window_days))
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
                      left_on='index_date', right_on='rec_date', by='pid',
                      direction='backward', tolerance=pd.Timedelta(days=window_days))
    fw['_diff'] = (fw['rec_date'] - fw['index_date']).abs()
    bw['_diff'] = (bw['rec_date'] - bw['index_date']).abs()
    # choose closer
    use_fw = (fw['_diff'].fillna(pd.Timedelta(days=window_days*10)) <=
              bw['_diff'].fillna(pd.Timedelta(days=window_days*10)))
    chosen = pd.Series(np.where(use_fw, fw[value_col].values, bw[value_col].values),
                      index=fw['orig_idx'].values)
    return chosen.reindex(query_df.index).astype(float)

def latest_smoke_before_vec(query_df, ci):
    smk = ci[['연구번호','기록일자_dt','흡연여부']].dropna(subset=['흡연여부','기록일자_dt']).copy()
    smk = smk.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], smk,
                      left_on='index_date', right_on='rec_date', by='pid',
                      direction='backward')
    res = pd.Series(bw['흡연여부'].map(SMOKE_MAP).fillna('Unknown').values,
                   index=bw['orig_idx'].values)
    return res.reindex(query_df.index).fillna('Unknown')

# ------------------------------ build per-patient frame with index_date ------------------------------
def build_frame(pids_with_index):
    """pids_with_index: DataFrame with columns ['pid','index_date','vaccinated', ...optional]"""
    df = pids_with_index.copy()
    df['index_date'] = pd.to_datetime(df['index_date'])
    df = df.merge(cohort[['연구번호','birth_date','death_date','last_follow','is_seoul']],
                  left_on='pid', right_on='연구번호', how='left')
    df['age_at_index'] = (df['index_date'] - df['birth_date']).dt.days/365.25
    df['birth_year'] = df['birth_date'].dt.year
    df['index_year'] = df['index_date'].dt.year
    df['follow_up_days'] = (df['last_follow'] - df['index_date']).dt.days
    df['died'] = df['death_date'].notna() & (df['death_date'] >= df['index_date'])
    # comorbidities at baseline
    df = df.merge(first_diag, left_on='pid', right_index=True, how='left')
    for c in CLASS_LABELS:
        df[f'baseline_{c}'] = (df[c].notna()) & (df[c] <= df['index_date'])
    df['baseline_any'] = df[[f'baseline_{c}' for c in CLASS_LABELS]].any(axis=1)

    # Vital signs (vectorized merge_asof)
    q = df[['pid','index_date']].copy()
    df['height'] = closest_clinical_vec(q, ci, '키')
    df['weight'] = closest_clinical_vec(q, ci, '몸무게')
    df['sbp'] = closest_clinical_vec(q, ci, '수축기혈압')
    df['dbp'] = closest_clinical_vec(q, ci, '이완기혈압')
    df['bmi'] = df['weight'] / (df['height']/100)**2
    df['smoke'] = latest_smoke_before_vec(q, ci)

    # Surgery info
    df = df.merge(first_surg, left_on='pid', right_on='연구번호', how='left', suffixes=('','_surg'))
    df['age_at_surgery'] = (df['first_surg_date'] - df['birth_date']).dt.days/365.25
    df['surgery_year'] = df['first_surg_date'].dt.year
    df['days_surg_to_index'] = (df['index_date'] - df['first_surg_date']).dt.days
    return df

# ------------------------------ Cohort A ------------------------------
print('\n=== Cohort A — pre-match (full cohort, pseudo-index for controls) ===')
cohort_drop_dob = cohort.dropna(subset=['birth_date']).copy()
A_vac_pre = cohort_drop_dob[cohort_drop_dob['first_vaccine_date'].notna()][['연구번호','first_vaccine_date']].copy()
A_vac_pre.columns = ['pid','index_date']
A_vac_pre['vaccinated'] = True

A_ctl_pre = cohort_drop_dob[cohort_drop_dob['first_vaccine_date'].isna()][['연구번호']].copy()
A_ctl_pre.columns = ['pid']
vac_dates = A_vac_pre['index_date'].dropna().values
A_ctl_pre['index_date'] = pd.to_datetime(rng.choice(vac_dates, size=len(A_ctl_pre)))
A_ctl_pre['vaccinated'] = False

A_pre_full = pd.concat([A_vac_pre, A_ctl_pre], ignore_index=True)
A_pre = build_frame(A_pre_full)
print(f'  pre: vac={A_pre["vaccinated"].sum()}, ctl={(~A_pre["vaccinated"]).sum()}')

print('=== Cohort A — post-match (1:1 propensity-score-matched) ===')
# Rebuild the 1:1 PSM matched cohort with strict eligibility (≥1 day follow-up applied
# BEFORE matching), mirroring scripts/rebuild_table2.py — so Table 1 (b) numbers align
# exactly with Table 2.
from sklearn.linear_model import LogisticRegression as _LR
from sklearn.preprocessing import StandardScaler as _SS
_rng = np.random.default_rng(RANDOM_SEED)
_df = cohort_drop_dob.rename(columns={'연구번호':'pid'}).copy().reset_index(drop=True)
_df['vaccinated'] = _df['first_vaccine_date'].notna()
_vac_dates = _df.loc[_df['vaccinated'], 'first_vaccine_date'].dropna().values
_df.loc[~_df['vaccinated'], 'index_date'] = pd.to_datetime(_rng.choice(_vac_dates, size=(~_df['vaccinated']).sum()))
_df.loc[_df['vaccinated'], 'index_date'] = _df.loc[_df['vaccinated'], 'first_vaccine_date']
_df['age_at_index'] = (_df['index_date'] - _df['birth_date']).dt.days/365.25
_df = _df[(_df['death_date'].isna()) | (_df['death_date'] > _df['index_date'])]
_df = _df[_df['last_follow'] > _df['index_date']].reset_index(drop=True)  # strict ≥1 day FU
_q = _df[['pid','index_date']].copy()
_df['height'] = closest_clinical_vec(_q, ci, '키')
_df['weight'] = closest_clinical_vec(_q, ci, '몸무게')
_df['sbp']    = closest_clinical_vec(_q, ci, '수축기혈압')
_df['dbp']    = closest_clinical_vec(_q, ci, '이완기혈압')
_df['bmi']    = _df['weight']/(_df['height']/100)**2
_df['smoke']  = latest_smoke_before_vec(_q, ci).values
for _c in ['bmi','sbp','dbp']:
    _df[f'{_c}_miss'] = _df[_c].isna().astype(int)
    _df[_c] = _df[_c].fillna(_df[_c].mean())
_sm = pd.get_dummies(_df['smoke'], prefix='smoke').astype(int)
_df = pd.concat([_df, _sm], axis=1)
_psf = ['age_at_index','bmi','bmi_miss','sbp','sbp_miss','dbp','dbp_miss','is_seoul',
       'smoke_Never','smoke_Former','smoke_Current']
_psf = [c for c in _psf if c in _df.columns]
_X = _df[_psf].astype(float).values
_y = _df['vaccinated'].astype(int).values
_Xs = _SS().fit_transform(_X)
_lr = _LR(max_iter=2000, C=1e6, solver='lbfgs').fit(_Xs, _y)
_df['ps'] = _lr.predict_proba(_Xs)[:,1]
_df['logit_ps'] = np.log(_df['ps']/(1-_df['ps']))
_caliper = 0.2 * _df['logit_ps'].std()
_vac_idx = _df.index[_df['vaccinated']].tolist()
_ctl_idx = np.array(_df.index[~_df['vaccinated']].tolist())
_ctl_logit = _df.loc[_ctl_idx,'logit_ps'].values
_order = np.argsort(_ctl_logit)
_ctl_sorted = _ctl_idx[_order]; _ctl_logit_sorted = _ctl_logit[_order]
_used = np.zeros(len(_ctl_sorted), dtype=bool)
_matched = []
_vo = np.array(_vac_idx); _rng2 = np.random.default_rng(RANDOM_SEED); _rng2.shuffle(_vo)
for _vi in _vo:
    _t = _df.loc[_vi,'logit_ps']
    _lo = np.searchsorted(_ctl_logit_sorted, _t-_caliper)
    _hi = np.searchsorted(_ctl_logit_sorted, _t+_caliper, side='right')
    _bj, _bd = -1, _caliper+1
    for _j in range(_lo, _hi):
        if _used[_j]: continue
        _d = abs(_ctl_logit_sorted[_j]-_t)
        if _d < _bd: _bd=_d; _bj=_j
    if _bj>=0:
        _used[_bj] = True
        _matched.append((_vi, _ctl_sorted[_bj]))
_pair_idx = []
for _pid_, (_vi, _cii) in enumerate(_matched):
    _pair_idx.extend([_vi, _cii])
_psm = _df.loc[_pair_idx].copy().reset_index(drop=True)
A_post_in = _psm[['pid','index_date','vaccinated']].copy()
A_post = build_frame(A_post_in)
print(f'  post (1:1 PSM, strict ≥1 day FU): vac={A_post["vaccinated"].sum()}, ctl={(~A_post["vaccinated"]).sum()}')

# ------------------------------ Cohort B ------------------------------
print('\n=== Cohort B — pre-match (all surgery patients, pseudo-index for controls) ===')
sg = first_surg.merge(cohort_drop_dob[['연구번호','first_vaccine_date','birth_date']], on='연구번호', how='left')
sg = sg.dropna(subset=['birth_date'])
B_vac_pre = sg[sg['first_vaccine_date'].notna()].copy()
B_vac_pre['index_date'] = B_vac_pre['first_vaccine_date']
B_vac_pre['vaccinated'] = True

# 수술-접종 간격 분포
intervals = (B_vac_pre['first_vaccine_date'] - B_vac_pre['first_surg_date']).dt.days.dropna().values
intervals = intervals[intervals >= 0]  # only positive intervals (vaccine after surgery)
print(f'  vaccinated surgery patients: {len(B_vac_pre)}, valid intervals: {len(intervals)}')

B_ctl_pre = sg[sg['first_vaccine_date'].isna()].copy()
sampled = rng.choice(intervals, size=len(B_ctl_pre))
B_ctl_pre['index_date'] = B_ctl_pre['first_surg_date'] + pd.to_timedelta(sampled, unit='D')
B_ctl_pre['vaccinated'] = False

B_pre_in = pd.concat([B_vac_pre[['연구번호','index_date','vaccinated']],
                     B_ctl_pre[['연구번호','index_date','vaccinated']]], ignore_index=True)
B_pre_in.columns = ['pid','index_date','vaccinated']
B_pre = build_frame(B_pre_in)
print(f'  pre: vac={B_pre["vaccinated"].sum()}, ctl={(~B_pre["vaccinated"]).sum()}')

print('=== Cohort B — post-match (fine-matched 1:4) ===')
B_post_in = pd.read_csv('Data/final_matched_cohort.csv', encoding='utf-8-sig')
B_post_in = B_post_in[['연구번호','index_date','접종여부']].rename(
    columns={'연구번호':'pid','접종여부':'vaccinated'})
B_post_in['vaccinated'] = B_post_in['vaccinated'].astype(bool)
B_post = build_frame(B_post_in)
print(f'  post: vac={B_post["vaccinated"].sum()}, ctl={(~B_post["vaccinated"]).sum()}')

# ------------------------------ Variable definitions (uniform) ------------------------------
# (label, kind, col, [extra]); kind in {'cont','bin','cat_pct'}
def variable_rows(df, vac_col='vaccinated', include_surgery=False):
    vac = df[df[vac_col]==True]; ctl = df[df[vac_col]==False]
    nv, nc = len(vac), len(ctl)

    rows = []

    def add_cont(label, col):
        rows.append((label,
                     fmt_cont(vac[col]),
                     fmt_cont(ctl[col]),
                     fmt_p(p_cont(vac[col], ctl[col])),
                     fmt_smd(smd_cont(vac[col], ctl[col]))))

    def add_bin_yn(label, col):
        v_yes = int(vac[col].fillna(0).astype(bool).sum())
        c_yes = int(ctl[col].fillna(0).astype(bool).sum())
        rows.append((label,
                     fmt_pct(v_yes, nv),
                     fmt_pct(c_yes, nc),
                     fmt_p(p_bin(v_yes, nv, c_yes, nc)),
                     fmt_smd(smd_bin(vac[col].astype(float), ctl[col].astype(float)))))

    def add_cat_levels(label, col, levels):
        # header row
        rows.append((label, '', '', '', ''))
        for lv in levels:
            v_yes = int((vac[col]==lv).sum())
            c_yes = int((ctl[col]==lv).sum())
            rows.append((f'  {lv}',
                         fmt_pct(v_yes, nv),
                         fmt_pct(c_yes, nc),
                         fmt_p(p_bin(v_yes, nv, c_yes, nc)),
                         fmt_smd(smd_bin((vac[col]==lv).astype(float), (ctl[col]==lv).astype(float)))))

    rows.append(('Demographics', '', '', '', ''))
    add_cont('  Age at index, years', 'age_at_index')
    add_cont('  Birth year', 'birth_year')
    add_cont('  Index year', 'index_year')
    add_bin_yn('  Residence in Seoul', 'is_seoul')

    rows.append(('Anthropometry & Vital signs (closest to index, ±1 yr)', '', '', '', ''))
    add_cont('  Height, cm', 'height')
    add_cont('  Weight, kg', 'weight')
    add_cont('  BMI, kg/m²', 'bmi')
    add_cont('  Systolic BP, mmHg', 'sbp')
    add_cont('  Diastolic BP, mmHg', 'dbp')

    add_cat_levels('Smoking status', 'smoke', ['Never','Former','Current','Unknown'])

    rows.append(('Pre-existing comorbidities (before index)', '', '', '', ''))
    for c, label in CLASS_LABELS.items():
        add_bin_yn(f'  {label}', f'baseline_{c}')
    add_bin_yn('  Composite (any of 5)', 'baseline_any')

    if include_surgery:
        add_cat_levels('Surgery type', 'first_surg_type', ['원추절제술','자궁절제술'])
        add_cont('  Age at surgery, years', 'age_at_surgery')
        add_cont('  Surgery year', 'surgery_year')
        add_cont('  Surgery → index, days', 'days_surg_to_index')

    rows.append(('Follow-up', '', '', '', ''))
    add_cont('  Follow-up, days', 'follow_up_days')
    add_bin_yn('  Mortality during follow-up', 'died')

    return rows, nv, nc

print('\nBuilding rows...')
A_pre_rows, A_pre_nv, A_pre_nc = variable_rows(A_pre, include_surgery=False)
A_post_rows, A_post_nv, A_post_nc = variable_rows(A_post, include_surgery=False)
B_pre_rows, B_pre_nv, B_pre_nc = variable_rows(B_pre, include_surgery=True)
B_post_rows, B_post_nv, B_post_nc = variable_rows(B_post, include_surgery=True)

# ------------------------------ docx ------------------------------
print('Writing docx...')
doc = Document()
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'
sty.font.size = Pt(10)

doc.add_heading('Table 1. Baseline Characteristics of Study Cohorts', level=0)

def add_table(title, header_n_v, header_n_c, rows):
    p = doc.add_paragraph()
    r = p.add_run(title); r.bold = True; r.font.size = Pt(11)
    table = doc.add_table(rows=2 + len(rows), cols=5)
    table.style = 'Light Grid Accent 1'
    hdr0 = table.rows[0].cells
    hdr0[0].text = 'Characteristic'
    hdr0[1].text = 'Vaccinated'
    hdr0[2].text = 'Non-vaccinated'
    hdr0[3].text = 'p-value'
    hdr0[4].text = '|SMD|'
    hdr1 = table.rows[1].cells
    hdr1[0].text = ''
    hdr1[1].text = f'(n={header_n_v})'
    hdr1[2].text = f'(n={header_n_c})'
    hdr1[3].text = ''
    hdr1[4].text = ''
    for c in list(hdr0) + list(hdr1):
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True; run.font.size = Pt(9)
    for i, row_data in enumerate(rows):
        cells = table.rows[i+2].cells
        # Section header rows are tuples where col 1-4 are empty
        is_section = all(x == '' for x in row_data[1:])
        for j, val in enumerate(row_data):
            cells[j].text = val
            for para in cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if is_section:
                        run.bold = True
                if j > 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

doc.add_heading('Cohort A — Whole-cohort comorbidity analysis', level=1)
add_table(f'(a) Pre-matching', A_pre_nv, A_pre_nc, A_pre_rows)
add_table(f'(b) Post-matching (1:1 propensity-score-matched)', A_post_nv, A_post_nc, A_post_rows)

doc.add_heading('Cohort B — Surgical efficacy analysis', level=1)
add_table(f'(c) Pre-matching (all surgery patients)', B_pre_nv, B_pre_nc, B_pre_rows)
add_table(f'(d) Post-matching (fine-matched 1:4)', B_post_nv, B_post_nc, B_post_rows)

# Footnote
foot = doc.add_paragraph()
foot.add_run('Footnotes. ').bold = True
foot_text = (
    'Continuous variables: mean ± SD (independent-sample Welch\'s t-test); '
    'categorical/binary: n (%) (Fisher\'s exact test). '
    '|SMD| = absolute standardized mean difference; <0.10 indicates good balance. '
    'Index date: vaccinated = first HPV vaccine prescription; non-vaccinated post-matching = matched index date. '
    'Pre-matching non-vaccinated controls were assigned a pseudo index date for symmetric variable evaluation: '
    'Cohort A — random sample from the vaccinated patients\' first-vaccine-date distribution (seed=42); '
    'Cohort B — surgery date plus a random sample from the vaccinated patients\' surgery-to-vaccine interval distribution (seed=42). '
    'BMI/SBP/DBP are taken from the clinical record closest to the index date within ±365 days; '
    'smoking status is the latest record before index, with no available record classified as "Unknown". '
    'Pre-existing comorbidities are defined as a diagnosis prior to or on the index date. '
    'Surgery type and age at surgery (Cohort B only) are defined at the first cervical surgery (conization or hysterectomy).'
)
foot.add_run(foot_text).font.size = Pt(8)

out = 'Data/Table1_BaselineCharacteristics_unified.docx'
doc.save(out)
print(f'Saved: {out}')

# CSV
all_rows = []
for label, rows in [('CohortA_pre', A_pre_rows), ('CohortA_post', A_post_rows),
                    ('CohortB_pre', B_pre_rows), ('CohortB_post', B_post_rows)]:
    for r in rows:
        all_rows.append({'block': label, 'variable': r[0], 'vaccinated': r[1],
                        'non_vaccinated': r[2], 'p_value': r[3], 'abs_SMD': r[4]})
pd.DataFrame(all_rows).to_csv('Data/Table1_BaselineCharacteristics_unified.csv', index=False, encoding='utf-8-sig')
print('Saved: Data/Table1_BaselineCharacteristics_unified.csv')
