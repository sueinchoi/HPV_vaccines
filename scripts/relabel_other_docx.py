"""
Relabel reinfection -> clearance / post-index detection (sensitivity) in:
  - docs/HPV_manuscript_csi.docx
  - docs/HPV_supplementary_csi.docx
  - Data/SupTableS5_number_at_risk.docx
  - Data/Methodology_Revisions.docx

Same approach as relabel_supplementary_clearance.py + relabel_supplementary_cells.py.
"""
from docx import Document
import re

FILES = [
    'docs/HPV_manuscript_csi.docx',
    'docs/HPV_supplementary_csi.docx',
    'Data/SupTableS5_number_at_risk.docx',
    'Data/Methodology_Revisions.docx',
]

# Replacement targets — applied in order.
# Long-phrase substitutions first (to win over short ones).
PARAGRAPH_REPLACEMENTS = [
    # paragraph-level descriptive phrases
    ('Lesion recurrence and high-risk HPV reinfection address distinct biological questions—reactivation or persistence of pre-existing transformed clones versus de novo viral acquisition',
     'Lesion recurrence and hr-HPV clearance address distinct biological questions—post-surgical regrowth of high-grade cervical neoplasia versus persistence or clearance of the pre-vaccine high-risk infection that drives it'),
    ('Lesion recurrence and high-risk HPV reinfection',
     'Lesion recurrence and hr-HPV clearance'),
    ('lesion recurrence and high-risk HPV reinfection',
     'lesion recurrence and hr-HPV clearance'),
    ('high-risk HPV reinfection', 'hr-HPV clearance'),
    ('HPV reinfection', 'hr-HPV clearance'),
    ('Gardasil-specific signal for hr-HPV clearance',
     'Gardasil-quadrivalent signal in the supplementary post-index hr-HPV detection sensitivity (not the co-primary clearance outcome)'),
]

# Cell-level label replacements (exact-match preferred for tables)
CELL_EXACT = {
    'HPV reinfection': 'Post-index hr-HPV detection (sensitivity)',
    'hr-HPV clearance': 'hr-HPV clearance',  # no-op (already correct)
}

def relabel_paragraph(p):
    txt = p.text
    new_txt = txt
    for old, new in PARAGRAPH_REPLACEMENTS:
        if old in new_txt:
            new_txt = new_txt.replace(old, new)
    if new_txt != txt:
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = new_txt
        else:
            p.add_run(new_txt)
        return True
    return False

def relabel_cell(c):
    s = c.text.strip()
    # tables: only relabel cells whose ENTIRE text is the label (avoid corrupting prose cells)
    if s in CELL_EXACT and CELL_EXACT[s] != s:
        for para in c.paragraphs:
            for run in list(para.runs):
                run.text = ''
        c.paragraphs[0].text = CELL_EXACT[s]
        return True
    return False

for f in FILES:
    doc = Document(f)
    n_par = sum(1 for p in doc.paragraphs if relabel_paragraph(p))
    n_cell = 0
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if relabel_cell(c):
                    n_cell += 1
    doc.save(f)
    print(f'{f}: paragraphs={n_par}, cells={n_cell}')
