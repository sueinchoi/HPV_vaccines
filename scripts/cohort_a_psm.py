"""
Cohort A — Propensity Score Matching
====================================

매칭 변수 (PS model에 포함):
  · Age at index (continuous)
  · BMI, SBP, DBP (continuous; mean imputation + missing indicator)
  · Smoking status (categorical with Unknown as a separate level)
  · Residence in Seoul (binary)

방법:
  1. Pseudo-index 부여 (비접종군: 접종군 vaccine_date 분포에서 random sample, seed=42)
  2. 각 환자에 대해 index_date ±365일 내 가장 가까운 BMI/SBP/DBP, latest 흡연 추출
  3. Logistic regression PS 추정 (statsmodels-style; sklearn LR with no penalty)
  4. 1:1 nearest neighbor matching on logit(PS), caliper = 0.2 × SD(logit PS), no replacement (Austin 2011)
  5. Pre/Post SMD, Love plot, 5개 기저질환 비교 결과 docx로 저장
"""
import pandas as pd
import numpy as np
import openpyxl
from scipy.stats import fisher_exact, ttest_ind
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.family'] = ['DejaVu Sans', 'AppleGothic']
plt.rcParams['axes.unicode_minus'] = False

RANDOM_SEED = 42
rng = np.random.default_rng(RANDOM_SEED)
np.random.seed(RANDOM_SEED)
CLASS_LABELS = {'1':'협심증/심근경색', '2':'고혈압', '3':'당뇨', '4':'뇌출혈/뇌경색', '5':'폐색전증'}
SMOKE_MAP = {'비흡연':'Never', '과거흡연':'Former', '현재흡연':'Current', '확인불능':'Unknown'}

# --------------------------- helpers ---------------------------
def smd_cont(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a)<2 or len(b)<2: return np.nan
    pooled = np.sqrt((a.var(ddof=1)+b.var(ddof=1))/2)
    return (a.mean()-b.mean())/pooled if pooled>0 else np.nan

def smd_bin(a, b):
    a = pd.Series(a).dropna(); b = pd.Series(b).dropna()
    if len(a)==0 or len(b)==0: return np.nan
    p1, p2 = a.mean(), b.mean()
    pooled = np.sqrt((p1*(1-p1)+p2*(1-p2))/2)
    return (p1-p2)/pooled if pooled>0 else np.nan

def closest_vec(query_df, ci, value_col, window_days=365):
    ci_v = ci[['연구번호','기록일자_dt', value_col]].dropna(subset=[value_col,'기록일자_dt']).copy()
    ci_v = ci_v.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    fw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid',
        direction='forward', tolerance=pd.Timedelta(days=window_days))
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], ci_v,
        left_on='index_date', right_on='rec_date', by='pid',
        direction='backward', tolerance=pd.Timedelta(days=window_days))
    fw['_d'] = (fw['rec_date']-fw['index_date']).abs()
    bw['_d'] = (bw['rec_date']-bw['index_date']).abs()
    use_fw = (fw['_d'].fillna(pd.Timedelta(days=window_days*10)) <=
              bw['_d'].fillna(pd.Timedelta(days=window_days*10)))
    chosen = pd.Series(np.where(use_fw, fw[value_col].values, bw[value_col].values),
                      index=fw['orig_idx'].values)
    return chosen.reindex(query_df.index).astype(float)

def smoke_vec(query_df, ci):
    smk = ci[['연구번호','기록일자_dt','흡연여부']].dropna(subset=['흡연여부','기록일자_dt']).copy()
    smk = smk.sort_values('기록일자_dt').rename(columns={'연구번호':'pid','기록일자_dt':'rec_date'})
    q = query_df.sort_values('index_date').reset_index().rename(columns={'index':'orig_idx'})
    bw = pd.merge_asof(q[['orig_idx','pid','index_date']], smk,
        left_on='index_date', right_on='rec_date', by='pid', direction='backward')
    res = pd.Series(bw['흡연여부'].map(SMOKE_MAP).fillna('Unknown').values,
                   index=bw['orig_idx'].values)
    return res.reindex(query_df.index).fillna('Unknown')

# --------------------------- load source ---------------------------
print('[1/7] Loading source data...')
cohort = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_코호트.csv', encoding='cp949', low_memory=False)
cohort['birth_date'] = pd.to_datetime(cohort['생년월'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['death_date'] = pd.to_datetime(cohort['사망일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['last_follow'] = pd.to_datetime(cohort['최종추적일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
cohort['is_seoul'] = (cohort['주소'].astype(str).str.split().str[0]=='서울').astype(int)

rx = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_처방정보.csv', encoding='cp949', low_memory=False)
mask = (rx['처방명'].astype(str).str.contains('Gardasil|Cervarix|HPV vaccine', case=False, na=False) |
        rx['처방한글명'].astype(str).str.contains('가다실|서바릭스', na=False))
rx_vac = rx[mask].copy()
rx_vac['처방일자'] = pd.to_datetime(rx_vac['처방일자'].astype('Int64').astype(str), format='%Y%m%d', errors='coerce')
first_vac = rx_vac.groupby('연구번호')['처방일자'].min().reset_index()
first_vac.columns = ['연구번호','first_vaccine_date']
cohort = cohort.merge(first_vac, on='연구번호', how='left')
cohort = cohort.dropna(subset=['birth_date'])

print('[2/7] Loading clinical info...')
ci = pd.read_csv('Data/한국 HPV 코호트 자료를 이용한 자_기초임상정보.csv', encoding='cp949', low_memory=False)
ci['기록일자_dt'] = pd.to_datetime(ci['기록일자'].astype(str).str.strip(), format='%Y%m%d', errors='coerce')

print('[3/7] Loading comorbidities...')
wb = openpyxl.load_workbook('Data/한국 HPV 코호트 자료를 이용한 자_진단정보_기저질환추가_unlocked.xlsx',
                          read_only=True, data_only=True)
ws = wb.active
recs=[]
for row in ws.iter_rows(min_row=2, values_only=True):
    pid, cls, dd = row[0], row[5], row[8]
    if cls is None or str(cls).strip()=='': continue
    cls = str(cls).strip()
    if cls not in CLASS_LABELS: continue
    d = pd.to_datetime(str(dd), format='%Y%m%d', errors='coerce')
    recs.append((pid, cls, d))
como = pd.DataFrame(recs, columns=['pid','class','diag_date'])
first_diag = como.groupby(['pid','class'])['diag_date'].min().unstack('class')
for c in CLASS_LABELS:
    if c not in first_diag.columns: first_diag[c] = pd.NaT

# --------------------------- build cohort ---------------------------
print('[4/7] Assigning index_date (pseudo for controls)...')
df = cohort.copy()
df['vaccinated'] = df['first_vaccine_date'].notna()
vac_dates = df.loc[df['vaccinated'], 'first_vaccine_date'].dropna().values
df.loc[~df['vaccinated'], 'index_date'] = pd.to_datetime(rng.choice(vac_dates, size=(~df['vaccinated']).sum()))
df.loc[df['vaccinated'], 'index_date'] = df.loc[df['vaccinated'], 'first_vaccine_date']
df = df.rename(columns={'연구번호':'pid'}).reset_index(drop=True)
df['age_at_index'] = (df['index_date'] - df['birth_date']).dt.days/365.25

# eligibility: alive at index_date, observable
df = df[df['birth_date'].notna() & df['index_date'].notna()]
df = df[(df['death_date'].isna()) | (df['death_date'] > df['index_date'])]
df = df[df['last_follow'] >= df['index_date']]
df = df.reset_index(drop=True)
print(f'  eligible: vac={df["vaccinated"].sum()}, ctl={(~df["vaccinated"]).sum()}')

print('[5/7] Extracting closest vital signs and smoking...')
q = df[['pid','index_date']].copy()
df['height'] = closest_vec(q, ci, '키')
df['weight'] = closest_vec(q, ci, '몸무게')
df['sbp'] = closest_vec(q, ci, '수축기혈압')
df['dbp'] = closest_vec(q, ci, '이완기혈압')
df['bmi'] = df['weight']/(df['height']/100)**2
df['smoke'] = smoke_vec(q, ci).values
# baseline comorbidities
df = df.merge(first_diag, left_on='pid', right_index=True, how='left')
for c in CLASS_LABELS:
    df[f'baseline_{c}'] = (df[c].notna()) & (df[c] <= df['index_date'])
df['baseline_any'] = df[[f'baseline_{c}' for c in CLASS_LABELS]].any(axis=1)
df['follow_up_days'] = (df['last_follow'] - df['index_date']).dt.days
df['died'] = df['death_date'].notna() & (df['death_date'] >= df['index_date'])

# --------------------------- PSM ---------------------------
print('[6/7] Fitting propensity score model...')
# Missing-indicator + mean imputation
for c in ['bmi','sbp','dbp']:
    df[f'{c}_miss'] = df[c].isna().astype(int)
    df[c] = df[c].fillna(df[c].mean())
# Smoking dummies (Unknown is reference)
sm = pd.get_dummies(df['smoke'], prefix='smoke').astype(int)
df = pd.concat([df, sm], axis=1)
ps_features = ['age_at_index','bmi','bmi_miss','sbp','sbp_miss','dbp','dbp_miss','is_seoul']
for col in ['smoke_Never','smoke_Former','smoke_Current']:  # Unknown 제외 = reference
    if col in df.columns: ps_features.append(col)

X = df[ps_features].astype(float).values
y = df['vaccinated'].astype(int).values
scaler = StandardScaler()
Xs = scaler.fit_transform(X)
lr = LogisticRegression(max_iter=2000, C=1e6, solver='lbfgs')
lr.fit(Xs, y)
df['ps'] = lr.predict_proba(Xs)[:,1]
df['logit_ps'] = np.log(df['ps']/(1-df['ps']))
print(f'  PS range: [{df["ps"].min():.4f}, {df["ps"].max():.4f}]; logit SD = {df["logit_ps"].std():.4f}')
print(f'  Coefficients (standardized):')
for f, c in zip(ps_features, lr.coef_[0]):
    print(f'    {f:18s} {c:+.4f}')

# 1:1 nearest neighbor matching with caliper
caliper = 0.2 * df['logit_ps'].std()
print(f'  Caliper = 0.2 × SD(logit PS) = {caliper:.4f}')

vac_idx = df.index[df['vaccinated']].tolist()
ctl_idx = df.index[~df['vaccinated']].tolist()
ctl_logit = df.loc[ctl_idx, 'logit_ps'].values
ctl_arr_idx = np.array(ctl_idx)
order = np.argsort(ctl_logit)  # sorted by logit_ps for fast nearest search
ctl_sorted = ctl_arr_idx[order]
ctl_logit_sorted = ctl_logit[order]
used = np.zeros(len(ctl_sorted), dtype=bool)
matched = []
# greedy: random order of vaccinated
vac_order = np.array(vac_idx)
rng.shuffle(vac_order)
for vi in vac_order:
    target = df.loc[vi, 'logit_ps']
    # binary search for nearest among unused
    pos = np.searchsorted(ctl_logit_sorted, target)
    # search outward for nearest unused within caliper
    best_j = -1
    best_diff = caliper + 1
    # scan a window — outward expansion
    for direction in [-1, 1, 0]:
        pass
    # simpler: scan whole range within caliper
    lo = np.searchsorted(ctl_logit_sorted, target - caliper)
    hi = np.searchsorted(ctl_logit_sorted, target + caliper, side='right')
    for j in range(lo, hi):
        if used[j]: continue
        d = abs(ctl_logit_sorted[j] - target)
        if d < best_diff:
            best_diff = d
            best_j = j
    if best_j >= 0:
        used[best_j] = True
        matched.append((vi, ctl_sorted[best_j]))

print(f'  Matched pairs: {len(matched)} / {len(vac_idx)} vaccinated')

vac_m = [m[0] for m in matched]; ctl_m = [m[1] for m in matched]
df['matched'] = False
df.loc[vac_m, 'matched'] = True
df.loc[ctl_m, 'matched'] = True

# --------------------------- balance check ---------------------------
print('[7/7] Computing SMDs and outputting...')

bal_vars_cont = [
    ('Age at index, years','age_at_index'),
    ('BMI, kg/m²','bmi'),
    ('Systolic BP, mmHg','sbp'),
    ('Diastolic BP, mmHg','dbp'),
    ('Logit(PS)','logit_ps'),
]
bal_vars_bin = [
    ('Residence in Seoul','is_seoul'),
    ('BMI missing','bmi_miss'),
    ('SBP missing','sbp_miss'),
    ('DBP missing','dbp_miss'),
    ('Smoking: Never','smoke_Never'),
    ('Smoking: Former','smoke_Former'),
    ('Smoking: Current','smoke_Current'),
    ('Smoking: Unknown','smoke_Unknown'),
]

def make_balance(df_sub):
    vac = df_sub[df_sub['vaccinated']]
    ctl = df_sub[~df_sub['vaccinated']]
    out = []
    for label, col in bal_vars_cont:
        v_mean = vac[col].mean(); v_sd = vac[col].std()
        c_mean = ctl[col].mean(); c_sd = ctl[col].std()
        s = abs(smd_cont(vac[col], ctl[col]))
        out.append((label, f'{v_mean:.2f} ± {v_sd:.2f}', f'{c_mean:.2f} ± {c_sd:.2f}', f'{s:.3f}'))
    for label, col in bal_vars_bin:
        if col not in df_sub.columns:
            out.append((label,'-','-','-')); continue
        v_p = vac[col].mean()*100; c_p = ctl[col].mean()*100
        s = abs(smd_bin(vac[col].astype(float), ctl[col].astype(float)))
        out.append((label, f'{v_p:.1f}%', f'{c_p:.1f}%', f'{s:.3f}'))
    return out, len(vac), len(ctl)

pre_rows, pre_nv, pre_nc = make_balance(df)
post_rows, post_nv, post_nc = make_balance(df[df['matched']])
print(f'  Pre: {pre_nv} vs {pre_nc}, Post: {post_nv} vs {post_nc}')

# Outcome comparison: 5 comorbidity baseline, new-onset
def outcome_table(df_sub):
    vac = df_sub[df_sub['vaccinated']]; ctl = df_sub[~df_sub['vaccinated']]
    rows = []
    rows.append(('Baseline comorbidity (before index)','','','','',''))
    for c, label in CLASS_LABELS.items():
        v_yes = int(vac[f'baseline_{c}'].sum()); c_yes = int(ctl[f'baseline_{c}'].sum())
        odds, p = fisher_exact([[v_yes, len(vac)-v_yes],[c_yes, len(ctl)-c_yes]])
        rows.append((f'  {label}',
                    f'{v_yes} ({100*v_yes/len(vac):.1f}%)',
                    f'{c_yes} ({100*c_yes/len(ctl):.1f}%)',
                    f'{odds:.2f}' if not np.isinf(odds) else '∞',
                    f'{p:.3f}' if p>=0.001 else '<0.001',
                    'Yes' if p<0.05 else ''))
    v_yes = int(vac['baseline_any'].sum()); c_yes = int(ctl['baseline_any'].sum())
    odds, p = fisher_exact([[v_yes, len(vac)-v_yes],[c_yes, len(ctl)-c_yes]])
    rows.append(('  Composite (any of 5)',
                f'{v_yes} ({100*v_yes/len(vac):.1f}%)',
                f'{c_yes} ({100*c_yes/len(ctl):.1f}%)',
                f'{odds:.2f}' if not np.isinf(odds) else '∞',
                f'{p:.3f}' if p>=0.001 else '<0.001',
                'Yes' if p<0.05 else ''))

    rows.append(('New-onset comorbidity (after index, baseline-free)','','','','',''))
    for c, label in CLASS_LABELS.items():
        v_elig = vac[~vac[f'baseline_{c}']]
        c_elig = ctl[~ctl[f'baseline_{c}']]
        v_yes = int(((v_elig[c].notna()) & (v_elig[c] > v_elig['index_date'])).sum())
        c_yes = int(((c_elig[c].notna()) & (c_elig[c] > c_elig['index_date'])).sum())
        if len(v_elig)==0 or len(c_elig)==0:
            rows.append((f'  {label}','-','-','-','-','')); continue
        odds, p = fisher_exact([[v_yes, len(v_elig)-v_yes],[c_yes, len(c_elig)-c_yes]])
        rows.append((f'  {label}',
                    f'{v_yes}/{len(v_elig)} ({100*v_yes/len(v_elig):.2f}%)',
                    f'{c_yes}/{len(c_elig)} ({100*c_yes/len(c_elig):.2f}%)',
                    f'{odds:.2f}' if not np.isinf(odds) else '∞',
                    f'{p:.3f}' if p>=0.001 else '<0.001',
                    'Yes' if p<0.05 else ''))
    return rows

post_outcomes = outcome_table(df[df['matched']])

# --------------------------- Love plot ---------------------------
print('  Creating Love plot...')
labels_all = [v[0] for v in bal_vars_cont] + [v[0] for v in bal_vars_bin]
pre_smd = []
post_smd = []
df_post = df[df['matched']]
for label, col in bal_vars_cont:
    pre_smd.append(abs(smd_cont(df[df['vaccinated']][col], df[~df['vaccinated']][col])))
    post_smd.append(abs(smd_cont(df_post[df_post['vaccinated']][col], df_post[~df_post['vaccinated']][col])))
for label, col in bal_vars_bin:
    if col not in df.columns:
        pre_smd.append(np.nan); post_smd.append(np.nan); continue
    pre_smd.append(abs(smd_bin(df[df['vaccinated']][col].astype(float), df[~df['vaccinated']][col].astype(float))))
    post_smd.append(abs(smd_bin(df_post[df_post['vaccinated']][col].astype(float), df_post[~df_post['vaccinated']][col].astype(float))))

fig, ax = plt.subplots(figsize=(8,6))
y = np.arange(len(labels_all))
ax.scatter(pre_smd, y, label='Pre-PSM', color='#9b2226', marker='o', s=60)
ax.scatter(post_smd, y, label='Post-PSM (1:1)', color='#1f6f8b', marker='s', s=60)
for i, (p_, q_) in enumerate(zip(pre_smd, post_smd)):
    if not (np.isnan(p_) or np.isnan(q_)):
        ax.plot([p_, q_], [i, i], color='gray', alpha=0.4, lw=0.8)
ax.axvline(0.1, color='black', linestyle='--', alpha=0.5, label='|SMD|=0.10')
ax.set_yticks(y); ax.set_yticklabels(labels_all)
ax.set_xlabel('|Standardized Mean Difference|')
ax.set_title('Cohort A — Love plot of covariate balance (PSM 1:1)')
ax.legend(loc='lower right')
ax.invert_yaxis()
plt.tight_layout()
plt.savefig('Data/cohort_a_psm_loveplot.png', dpi=150)
plt.close()

# Save matched cohort
df[df['matched']].to_csv('Data/cohort_a_psm_matched.csv', index=False, encoding='utf-8-sig')

# --------------------------- docx ---------------------------
doc = Document()
sty = doc.styles['Normal']; sty.font.name = 'Times New Roman'; sty.font.size = Pt(10)
doc.add_heading('Cohort A — Propensity Score Matching (1:1)', level=0)

# Methods box
mp = doc.add_paragraph()
mp.add_run('Methods. ').bold = True
mp.add_run(
    'Propensity scores were estimated by logistic regression with vaccination status as the outcome '
    'and the following covariates: age at index, BMI, SBP, DBP (with mean imputation and missing indicators), '
    'smoking status (Never/Former/Current/Unknown — Unknown as reference), and residence in Seoul. '
    'Pre-matching non-vaccinated controls were assigned a pseudo index date randomly sampled from the vaccinated '
    'group\'s vaccine-date distribution (seed=42). One-to-one nearest-neighbour matching on the logit of the '
    'propensity score without replacement was performed using a caliper of 0.2 × SD(logit PS) (Austin 2011). '
    'Covariate balance was assessed by absolute standardized mean differences (|SMD|<0.10 considered well-balanced).'
).font.size = Pt(9)

# Balance table
doc.add_heading('Table 1. Covariate balance — Pre vs Post PSM', level=1)
t = doc.add_table(rows=2+len(bal_vars_cont)+len(bal_vars_bin), cols=7)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'Variable'
hdr[1].text = 'Pre-matching'; hdr[2].text = ''
hdr[3].text = 'Pre |SMD|'
hdr[4].text = 'Post-matching'; hdr[5].text = ''
hdr[6].text = 'Post |SMD|'
sub = t.rows[1].cells
sub[0].text = ''
sub[1].text = f'Vac (n={pre_nv})'; sub[2].text = f'Ctl (n={pre_nc})'; sub[3].text = ''
sub[4].text = f'Vac (n={post_nv})'; sub[5].text = f'Ctl (n={post_nc})'; sub[6].text = ''
for c in list(hdr)+list(sub):
    for para in c.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold = True; r.font.size = Pt(9)

for i, (pre, post) in enumerate(zip(pre_rows, post_rows)):
    cells = t.rows[i+2].cells
    cells[0].text = pre[0]
    cells[1].text = pre[1]; cells[2].text = pre[2]; cells[3].text = pre[3]
    cells[4].text = post[1]; cells[5].text = post[2]; cells[6].text = post[3]
    for j, c in enumerate(cells):
        for para in c.paragraphs:
            for r in para.runs: r.font.size = Pt(9)
            if j>0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Outcome table
doc.add_heading('Table 2. Comorbidity comparison after PSM', level=1)
t2 = doc.add_table(rows=1+len(post_outcomes), cols=6)
t2.style = 'Light Grid Accent 1'
h = t2.rows[0].cells
h[0].text = 'Comorbidity'
h[1].text = f'Vaccinated (n={post_nv})'
h[2].text = f'Non-vaccinated (n={post_nc})'
h[3].text = 'Odds ratio'
h[4].text = 'p-value'
h[5].text = 'Significant (α=0.05)'
for c in h:
    for para in c.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs: r.bold = True; r.font.size = Pt(9)
for i, row in enumerate(post_outcomes):
    cells = t2.rows[i+1].cells
    is_section = all(x=='' for x in row[1:])
    for j, val in enumerate(row):
        cells[j].text = val
        for para in cells[j].paragraphs:
            for r in para.runs:
                r.font.size = Pt(9)
                if is_section: r.bold = True
            if j>0: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Love plot
doc.add_heading('Figure 1. Love plot', level=1)
doc.add_picture('Data/cohort_a_psm_loveplot.png', width=Inches(6.0))

doc.add_paragraph().add_run(
    'Footnote: Continuous variables shown as mean ± SD; binary variables as %. '
    '|SMD|: absolute standardized mean difference. Caliper = 0.2 × SD(logit PS). '
    'Logit(PS) is shown as a sanity check that the propensity score itself is balanced after matching.'
).font.size = Pt(8)

out = 'Data/CohortA_PSM_report.docx'
doc.save(out)
print(f'\nSaved:')
print(f'  {out}')
print(f'  Data/cohort_a_psm_loveplot.png')
print(f'  Data/cohort_a_psm_matched.csv (n={len(df[df["matched"]])})')

# Export balance + outcome CSV
pd.DataFrame({'variable':[r[0] for r in pre_rows],
              'vac_pre':[r[1] for r in pre_rows],'ctl_pre':[r[2] for r in pre_rows],'smd_pre':[r[3] for r in pre_rows],
              'vac_post':[r[1] for r in post_rows],'ctl_post':[r[2] for r in post_rows],'smd_post':[r[3] for r in post_rows]
             }).to_csv('Data/cohort_a_psm_balance.csv', index=False, encoding='utf-8-sig')
print('  Data/cohort_a_psm_balance.csv')
